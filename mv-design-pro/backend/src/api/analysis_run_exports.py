from __future__ import annotations

import io
import json
import textwrap
from typing import Any, Literal, cast

from api.analysis_case_context import build_analysis_case_context
from api.canonical_run_views import (
    build_analysis_run_summary,
    build_automation_trace_results_response,
    build_branch_results_response,
    build_bus_results_response,
    build_dynamic_stability_results_response,
    build_extended_trace_response,
    build_phase_state_results_response,
    build_power_flow_export_bundle,
    build_results_index_response,
    build_short_circuit_results_response,
    build_source_compliance_results_response,
)
from api.v125_contracts import build_export_artifact, build_export_policy, resolve_proof_pack_ref
from application.analysis_run.read_model import build_trace_summary, canonicalize_json
from enm.canonical_analysis import CanonicalRun
from fastapi.responses import Response
from network_model.reporting.czcionki import zarejestruj_czcionki

ReportProfile = Literal["osd", "wykonawczy", "audytowy"]
ReportDetailLevel = Literal["minimalny", "standardowy", "pelny"]
ReportScope = Literal["whole_run", "active_table"]
ReportSection = Literal["summary", "results", "catalog", "trace"]
ReportFocusTable = (
    Literal[
        "buses",
        "branches",
        "short_circuit",
        "phase_state",
        "dynamic_stability",
        "automation_trace",
        "source_compliance",
        "trace",
    ]
    | None
)

DEFAULT_REPORT_SECTIONS_BY_DETAIL: dict[ReportDetailLevel, tuple[ReportSection, ...]] = {
    "minimalny": ("summary", "results"),
    "standardowy": ("summary", "results", "catalog"),
    "pelny": ("summary", "results", "catalog", "trace"),
}

REPORT_SECTION_LABELS: dict[ReportSection, str] = {
    "summary": "Podsumowanie",
    "results": "Wyniki tabelaryczne",
    "catalog": "Kontekst katalogowy",
    "trace": "Wywód szczegółowy",
}

REPORT_PROFILE_LABELS: dict[ReportProfile, str] = {
    "osd": "OSD",
    "wykonawczy": "Wykonawczy",
    "audytowy": "Audytowy",
}

REPORT_DETAIL_LABELS: dict[ReportDetailLevel, str] = {
    "minimalny": "Minimalny",
    "standardowy": "Standardowy",
    "pelny": "Pełny",
}

REPORT_SCOPE_LABELS: dict[ReportScope, str] = {
    "whole_run": "Cały model",
    "active_table": "Aktywna tabela",
}


def normalize_report_options(
    *,
    profile: str | None = None,
    detail_level: str | None = None,
    scope: str | None = None,
    sections: list[str] | tuple[str, ...] | None = None,
    focus_table: str | None = None,
) -> dict[str, Any]:
    normalized_profile: ReportProfile = (
        cast(ReportProfile, profile) if profile in REPORT_PROFILE_LABELS else "osd"
    )
    normalized_detail: ReportDetailLevel = (
        cast(ReportDetailLevel, detail_level)
        if detail_level in DEFAULT_REPORT_SECTIONS_BY_DETAIL
        else "standardowy"
    )
    normalized_scope: ReportScope = (
        cast(ReportScope, scope) if scope in REPORT_SCOPE_LABELS else "whole_run"
    )
    normalized_focus_table: ReportFocusTable = (
        cast(ReportFocusTable, focus_table)
        if focus_table
        in {
            "buses",
            "branches",
            "short_circuit",
            "phase_state",
            "dynamic_stability",
            "automation_trace",
            "source_compliance",
            "trace",
        }
        else None
    )

    default_sections = list(DEFAULT_REPORT_SECTIONS_BY_DETAIL[normalized_detail])
    raw_sections = list(sections or [])
    all_sections: tuple[ReportSection, ...] = ("summary", "results", "catalog", "trace")
    normalized_sections: list[ReportSection] = [
        section for section in all_sections if section in raw_sections
    ]
    if not normalized_sections:
        normalized_sections = default_sections

    return {
        "profile": normalized_profile,
        "profile_label": REPORT_PROFILE_LABELS[normalized_profile],
        "detail_level": normalized_detail,
        "detail_level_label": REPORT_DETAIL_LABELS[normalized_detail],
        "scope": normalized_scope,
        "scope_label": REPORT_SCOPE_LABELS[normalized_scope],
        "sections": normalized_sections,
        "section_labels": [REPORT_SECTION_LABELS[section] for section in normalized_sections],
        "focus_table": normalized_focus_table,
    }


def _analysis_title(run: CanonicalRun) -> str:
    if run.analysis_type == "PF":
        return "Raport rozpływu mocy"
    if run.analysis_type == "short_circuit_sn":
        return "Raport analizy zwarciowej"
    if run.analysis_type == "phase_state_sn":
        return "Raport stanu fazowego SN"
    if run.analysis_type == "dynamic_stability":
        return "Raport stabilności dynamicznej"
    if run.analysis_type == "source_compliance":
        return "Raport zgodności źródła"
    return "Raport analizy sieci"


def _detail_limits(detail_level: ReportDetailLevel) -> dict[str, int]:
    if detail_level == "minimalny":
        return {"catalog_items": 8, "trace_steps": 4, "rows": 12}
    if detail_level == "pelny":
        return {"catalog_items": 40, "trace_steps": 24, "rows": 80}
    return {"catalog_items": 20, "trace_steps": 12, "rows": 30}


def _build_report_results_section(
    run: CanonicalRun,
    *,
    scope: ReportScope,
    focus_table: ReportFocusTable,
) -> dict[str, Any]:
    results_index = build_results_index_response(run)
    bus_results = build_bus_results_response(run)
    branch_results = build_branch_results_response(run)
    short_circuit_results = build_short_circuit_results_response(run)
    phase_state_results = build_phase_state_results_response(run)
    dynamic_stability_results = build_dynamic_stability_results_response(run)
    automation_trace_results = build_automation_trace_results_response(run)
    source_compliance_results = build_source_compliance_results_response(run)

    if scope != "active_table" or focus_table is None:
        return {
            "index": results_index,
            "buses": bus_results,
            "branches": branch_results,
            "short_circuit": short_circuit_results,
            "phase_state": phase_state_results,
            "dynamic_stability": dynamic_stability_results,
            "automation_trace": automation_trace_results,
            "source_compliance": source_compliance_results,
        }

    filtered_tables = [
        table for table in results_index.get("tables", []) if table.get("table_id") == focus_table
    ]
    return {
        "index": {
            **results_index,
            "tables": filtered_tables,
        },
        "buses": bus_results if focus_table == "buses" else {"run_id": str(run.id), "rows": []},
        "branches": (
            branch_results if focus_table == "branches" else {"run_id": str(run.id), "rows": []}
        ),
        "short_circuit": (
            short_circuit_results
            if focus_table == "short_circuit"
            else {"run_id": str(run.id), "rows": []}
        ),
        "phase_state": (
            phase_state_results
            if focus_table == "phase_state"
            else {"run_id": str(run.id), "rows": []}
        ),
        "dynamic_stability": (
            dynamic_stability_results
            if focus_table == "dynamic_stability"
            else {"run_id": str(run.id), "rows": []}
        ),
        "automation_trace": (
            automation_trace_results
            if focus_table == "automation_trace"
            else {"run_id": str(run.id), "rows": []}
        ),
        "source_compliance": (
            source_compliance_results
            if focus_table == "source_compliance"
            else {"run_id": str(run.id), "rows": []}
        ),
    }


def build_analysis_run_report_payload(
    run: CanonicalRun,
    *,
    report_options: dict[str, Any] | None = None,
) -> dict[str, Any]:
    analysis_case_context = build_analysis_case_context(run)
    normalized_options = normalize_report_options(**(report_options or {}))
    trace_payload = canonicalize_json(build_extended_trace_response(run))
    results_section = _build_report_results_section(
        run,
        scope=normalized_options["scope"],
        focus_table=normalized_options["focus_table"],
    )

    payload: dict[str, Any] = {
        "report_type": "analysis_run_report",
        "report_version": "2.0.0",
        "title": _analysis_title(run),
        "report_options": normalized_options,
        "analysis_case_context": analysis_case_context,
        "proof_pack_ref": resolve_proof_pack_ref(run),
        "export_artifact": build_export_artifact(run, export_kind="json"),
        "export_policy": build_export_policy("json"),
        "run": build_analysis_run_summary(run),
        "summary": build_analysis_run_summary(run).get("summary_json", {}),
    }

    selected_sections = set(normalized_options["sections"])
    if "results" in selected_sections:
        payload["results"] = results_section
    if "catalog" in selected_sections:
        payload["catalog_context"] = trace_payload.get("catalog_context", [])
        payload["catalog_context_summary"] = trace_payload.get("catalog_context_summary", {})
    if "trace" in selected_sections:
        payload["trace"] = {
            "trace_summary": build_trace_summary(trace_payload.get("white_box_trace") or []),
            "white_box_trace": trace_payload.get("white_box_trace", []),
        }

    return payload


def _require_power_flow_bundle(run: CanonicalRun) -> dict[str, Any]:
    if run.analysis_type != "PF":
        raise ValueError(
            "Eksport raportu z kanonicznego obliczenia jest obecnie dostępny tylko dla rozpływu mocy.",
        )
    return build_power_flow_export_bundle(run)


def _build_short_circuit_export_bundle(run: CanonicalRun) -> dict[str, Any]:
    if run.analysis_type != "short_circuit_sn":
        raise ValueError("Przebieg nie jest analiza zwarciowa")
    analysis_case_context = build_analysis_case_context(run)
    trace_payload = canonicalize_json(build_extended_trace_response(run))
    return {
        "analysis_case_context": analysis_case_context,
        "short_circuit_results": build_short_circuit_results_response(run),
        "results_index": build_results_index_response(run),
        "catalog_context": trace_payload.get("catalog_context", []),
        "catalog_context_summary": trace_payload.get("catalog_context_summary", {}),
        "white_box_trace": trace_payload.get("white_box_trace", []),
        "metadata": {
            "run_id": str(run.id),
            "project_id": run.project_id,
            "study_case_id": run.case_id,
            "created_at": run.created_at.isoformat(),
            "input_hash": run.input_hash,
            "snapshot_hash": run.snapshot_hash,
            "analysis_type": (run.raw_result or {}).get("analysis_type"),
            "short_circuit_type": (run.raw_result or {}).get("short_circuit_type"),
            "reporting_status": (run.raw_result or {}).get("reporting_status"),
            "proof_status": (run.raw_result or {}).get("proof_status"),
            "proof_pack_ref": resolve_proof_pack_ref(run),
            "catalog_context_count": len(trace_payload.get("catalog_context", [])),
            "analysis_case_context": analysis_case_context,
        },
    }


def _build_generic_export_bundle(run: CanonicalRun) -> dict[str, Any]:
    analysis_case_context = build_analysis_case_context(run)
    trace_payload = canonicalize_json(build_extended_trace_response(run))
    results_index = build_results_index_response(run)
    bundle: dict[str, Any] = {
        "analysis_case_context": analysis_case_context,
        "results_index": results_index,
        "catalog_context": trace_payload.get("catalog_context", []),
        "catalog_context_summary": trace_payload.get("catalog_context_summary", {}),
        "white_box_trace": trace_payload.get("white_box_trace", []),
        "metadata": {
            "run_id": str(run.id),
            "project_id": run.project_id,
            "study_case_id": run.case_id,
            "created_at": run.created_at.isoformat(),
            "input_hash": run.input_hash,
            "snapshot_hash": run.snapshot_hash,
            "analysis_type": (run.raw_result or {}).get("analysis_type", run.analysis_type),
            "reporting_status": (run.raw_result or {}).get("reporting_status"),
            "proof_status": (run.raw_result or {}).get("proof_status"),
            "proof_pack_ref": resolve_proof_pack_ref(run),
            "catalog_context_count": len(trace_payload.get("catalog_context", [])),
            "analysis_case_context": analysis_case_context,
        },
    }
    if run.analysis_type == "phase_state_sn":
        bundle["phase_state"] = build_phase_state_results_response(run)
    elif run.analysis_type == "dynamic_stability":
        bundle["dynamic_stability"] = build_dynamic_stability_results_response(run)
        bundle["automation_trace"] = build_automation_trace_results_response(run)
    elif run.analysis_type == "source_compliance":
        bundle["source_compliance"] = build_source_compliance_results_response(run)
    return bundle


def _format_catalog_binding(entry: dict[str, Any]) -> str:
    binding = entry.get("catalog_binding") or {}
    namespace = binding.get("catalog_namespace") or "-"
    item_id = binding.get("catalog_item_id") or "-"
    version = binding.get("catalog_item_version")
    if version:
        return f"{namespace}:{item_id} ({version})"
    return f"{namespace}:{item_id}"


def _truncate_catalog_params(value: Any, max_chars: int = 220) -> str:
    if value is None:
        return "—"
    payload = canonicalize_json(value)
    text = str(payload)
    if isinstance(payload, dict | list):
        text = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    if len(text) <= max_chars:
        return text
    return f"{text[:max_chars]}..."


def _catalog_context_lines(bundle: dict[str, Any], *, max_items: int = 20) -> list[str]:
    lines: list[str] = []
    for entry in (bundle.get("catalog_context") or [])[:max_items]:
        lines.append(
            " | ".join(
                [
                    str(entry.get("element_id") or "-"),
                    str(entry.get("element_type") or "-"),
                    _format_catalog_binding(entry),
                    str(entry.get("parameter_source") or entry.get("parameter_origin") or "-"),
                    _truncate_catalog_params(entry.get("materialized_params")),
                ]
            )
        )
    return lines


def build_analysis_run_export_payload(run: CanonicalRun) -> dict[str, Any]:
    if run.analysis_type == "short_circuit_sn":
        bundle = _build_short_circuit_export_bundle(run)
        analysis_case_context = bundle.get("analysis_case_context") or build_analysis_case_context(
            run
        )
        return {
            "report_type": "short_circuit_result",
            "report_version": "1.1.0",
            "analysis_case_context": analysis_case_context,
            "proof_pack_ref": resolve_proof_pack_ref(run),
            "export_artifact": build_export_artifact(run, export_kind="json"),
            "export_policy": build_export_policy("json"),
            "metadata": bundle["metadata"],
            "short_circuit_results": bundle.get("short_circuit_results", {}),
            "results_index": bundle.get("results_index", {}),
            "catalog_context": bundle.get("catalog_context", []),
            "catalog_context_summary": bundle.get("catalog_context_summary", {}),
            "white_box_trace": bundle.get("white_box_trace", []),
            "trace_summary": build_trace_summary(bundle.get("white_box_trace") or []),
        }

    if run.analysis_type == "PF":
        bundle = _require_power_flow_bundle(run)
        analysis_case_context = bundle.get("analysis_case_context") or build_analysis_case_context(
            run
        )
        return {
            "report_type": "power_flow_result",
            "report_version": "1.1.0",
            "analysis_case_context": analysis_case_context,
            "proof_pack_ref": resolve_proof_pack_ref(run),
            "export_artifact": build_export_artifact(run, export_kind="json"),
            "export_policy": build_export_policy("json"),
            "metadata": bundle["metadata"],
            "result": bundle["result"],
            "bus_results": bundle.get("bus_results", {}),
            "branch_results": bundle.get("branch_results", {}),
            "results_index": bundle.get("results_index", {}),
            "catalog_context": bundle.get("catalog_context", []),
            "white_box_trace": bundle.get("white_box_trace", []),
            "trace_summary": {
                "solver_version": bundle["trace"].get("solver_version"),
                "input_hash": bundle["trace"].get("input_hash"),
                "converged": bundle["trace"].get("converged"),
                "final_iterations_count": bundle["trace"].get("final_iterations_count"),
            },
        }

    bundle = _build_generic_export_bundle(run)
    analysis_case_context = bundle.get("analysis_case_context") or build_analysis_case_context(run)
    payload = {
        "report_type": str(bundle["metadata"].get("analysis_type") or run.analysis_type),
        "report_version": "1.1.0",
        "analysis_case_context": analysis_case_context,
        "proof_pack_ref": resolve_proof_pack_ref(run),
        "export_artifact": build_export_artifact(run, export_kind="json"),
        "export_policy": build_export_policy("json"),
        "metadata": bundle["metadata"],
        "results_index": bundle.get("results_index", {}),
        "catalog_context": bundle.get("catalog_context", []),
        "catalog_context_summary": bundle.get("catalog_context_summary", {}),
        "white_box_trace": bundle.get("white_box_trace", []),
        "trace_summary": build_trace_summary(bundle.get("white_box_trace") or []),
    }
    if "phase_state" in bundle:
        payload["phase_state"] = bundle["phase_state"]
    if "dynamic_stability" in bundle:
        payload["dynamic_stability"] = bundle["dynamic_stability"]
    if "automation_trace" in bundle:
        payload["automation_trace"] = bundle["automation_trace"]
    if "source_compliance" in bundle:
        payload["source_compliance"] = bundle["source_compliance"]
    return payload


def build_analysis_run_trace_export_payload(run: CanonicalRun) -> dict[str, Any]:
    trace_payload = canonicalize_json(build_extended_trace_response(run))
    if not trace_payload.get("white_box_trace"):
        raise ValueError("Ślad obliczeniowy niedostępny dla tego obliczenia.")
    short_circuit_currents = _build_short_circuit_proof_currents(run)
    if short_circuit_currents:
        trace_payload["short_circuit_proof_currents"] = short_circuit_currents
    trace_payload["proof_pack_ref"] = resolve_proof_pack_ref(run)
    trace_payload["export_artifact"] = build_export_artifact(run, export_kind="whitebox_package")
    trace_payload["export_policy"] = build_export_policy("whitebox_package")
    return trace_payload


def _amps_to_ka(value: Any) -> float | None:
    if value is None:
        return None
    try:
        return float(value) / 1000.0
    except (TypeError, ValueError):
        return None


def _fmt_liczba(value: Any) -> str:
    """Deterministyczny format liczby do raportu: None -> em-dash, liczba -> %.6g.

    Uzywany zamiast wzorca ``value or "—"`` (ktory falszywie maskowal 0.0,
    np. kat 0.0 deg wezla bilansujacego) — uczciwe braki tylko dla None.
    """
    if value is None:
        return "—"
    if isinstance(value, int | float):
        return f"{float(value):.6g}"
    return str(value)


# Pelny bilans IEC 60909 punktu zwarcia w raportach PDF/DOCX — TEN SAM zakres co
# panel "Bilans IEC 60909" w UI (ZWARCIA-PRO F5, pkt 13 karty wlasciciela).
# Pola czytane 1:1 z wierszy kanonicznych `build_short_circuit_results`
# (enm/canonical_analysis) — zero fizyki w warstwie raportowej.
_SC_BILANS_POLA: tuple[tuple[str, str, str], ...] = (
    ("Rk", "rk_ohm", "Ohm"),
    ("Xk", "xk_ohm", "Ohm"),
    ("|Zk|", "zk_ohm", "Ohm"),
    ("X/R", "xr_ratio", ""),
    ("kappa", "kappa", ""),
    ("c", "c_factor", ""),
    ("Un", "un_kv", "kV"),
    ("tk", "tk_s", "s"),
    ("tb", "tb_s", "s"),
    ("I2t", "i2t_ka2s", "kA2s"),
)


def _krok_wywodu_linie(step: dict[str, Any], index: int) -> list[str]:
    """Krok wywodu jako linie tekstu (wzor LaTeX/podstawienie/wynik/uwagi).

    ZWARCIA-PRO F5 pkt 13: sekcja wywodu w raportach PDF/DOCX prezentuje kroki
    czytelnie (tekst kroku; LaTeX jako tekst wzoru — generator canvas/docx nie
    renderuje LaTeX), zamiast surowego zrzutu JSON. Format deterministyczny.
    """
    title = step.get("title") or step.get("key") or f"Krok {index}"
    lines = [f"{index}. {title}"]
    if step.get("formula_latex"):
        lines.append(f"Wzór: {step.get('formula_latex')}")
    substitution = step.get("substitution_latex") or step.get("substitution")
    if substitution:
        lines.append(f"Podstawienie: {substitution}")
    if step.get("result") is not None:
        lines.append(f"Wynik: {json.dumps(step.get('result'), ensure_ascii=False, sort_keys=True)}")
    if step.get("notes"):
        lines.append(f"Uwagi: {step.get('notes')}")
    return lines


def _sc_row_glowne_linia(row_data: dict[str, Any]) -> str:
    """Glowna linia wiersza zwarciowego: prady i moc zwarciowa (kanoniczne kA/MVA)."""
    return " | ".join(
        [
            str(row_data.get("target_name") or row_data.get("target_id") or "—"),
            f"Ik''={_fmt_liczba(row_data.get('ikss_ka'))} kA",
            f"ip={_fmt_liczba(row_data.get('ip_ka'))} kA",
            f"Ith={_fmt_liczba(row_data.get('ith_ka'))} kA",
            f"Ib={_fmt_liczba(row_data.get('ib_ka'))} kA",
            f"Ik={_fmt_liczba(row_data.get('ik_ka'))} kA",
            f"Sk''={_fmt_liczba(row_data.get('sk_mva'))} MVA",
        ]
    )


def _sc_row_bilans_linie(row_data: dict[str, Any]) -> tuple[str, str]:
    """Dwie linie pelnego bilansu IEC 60909 wiersza zwarciowego (kolumny addytywne).

    Podzial staly (impedancje / wielkosci normowe), zeby linia PDF nie przekraczala
    limitu szerokosci strony — format deterministyczny.
    """

    def _czesc(pola: tuple[tuple[str, str, str], ...]) -> str:
        parts = []
        for label, key, unit in pola:
            value = _fmt_liczba(row_data.get(key))
            parts.append(f"{label}={value} {unit}".rstrip())
        return " | ".join(parts)

    return (
        "Bilans IEC 60909: " + _czesc(_SC_BILANS_POLA[:5]),
        _czesc(_SC_BILANS_POLA[5:]),
    )


def _build_short_circuit_proof_currents(run: CanonicalRun) -> dict[str, Any] | None:
    """Expose SC3F I_dyn/I_th proof aliases from solver result fields."""
    if run.analysis_type != "short_circuit_sn":
        return None
    raw_result = run.raw_result or {}
    if raw_result.get("short_circuit_type") != "3F":
        return None
    graph_nodes = (raw_result.get("graph") or {}).get("nodes", {})
    rows: list[dict[str, Any]] = []
    for item in raw_result.get("results", []):
        target_id = item.get("fault_node_id")
        node = graph_nodes.get(target_id, {}) if isinstance(graph_nodes, dict) else {}
        rows.append(
            {
                "target_id": target_id,
                "element_id": node.get("element_id") or target_id,
                "target_name": node.get("name") or node.get("element_id") or target_id,
                "fault_type": item.get("short_circuit_type")
                or raw_result.get("short_circuit_type"),
                "I_dyn": {
                    "symbol": "I_dyn",
                    "label_pl": "Prad dynamiczny do sprawdzenia aparatury",
                    "value_ka": _amps_to_ka(item.get("ip_a")),
                    "source_field": "ip_a",
                    "source_standard": "IEC 60909",
                },
                "I_th": {
                    "symbol": "I_th",
                    "label_pl": "Prad cieplny zastepczy do sprawdzenia aparatury",
                    "value_ka": _amps_to_ka(item.get("ith_a")),
                    "source_field": "ith_a",
                    "source_standard": "IEC 60909",
                },
                "proof_ref": item.get("proof_ref"),
            }
        )
    if not rows:
        return None
    rows.sort(key=lambda row: str(row.get("target_id") or ""))
    return {
        "standard_basis": "IEC 60909",
        "scope": "SC3F",
        "symbols": {
            "I_dyn": "Prad dynamiczny porownywany z wytrzymaloscia dynamiczna aparatury.",
            "I_th": "Prad cieplny zastepczy porownywany z wytrzymaloscia cieplna aparatury.",
        },
        "rows": rows,
    }


def export_run_report_json_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
    report_options: dict[str, Any] | None = None,
) -> Response:
    json_content = json.dumps(
        build_analysis_run_report_payload(run, report_options=report_options),
        indent=2,
        ensure_ascii=False,
        sort_keys=True,
    )
    return Response(
        content=json_content,
        media_type="application/json",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.json"',
        },
    )


def _trace_jsonl_lines(trace_payload: dict[str, Any]) -> list[str]:
    exported_at = canonicalize_json(trace_payload.get("exported_at")) or None
    if not isinstance(exported_at, str):
        from datetime import UTC, datetime

        exported_at = datetime.now(UTC).isoformat()

    white_box_trace = trace_payload.get("white_box_trace") or []
    lines = [
        json.dumps(
            {
                "type": "header",
                "seq": 1,
                "exported_at": exported_at,
                "data": {
                    "run_id": trace_payload.get("run_id"),
                    "snapshot_id": trace_payload.get("snapshot_id"),
                    "input_hash": trace_payload.get("input_hash"),
                    "total_steps": len(white_box_trace),
                    "catalog_context_count": len(trace_payload.get("catalog_context") or []),
                    "catalog_context": trace_payload.get("catalog_context") or [],
                    "catalog_context_by_element": trace_payload.get("catalog_context_by_element"),
                    "catalog_context_summary": trace_payload.get("catalog_context_summary"),
                    "analysis_case_context": trace_payload.get("analysis_case_context"),
                    "proof_pack_ref": trace_payload.get("proof_pack_ref"),
                    "export_artifact": trace_payload.get("export_artifact"),
                    "export_policy": trace_payload.get("export_policy"),
                    "export_version": "1.0.0",
                },
            },
            ensure_ascii=False,
        )
    ]

    for index, step in enumerate(white_box_trace):
        lines.append(
            json.dumps(
                {
                    "type": "step",
                    "seq": index + 2,
                    "exported_at": exported_at,
                    "data": {
                        "step_index": index,
                        "step_number": step.get("step"),
                        "key": step.get("key") or step.get("step_id"),
                        "title": step.get("title") or step.get("description"),
                        "phase": step.get("phase"),
                        "element_id": step.get("element_id")
                        or (step.get("catalog_context_entry") or {}).get("element_id"),
                        "target_id": step.get("target_id"),
                        "solver_ref": step.get("solver_ref"),
                        "catalog_binding": step.get("catalog_binding")
                        or (step.get("catalog_context_entry") or {}).get("catalog_binding"),
                        "source_catalog": step.get("source_catalog")
                        or (step.get("catalog_context_entry") or {}).get("source_catalog"),
                        "source_catalog_label": step.get("source_catalog_label")
                        or (step.get("catalog_context_entry") or {}).get("source_catalog_label"),
                        "parameter_origin": step.get("parameter_origin")
                        or (step.get("catalog_context_entry") or {}).get("parameter_origin"),
                        "parameter_source": step.get("parameter_source")
                        or (step.get("catalog_context_entry") or {}).get("parameter_source"),
                        "formula_latex": step.get("formula_latex"),
                        "inputs": step.get("inputs"),
                        "substitution": step.get("substitution"),
                        "result": step.get("result"),
                        "proof_ref": step.get("proof_ref"),
                        "proof_status": step.get("proof_status"),
                        "reporting_status": step.get("reporting_status"),
                        "method_basis": step.get("method_basis"),
                        "materialized_params": step.get("materialized_params")
                        or (step.get("catalog_context_entry") or {}).get("materialized_params"),
                        "manual_overrides": step.get("manual_overrides")
                        or (step.get("catalog_context_entry") or {}).get("manual_overrides"),
                        "manual_override_count": step.get("manual_override_count")
                        or (step.get("catalog_context_entry") or {}).get("manual_override_count"),
                        "notes": step.get("notes"),
                    },
                },
                ensure_ascii=False,
            )
        )

    return lines


def export_run_json_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
) -> Response:
    json_content = json.dumps(
        build_analysis_run_export_payload(run),
        indent=2,
        ensure_ascii=False,
        sort_keys=True,
    )
    return Response(
        content=json_content,
        media_type="application/json",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.json"',
        },
    )


def export_run_docx_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
) -> Response:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise ValueError(
            "Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx",
        ) from exc

    bundle = _require_power_flow_bundle(run)
    result = bundle["result"]
    metadata = bundle["metadata"]
    catalog_context_lines = _catalog_context_lines(bundle)
    white_box_trace = bundle.get("white_box_trace") or []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading("Raport rozpływu mocy", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_parts = [
        f"Status: {'Zbieżny' if result.get('converged') else 'Niezbieżny'}",
        f"Iteracje: {result.get('iterations_count', '—')}",
        f"Uruchomienie: {str(metadata.get('run_id', '—'))[:8]}...",
    ]
    subtitle = doc.add_paragraph(" | ".join(status_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Podsumowanie", level=1)
    summary = result.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartość"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run_obj in paragraph.runs:
                run_obj.bold = True

    def add_row(label: str, value: Any) -> None:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value is not None else "—"

    add_row("Status zbieżności", "Zbieżny" if result.get("converged") else "Niezbieżny")
    add_row("Liczba iteracji", result.get("iterations_count"))
    add_row("Węzeł bilansujący", result.get("slack_bus_id"))
    add_row("Całkowite straty P [MW]", f"{summary.get('total_losses_p_mw', 0):.4g}")
    add_row("Całkowite straty Q [Mvar]", f"{summary.get('total_losses_q_mvar', 0):.4g}")
    add_row("Min. napięcie [pu]", f"{summary.get('min_v_pu', 0):.4g}")
    add_row("Max. napięcie [pu]", f"{summary.get('max_v_pu', 0):.4g}")
    add_row("Elementy z katalogiem", metadata.get("catalog_context_count"))

    doc.add_paragraph()
    doc.add_heading("Kontekst katalogowy", level=1)
    if catalog_context_lines:
        doc.add_paragraph(
            "Format: element_id | typ | katalog | pochodzenie parametrów | materialized_params"
        )
        for line in catalog_context_lines:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("Brak jawnego kontekstu katalogowego.")

    doc.add_paragraph()
    doc.add_heading("Wywód szczegółowy", level=1)
    if white_box_trace:
        for step in white_box_trace[:12]:
            title = step.get("title") or step.get("key") or "Krok"
            doc.add_paragraph(f"{title}: {json.dumps(step, ensure_ascii=False, sort_keys=True)}")
        if len(white_box_trace) > 12:
            doc.add_paragraph(f"... oraz {len(white_box_trace) - 12} kolejnych kroków")
    else:
        doc.add_paragraph("Brak jawnego śladu obliczeń.")

    doc.add_paragraph()
    doc.add_heading("Wyniki węzłowe (szyny)", level=1)
    bus_results = result.get("bus_results", [])
    if bus_results:
        bus_table = doc.add_table(rows=1, cols=5)
        bus_table.style = "Table Grid"
        header = bus_table.rows[0].cells
        header[0].text = "ID szyny"
        header[1].text = "V [pu]"
        header[2].text = "Kąt [deg]"
        header[3].text = "P_inj [MW]"
        header[4].text = "Q_inj [Mvar]"
        for cell in header:
            for paragraph in cell.paragraphs:
                for run_obj in paragraph.runs:
                    run_obj.bold = True
        for bus in bus_results[:30]:
            row = bus_table.add_row().cells
            row[0].text = str(bus.get("bus_id", "—"))[:16]
            row[1].text = f"{bus.get('v_pu', 0):.4g}"
            row[2].text = f"{bus.get('angle_deg', 0):.2f}"
            row[3].text = f"{bus.get('p_injected_mw', 0):.3g}"
            row[4].text = f"{bus.get('q_injected_mvar', 0):.3g}"
        if len(bus_results) > 30:
            doc.add_paragraph(f"... oraz {len(bus_results) - 30} dodatkowych węzłów")
    else:
        doc.add_paragraph("Brak wyników węzłowych.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.docx"',
        },
    )


def export_run_pdf_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
) -> Response:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ValueError(
            "Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab",
        ) from exc

    bundle = _require_power_flow_bundle(run)
    result = bundle["result"]
    metadata = bundle["metadata"]
    catalog_context_lines = _catalog_context_lines(bundle)
    white_box_trace = bundle.get("white_box_trace") or []

    buffer = io.BytesIO()
    zarejestruj_czcionki()
    canvas_obj = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 25 * mm
    top_margin = page_height - 25 * mm
    y = top_margin
    line_height = 5 * mm

    canvas_obj.setFont("DejaVuSans-Bold", 16)
    title = "Raport rozpływu mocy"
    canvas_obj.drawString(
        (page_width - canvas_obj.stringWidth(title, "DejaVuSans-Bold", 16)) / 2,
        y,
        title,
    )
    y -= 10 * mm

    canvas_obj.setFont("DejaVuSans", 10)
    status_text = (
        f"Status: {'Zbieżny' if result.get('converged') else 'Niezbieżny'} | "
        f"Iteracje: {result.get('iterations_count', '—')} | "
        f"Uruchomienie: {str(metadata.get('run_id', '—'))[:8]}..."
    )
    canvas_obj.drawString(left_margin, y, status_text)
    y -= 8 * mm

    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Podsumowanie")
    y -= 6 * mm

    canvas_obj.setFont("DejaVuSans", 10)
    summary = result.get("summary", {})
    summary_lines = [
        f"Węzeł bilansujący: {result.get('slack_bus_id', '—')}",
        f"Całkowite straty P: {summary.get('total_losses_p_mw', 0):.4g} MW",
        f"Całkowite straty Q: {summary.get('total_losses_q_mvar', 0):.4g} Mvar",
        f"Min. napięcie: {summary.get('min_v_pu', 0):.4g} pu",
        f"Max. napięcie: {summary.get('max_v_pu', 0):.4g} pu",
        f"Elementy z katalogiem: {metadata.get('catalog_context_count', 0)}",
    ]
    for line in summary_lines:
        canvas_obj.drawString(left_margin, y, line)
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Kontekst katalogowy")
    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans", 8)
    if catalog_context_lines:
        canvas_obj.drawString(
            left_margin,
            y,
            "Format: element_id | typ | katalog | pochodzenie | materialized_params",
        )
        y -= line_height
        for line in catalog_context_lines:
            canvas_obj.drawString(left_margin, y, line[:160])
            y -= line_height
            if y < 30 * mm:
                canvas_obj.showPage()
                y = top_margin
                canvas_obj.setFont("DejaVuSans", 8)
    else:
        canvas_obj.drawString(left_margin, y, "Brak jawnego kontekstu katalogowego.")
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Wywód szczegółowy")
    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans", 8)
    if white_box_trace:
        for step in white_box_trace[:10]:
            title = step.get("title") or step.get("key") or "Krok"
            canvas_obj.drawString(
                left_margin,
                y,
                f"{title}: {json.dumps(step, ensure_ascii=False, sort_keys=True)[:150]}",
            )
            y -= line_height
            if y < 30 * mm:
                canvas_obj.showPage()
                y = top_margin
                canvas_obj.setFont("DejaVuSans", 8)
        if len(white_box_trace) > 10:
            canvas_obj.drawString(
                left_margin,
                y,
                f"... oraz {len(white_box_trace) - 10} kolejnych kroków",
            )
            y -= line_height
    else:
        canvas_obj.drawString(left_margin, y, "Brak jawnego śladu obliczeń.")
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Wyniki węzłowe (top 20)")
    y -= 5 * mm

    canvas_obj.setFont("DejaVuSans", 9)
    for bus in result.get("bus_results", [])[:20]:
        text = (
            f"{str(bus.get('bus_id', '—'))[:12]}: "
            f"V={bus.get('v_pu', 0):.4g} pu, "
            f"kat={bus.get('angle_deg', 0):.2f} deg"
        )
        canvas_obj.drawString(left_margin, y, text)
        y -= line_height
        if y < 30 * mm:
            canvas_obj.showPage()
            y = top_margin

    canvas_obj.save()
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.pdf"',
        },
    )


def export_run_report_docx_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
    report_options: dict[str, Any] | None = None,
) -> Response:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise ValueError(
            "Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx",
        ) from exc

    payload = build_analysis_run_report_payload(run, report_options=report_options)
    options = payload["report_options"]
    limits = _detail_limits(options["detail_level"])
    results_section = payload.get("results", {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading(payload["title"], level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        " | ".join(
            [
                f"Profil: {options['profile_label']}",
                f"Poziom: {options['detail_level_label']}",
                f"Zakres: {options['scope_label']}",
                f"Uruchomienie: {str(run.id)[:8]}...",
            ]
        )
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Parametry raportu", level=1)
    options_table = doc.add_table(rows=1, cols=2)
    options_table.style = "Table Grid"
    header = options_table.rows[0].cells
    header[0].text = "Parametr"
    header[1].text = "Wartość"

    def add_pair(label: str, value: Any) -> None:
        row = options_table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value is not None else "—"

    add_pair("Typ analizy", run.analysis_type)
    add_pair("Status obliczenia", run.status)
    add_pair("Status wyników", run.result_status)
    add_pair("Profil raportu", options["profile_label"])
    add_pair("Poziom szczegółowości", options["detail_level_label"])
    add_pair("Zakres", options["scope_label"])
    add_pair("Sekcje", ", ".join(options["section_labels"]))

    if "summary" in options["sections"]:
        doc.add_paragraph()
        doc.add_heading("Podsumowanie", level=1)
        summary = payload.get("summary", {})
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = "Table Grid"
        summary_header = summary_table.rows[0].cells
        summary_header[0].text = "Pole"
        summary_header[1].text = "Wartość"
        for key, value in summary.items():
            row = summary_table.add_row().cells
            row[0].text = str(key)
            row[1].text = (
                json.dumps(value, ensure_ascii=False)
                if isinstance(value, dict | list)
                else str(value)
            )

    if "results" in options["sections"]:
        doc.add_paragraph()
        doc.add_heading("Wyniki tabelaryczne", level=1)
        for table in results_section.get("index", {}).get("tables", []):
            table_id = table.get("table_id")
            label = table.get("label_pl") or table_id or "Tabela"
            doc.add_heading(str(label), level=2)
            if table_id == "buses":
                rows = (results_section.get("buses", {}) or {}).get("rows", [])[: limits["rows"]]
                for row_data in rows:
                    doc.add_paragraph(
                        " | ".join(
                            [
                                str(row_data.get("name") or row_data.get("bus_id") or "—"),
                                f"U={_fmt_liczba(row_data.get('u_pu'))} pu",
                                f"kąt={_fmt_liczba(row_data.get('angle_deg'))} deg",
                            ]
                        )
                    )
            elif table_id == "branches":
                rows = (results_section.get("branches", {}) or {}).get("rows", [])[: limits["rows"]]
                for row_data in rows:
                    doc.add_paragraph(
                        " | ".join(
                            [
                                str(row_data.get("name") or row_data.get("branch_id") or "—"),
                                f"I={row_data.get('i_a') or '—'} A",
                                f"P={row_data.get('p_mw') or '—'} MW",
                                f"Q={row_data.get('q_mvar') or '—'} Mvar",
                            ]
                        )
                    )
            elif table_id == "short_circuit":
                rows = (results_section.get("short_circuit", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]
                for row_data in rows:
                    # ZWARCIA-PRO F5: pelny bilans IEC 60909 1:1 z wierszy
                    # kanonicznych (TEN SAM zakres co UI) — trzy linie na punkt.
                    doc.add_paragraph(_sc_row_glowne_linia(row_data))
                    bilans_1, bilans_2 = _sc_row_bilans_linie(row_data)
                    doc.add_paragraph(bilans_1)
                    doc.add_paragraph(bilans_2)
            elif table_id == "phase_state":
                rows = (results_section.get("phase_state", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]
                for row_data in rows:
                    doc.add_paragraph(
                        " | ".join(
                            [
                                str(
                                    row_data.get("target_name") or row_data.get("element_id") or "—"
                                ),
                                f"UA={row_data.get('ua_kv') or '—'} kV",
                                f"UB={row_data.get('ub_kv') or '—'} kV",
                                f"UC={row_data.get('uc_kv') or '—'} kV",
                                f"Asymetria U={row_data.get('voltage_unbalance_percent') or '—'} %",
                            ]
                        )
                    )
            elif table_id == "dynamic_stability":
                rows = (results_section.get("dynamic_stability", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]
                for row_data in rows:
                    doc.add_paragraph(
                        " | ".join(
                            [
                                str(row_data.get("source_id") or "—"),
                                f"Status={row_data.get('status') or '—'}",
                                f"t_wyl={row_data.get('clearing_time_ms') or '—'} ms",
                                f"Margines={row_data.get('clearing_margin_ms') or '—'} ms",
                                f"Indeks={row_data.get('stability_index') or '—'}",
                            ]
                        )
                    )
            elif table_id == "automation_trace":
                rows = (results_section.get("automation_trace", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]
                for row_data in rows:
                    doc.add_paragraph(
                        " | ".join(
                            [
                                str(row_data.get("event_seq") or "—"),
                                str(row_data.get("event_type") or "—"),
                                str(row_data.get("element_id") or "—"),
                                str(row_data.get("detail") or "—"),
                            ]
                        )
                    )
            elif table_id == "source_compliance":
                rows = (results_section.get("source_compliance", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]
                for row_data in rows:
                    doc.add_paragraph(
                        " | ".join(
                            [
                                str(row_data.get("source_ref") or "—"),
                                str(row_data.get("source_type") or "—"),
                                str(row_data.get("verdict") or "—"),
                                str(row_data.get("reporting_status") or "—"),
                            ]
                        )
                    )
            else:
                doc.add_paragraph("Brak danych tabelarycznych dla wybranego zakresu.")

    if "catalog" in options["sections"]:
        doc.add_paragraph()
        doc.add_heading("Kontekst katalogowy", level=1)
        catalog_context = payload.get("catalog_context", [])
        if catalog_context:
            for entry in catalog_context[: limits["catalog_items"]]:
                doc.add_paragraph(
                    " | ".join(
                        [
                            str(entry.get("element_id") or "—"),
                            str(entry.get("element_type") or "—"),
                            _format_catalog_binding(entry),
                            str(
                                entry.get("parameter_source")
                                or entry.get("parameter_origin")
                                or "—"
                            ),
                        ]
                    )
                )
        else:
            doc.add_paragraph("Brak jawnego kontekstu katalogowego.")

    if "trace" in options["sections"]:
        doc.add_paragraph()
        doc.add_heading("Wywód szczegółowy", level=1)
        white_box_trace = (payload.get("trace", {}) or {}).get("white_box_trace", [])
        if white_box_trace:
            for index, step in enumerate(white_box_trace[: limits["trace_steps"]], start=1):
                for line in _krok_wywodu_linie(step, index):
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph("Brak jawnego śladu obliczeń.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.docx"',
        },
    )


def export_run_report_pdf_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
    report_options: dict[str, Any] | None = None,
) -> Response:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ValueError(
            "Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab",
        ) from exc

    payload = build_analysis_run_report_payload(run, report_options=report_options)
    options = payload["report_options"]
    limits = _detail_limits(options["detail_level"])
    results_section = payload.get("results", {})

    buffer = io.BytesIO()
    zarejestruj_czcionki()
    canvas_obj = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 25 * mm
    top_margin = page_height - 25 * mm
    bottom_margin = 18 * mm
    y = top_margin
    line_height = 5 * mm

    def ensure_page(required_lines: int = 1) -> None:
        nonlocal y
        if y - (required_lines * line_height) < bottom_margin:
            canvas_obj.showPage()
            y = top_margin

    def draw_line(text: str, *, font: str = "DejaVuSans", size: int = 9) -> None:
        nonlocal y
        ensure_page(1)
        canvas_obj.setFont(font, size)
        canvas_obj.drawString(left_margin, y, text[:155])
        y -= line_height

    canvas_obj.setFont("DejaVuSans-Bold", 16)
    canvas_obj.drawString(left_margin, y, payload["title"])
    y -= 8 * mm

    draw_line(
        f"Profil: {options['profile_label']} | Poziom: {options['detail_level_label']} | Zakres: {options['scope_label']}",
        size=10,
    )
    draw_line(
        f"Uruchomienie: {run.id} | Analiza: {run.analysis_type} | Status: {run.status}",
        size=10,
    )
    y -= 2 * mm

    if "summary" in options["sections"]:
        canvas_obj.setFont("DejaVuSans-Bold", 12)
        canvas_obj.drawString(left_margin, y, "Podsumowanie")
        y -= line_height
        for key, value in (payload.get("summary", {}) or {}).items():
            draw_line(f"{key}: {value}")
        y -= 2 * mm

    if "results" in options["sections"]:
        canvas_obj.setFont("DejaVuSans-Bold", 12)
        canvas_obj.drawString(left_margin, y, "Wyniki tabelaryczne")
        y -= line_height
        for table in results_section.get("index", {}).get("tables", []):
            draw_line(
                str(table.get("label_pl") or table.get("table_id") or "Tabela"),
                font="DejaVuSans-Bold",
                size=10,
            )
            if table.get("table_id") == "buses":
                for row_data in (results_section.get("buses", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]:
                    draw_line(
                        f"{row_data.get('name') or row_data.get('bus_id')}: U={_fmt_liczba(row_data.get('u_pu'))} pu, kąt={_fmt_liczba(row_data.get('angle_deg'))} deg"
                    )
            elif table.get("table_id") == "branches":
                for row_data in (results_section.get("branches", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]:
                    draw_line(
                        f"{row_data.get('name') or row_data.get('branch_id')}: I={row_data.get('i_a') or '—'} A, P={row_data.get('p_mw') or '—'} MW"
                    )
            elif table.get("table_id") == "short_circuit":
                for row_data in (results_section.get("short_circuit", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]:
                    # ZWARCIA-PRO F5: pelny bilans IEC 60909 1:1 z wierszy
                    # kanonicznych (TEN SAM zakres co UI) — trzy linie na punkt.
                    draw_line(_sc_row_glowne_linia(row_data))
                    bilans_1, bilans_2 = _sc_row_bilans_linie(row_data)
                    draw_line(bilans_1, size=8)
                    draw_line(bilans_2, size=8)
            elif table.get("table_id") == "phase_state":
                for row_data in (results_section.get("phase_state", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]:
                    draw_line(
                        f"{row_data.get('target_name') or row_data.get('element_id')}: UA={row_data.get('ua_kv') or '—'} kV, UB={row_data.get('ub_kv') or '—'} kV, UC={row_data.get('uc_kv') or '—'} kV"
                    )
            elif table.get("table_id") == "dynamic_stability":
                for row_data in (results_section.get("dynamic_stability", {}) or {}).get(
                    "rows", []
                )[: limits["rows"]]:
                    draw_line(
                        f"{row_data.get('source_id')}: status={row_data.get('status') or '—'}, t_wyl={row_data.get('clearing_time_ms') or '—'} ms, indeks={row_data.get('stability_index') or '—'}"
                    )
            elif table.get("table_id") == "automation_trace":
                for row_data in (results_section.get("automation_trace", {}) or {}).get("rows", [])[
                    : limits["rows"]
                ]:
                    draw_line(
                        f"{row_data.get('event_seq') or '—'} | {row_data.get('event_type') or '—'} | {row_data.get('element_id') or '—'} | {row_data.get('detail') or '—'}"
                    )
            elif table.get("table_id") == "source_compliance":
                for row_data in (results_section.get("source_compliance", {}) or {}).get(
                    "rows", []
                )[: limits["rows"]]:
                    draw_line(
                        f"{row_data.get('source_ref') or '—'} | {row_data.get('source_type') or '—'} | {row_data.get('verdict') or '—'} | {row_data.get('reporting_status') or '—'}"
                    )
        y -= 2 * mm

    if "catalog" in options["sections"]:
        canvas_obj.setFont("DejaVuSans-Bold", 12)
        canvas_obj.drawString(left_margin, y, "Kontekst katalogowy")
        y -= line_height
        catalog_context = payload.get("catalog_context", [])
        if catalog_context:
            for entry in catalog_context[: limits["catalog_items"]]:
                draw_line(
                    " | ".join(
                        [
                            str(entry.get("element_id") or "—"),
                            str(entry.get("element_type") or "—"),
                            _format_catalog_binding(entry),
                        ]
                    ),
                    size=8,
                )
        else:
            draw_line("Brak jawnego kontekstu katalogowego.")
        y -= 2 * mm

    if "trace" in options["sections"]:
        canvas_obj.setFont("DejaVuSans-Bold", 12)
        canvas_obj.drawString(left_margin, y, "Wywód szczegółowy")
        y -= line_height
        white_box_trace = (payload.get("trace", {}) or {}).get("white_box_trace", [])
        if white_box_trace:
            for index, step in enumerate(white_box_trace[: limits["trace_steps"]], start=1):
                for line in _krok_wywodu_linie(step, index):
                    draw_line(line, size=8)
        else:
            draw_line("Brak jawnego śladu obliczeń.")

    canvas_obj.save()
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.pdf"',
        },
    )


def export_run_trace_jsonl_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
) -> Response:
    lines = _trace_jsonl_lines(build_analysis_run_trace_export_payload(run))
    return Response(
        content="\n".join(lines),
        media_type="application/x-ndjson",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.jsonl"',
        },
    )


def export_run_trace_pdf_response(
    run: CanonicalRun,
    *,
    filename_stem: str,
) -> Response:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ValueError(
            "Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab",
        ) from exc

    trace_payload = build_analysis_run_trace_export_payload(run)
    summary = build_trace_summary(trace_payload.get("white_box_trace") or [])
    buffer = io.BytesIO()
    zarejestruj_czcionki()
    canvas_obj = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 20 * mm
    right_margin = page_width - 20 * mm
    top_margin = page_height - 20 * mm
    bottom_margin = 18 * mm
    y = top_margin

    def ensure_page(required_lines: int = 1, line_height: float = 4.5 * mm) -> None:
        nonlocal y
        if y - (required_lines * line_height) < bottom_margin:
            canvas_obj.showPage()
            y = top_margin

    def draw_wrapped(
        text: str,
        *,
        font_name: str = "DejaVuSans",
        font_size: int = 9,
        max_chars: int = 110,
        line_height: float = 4.5 * mm,
    ) -> None:
        nonlocal y
        canvas_obj.setFont(font_name, font_size)
        lines = textwrap.wrap(text, width=max_chars) or [text]
        ensure_page(len(lines), line_height)
        for line in lines:
            canvas_obj.drawString(left_margin, y, line)
            y -= line_height

    canvas_obj.setFont("DejaVuSans-Bold", 15)
    title = "Wywód obliczeń"
    canvas_obj.drawString(
        (page_width - canvas_obj.stringWidth(title, "DejaVuSans-Bold", 15)) / 2,
        y,
        title,
    )
    y -= 7 * mm

    draw_wrapped(
        f"Uruchomienie: {trace_payload.get('run_id')}", font_name="DejaVuSans", font_size=10
    )
    draw_wrapped(
        f"Wersja modelu: {trace_payload.get('snapshot_id') or '—'} | Hash wejścia: {trace_payload.get('input_hash') or '—'}",
        font_name="DejaVuSans",
        font_size=9,
    )
    draw_wrapped(
        f"Liczba kroków: {summary.get('count', 0)} | Fazy: {', '.join(summary.get('phases', []) or []) or '—'}",
        font_name="DejaVuSans",
        font_size=9,
    )
    y -= 3 * mm

    white_box_trace = trace_payload.get("white_box_trace") or []
    for index, step in enumerate(white_box_trace, start=1):
        ensure_page(6)
        canvas_obj.setFont("DejaVuSans-Bold", 11)
        heading = step.get("title") or step.get("key") or f"Krok {index}"
        canvas_obj.drawString(left_margin, y, f"{index}. {heading}")
        y -= 5 * mm

        meta = " | ".join(
            [
                f"Faza: {step.get('phase') or '—'}",
                f"Element: {step.get('element_id') or '—'}",
                f"Cel: {step.get('target_id') or '—'}",
            ]
        )
        draw_wrapped(meta, font_name="DejaVuSans", font_size=8, max_chars=120, line_height=4 * mm)

        if step.get("formula_latex"):
            draw_wrapped(
                f"Wzór: {step.get('formula_latex')}",
                font_name="DejaVuSans-Oblique",
                font_size=8,
                max_chars=110,
                line_height=4 * mm,
            )
        if step.get("substitution"):
            draw_wrapped(
                f"Podstawienie: {step.get('substitution')}",
                font_name="DejaVuSans",
                font_size=8,
                max_chars=110,
                line_height=4 * mm,
            )
        if step.get("result") is not None:
            draw_wrapped(
                f"Wynik: {json.dumps(step.get('result'), ensure_ascii=False, sort_keys=True)}",
                font_name="DejaVuSans",
                font_size=8,
                max_chars=110,
                line_height=4 * mm,
            )
        if step.get("notes"):
            draw_wrapped(
                f"Uwagi: {step.get('notes')}",
                font_name="DejaVuSans",
                font_size=8,
                max_chars=110,
                line_height=4 * mm,
            )
        y -= 2 * mm

    ensure_page(2)
    canvas_obj.setFont("DejaVuSans", 8)
    canvas_obj.drawRightString(
        right_margin, bottom_margin - 2 * mm, "Wygenerowano przez MV-DESIGN-PRO"
    )
    canvas_obj.save()
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename_stem}_{run.id}.pdf"',
        },
    )
