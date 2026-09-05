"""Canonical-only API endpoints for power flow runs."""

from __future__ import annotations

import io
import math
import zipfile
from collections.abc import Callable
from typing import Any
from uuid import UUID
from xml.sax.saxutils import escape

from api.analysis_run_exports import export_run_json_response
from api.canonical_run_views import (
    build_power_flow_export_bundle,
    build_power_flow_interpretation,
    build_power_flow_run_header,
)
from api.canonical_run_views import (
    get_power_flow_result as get_canonical_power_flow_result,
)
from api.canonical_run_views import (
    get_power_flow_trace as get_canonical_power_flow_trace,
)
from api.dependencies import get_uow_factory
from api.klucz_twin_dep import klucz_twin_z_uow
from application.analysis_run.read_model import canonicalize_json
from enm.canonical_analysis import CanonicalRun
from enm.canonical_analysis import create_run as create_canonical_run
from enm.canonical_analysis import execute_run as execute_canonical_run
from enm.canonical_analysis import get_run as get_canonical_run
from enm.canonical_analysis import list_runs_for_project as list_canonical_runs_for_project
from fastapi import APIRouter, Depends, HTTPException, Query, Response, status
from infrastructure.persistence.unit_of_work import UnitOfWork
from network_model.reporting.czcionki import zarejestruj_czcionki
from network_model.reporting.missing_value import format_wynik
from pydantic import BaseModel, Field

router = APIRouter(tags=["power-flow"])


class PowerFlowRunCreateRequest(BaseModel):
    study_case_id: UUID | None = Field(
        default=None,
        description="ID przypadku obliczeniowego. Jeżeli None, użyje aktywnego przypadku.",
    )

    options: dict[str, Any] | None = Field(
        default=None,
        description="Opcje solvera (tolerance, max_iter, trace_level, etc.)",
    )


class PowerFlowRunResponse(BaseModel):
    id: str
    deterministic_id: str
    project_id: str
    study_case_id: str
    analysis_type: str
    status: str
    result_status: str
    created_at: str
    started_at: str | None
    finished_at: str | None
    input_hash: str
    converged: bool | None = None
    iterations: int | None = None


class PowerFlowExecuteResponse(BaseModel):
    id: str
    status: str
    converged: bool | None = None
    iterations: int | None = None
    error_message: str | None = None


def _require_canonical_run(run_id: UUID) -> CanonicalRun:
    run = get_canonical_run(run_id)
    if run is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Run {run_id} not found",
        )
    return run


def _resolve_case_id(
    project_id: UUID,
    study_case_id: UUID | None,
    uow_factory: Any,
) -> str:
    if study_case_id is not None:
        return str(study_case_id)
    with uow_factory() as uow:
        active_case = uow.cases.get_active_study_case(project_id)
    if active_case is None:
        raise ValueError(f"Brak aktywnego przypadku dla projektu {project_id}")
    return str(active_case.id)


def _build_export_bundle(run_id: UUID) -> dict[str, Any]:
    return build_power_flow_export_bundle(_require_canonical_run(run_id))


def _format_catalog_binding(entry: dict[str, Any]) -> str:
    binding = entry.get("catalog_binding") or {}
    namespace = binding.get("catalog_namespace") or "-"
    item_id = binding.get("catalog_item_id") or "-"
    version = binding.get("catalog_item_version")
    if version:
        return f"{namespace}:{item_id} ({version})"
    return f"{namespace}:{item_id}"


def _truncate_catalog_params(value: Any, max_chars: int = 220) -> str:
    if value is None:
        return "—"
    payload = canonicalize_json(value)
    text = str(payload)
    if isinstance(payload, dict | list):
        import json

        text = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    if len(text) <= max_chars:
        return text
    return f"{text[:max_chars]}..."


def _catalog_context_lines(bundle: dict[str, Any], *, max_items: int = 20) -> list[str]:
    lines: list[str] = []
    for entry in (bundle.get("catalog_context") or [])[:max_items]:
        lines.append(
            " | ".join(
                [
                    str(entry.get("element_id") or "-"),
                    str(entry.get("element_type") or "-"),
                    _format_catalog_binding(entry),
                    str(entry.get("parameter_source") or entry.get("parameter_origin") or "-"),
                    _truncate_catalog_params(entry.get("materialized_params")),
                ]
            )
        )
    return lines


def _xlsx_column_name(index: int) -> str:
    name = ""
    while index > 0:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def _xlsx_cell_xml(value: Any, row_index: int, column_index: int) -> str:
    cell_ref = f"{_xlsx_column_name(column_index)}{row_index}"
    if isinstance(value, bool):
        value = "TAK" if value else "NIE"
    if isinstance(value, int | float) and not isinstance(value, bool):
        if math.isfinite(float(value)):
            return f'<c r="{cell_ref}"><v>{value}</v></c>'
    text = "" if value is None else escape(str(value))
    return f'<c r="{cell_ref}" t="inlineStr"><is><t xml:space="preserve">{text}</t></is></c>'


def _xlsx_sheet_xml(rows: list[list[Any]]) -> str:
    row_xml_parts: list[str] = []
    for row_index, row in enumerate(rows, start=1):
        cells = "".join(
            _xlsx_cell_xml(value, row_index, column_index)
            for column_index, value in enumerate(row, start=1)
        )
        row_xml_parts.append(f'<row r="{row_index}">{cells}</row>')
    sheet_data = "".join(row_xml_parts)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f"<sheetData>{sheet_data}</sheetData>"
        "</worksheet>"
    )


def _build_power_flow_xlsx(bundle: dict[str, Any]) -> bytes:
    result = bundle["result"]
    metadata = bundle["metadata"]
    summary = result.get("summary", {}) or {}
    bus_results = result.get("bus_results", []) or []
    branch_results = result.get("branch_results", []) or []
    catalog_context = bundle.get("catalog_context", []) or []
    white_box_trace = bundle.get("white_box_trace", []) or []

    sheets: list[tuple[str, list[list[Any]]]] = [
        (
            "Podsumowanie",
            [
                ["Pole", "Wartosc"],
                ["Run ID", metadata.get("run_id")],
                ["StudyCase ID", metadata.get("study_case_id")],
                ["Status zbieznosci", "Zbiezny" if result.get("converged") else "Niezbiezny"],
                ["Liczba iteracji", result.get("iterations_count")],
                ["Wezel bilansujacy", result.get("slack_bus_id")],
                ["Straty P [MW]", summary.get("total_losses_p_mw")],
                ["Straty Q [Mvar]", summary.get("total_losses_q_mvar")],
                ["Min napiecie [pu]", summary.get("min_v_pu")],
                ["Max napiecie [pu]", summary.get("max_v_pu")],
                ["Snapshot hash", metadata.get("snapshot_hash")],
                ["Input hash", metadata.get("input_hash")],
            ],
        ),
        (
            "Szyny",
            [["Bus ID", "V [pu]", "Kat [deg]", "P inj [MW]", "Q inj [Mvar]"]]
            + [
                [
                    row.get("bus_id"),
                    row.get("v_pu"),
                    row.get("angle_deg"),
                    row.get("p_injected_mw"),
                    row.get("q_injected_mvar"),
                ]
                for row in bus_results
            ],
        ),
        (
            "Odcinki",
            [["Branch ID", "P from [MW]", "Q from [Mvar]", "P to [MW]", "Q to [Mvar]"]]
            + [
                [
                    row.get("branch_id"),
                    row.get("p_from_mw"),
                    row.get("q_from_mvar"),
                    row.get("p_to_mw"),
                    row.get("q_to_mvar"),
                ]
                for row in branch_results
            ],
        ),
        (
            "Katalog",
            [["Element", "Typ", "Katalog", "Pochodzenie", "Parametry"]]
            + [
                [
                    entry.get("element_id"),
                    entry.get("element_type"),
                    _format_catalog_binding(entry),
                    entry.get("parameter_source") or entry.get("parameter_origin"),
                    _truncate_catalog_params(entry.get("materialized_params")),
                ]
                for entry in catalog_context
            ],
        ),
        (
            "WhiteBox",
            [["Krok", "Tytul", "Proof ref", "Proof status", "Reporting status"]]
            + [
                [
                    step.get("step") or step.get("key"),
                    step.get("title"),
                    step.get("proof_ref"),
                    step.get("proof_status"),
                    step.get("reporting_status"),
                ]
                for step in white_box_trace
            ],
        ),
    ]

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/styles.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
        + "".join(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            for index in range(1, len(sheets) + 1)
        )
        + "</Types>"
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        "</Relationships>"
    )
    workbook_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "<sheets>"
        + "".join(
            f'<sheet name="{escape(name[:31])}" sheetId="{index}" r:id="rId{index}"/>'
            for index, (name, _) in enumerate(sheets, start=1)
        )
        + "</sheets></workbook>"
    )
    workbook_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(
            f'<Relationship Id="rId{index}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
            f'Target="worksheets/sheet{index}.xml"/>'
            for index in range(1, len(sheets) + 1)
        )
        + f'<Relationship Id="rId{len(sheets) + 1}" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" '
        'Target="styles.xml"/>'
        "</Relationships>"
    )
    styles_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<fonts count="1"><font><sz val="11"/><name val="Calibri"/></font></fonts>'
        '<fills count="1"><fill><patternFill patternType="none"/></fill></fills>'
        '<borders count="1"><border/></borders>'
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
        '<cellXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/></cellXfs>'
        '<cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>'
        "</styleSheet>"
    )

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook_xml)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/styles.xml", styles_xml)
        for index, (_, rows) in enumerate(sheets, start=1):
            archive.writestr(f"xl/worksheets/sheet{index}.xml", _xlsx_sheet_xml(rows))
    return buffer.getvalue()


@router.get("/projects/{project_id}/power-flow-runs")
def list_power_flow_runs(
    project_id: UUID,
    status: str | None = Query(
        default=None,
        description="Filtruj po statusie (CREATED, RUNNING, FINISHED, FAILED)",
    ),
) -> dict[str, Any]:
    runs = [
        {
            "id": str(run.id),
            "project_id": run.project_id,
            "study_case_id": run.case_id,
            "analysis_type": run.analysis_type,
            "status": run.status,
            "result_status": run.result_status,
            "created_at": run.created_at.isoformat(),
            "finished_at": run.finished_at.isoformat() if run.finished_at else None,
            "input_hash": run.input_hash,
            # B5 (karta CV-3.3-B): etykieta wyboru biegu w porownaniu A/B —
            # rodzaj + rewizja/scenariusz + krotki snapshot_hash — czyta te
            # same trzy pola co koperta rewizji (`enm/envelope.py`), zamiast
            # samego UUID biegu bez dowodu KTORY stan modelu opisuje.
            "snapshot_hash": run.snapshot_hash,
            "model_revision": (run.envelope or {}).get("model_revision"),
            "scenario_ref": (run.envelope or {}).get("scenario_ref"),
            "converged": ((run.raw_result or {}).get("result_v1") or {}).get("converged"),
            "iterations": ((run.raw_result or {}).get("result_v1") or {}).get("iterations_count"),
        }
        for run in list_canonical_runs_for_project(str(project_id), analysis_type="PF")
        if status is None or run.status == status
    ]
    runs.sort(key=lambda run: run.get("created_at") or "", reverse=True)
    return canonicalize_json({"runs": runs, "total": len(runs)})


@router.post("/projects/{project_id}/power-flow-runs")
def create_power_flow_run(
    project_id: UUID,
    request: PowerFlowRunCreateRequest,
    uow_factory: Callable[[], UnitOfWork] = Depends(get_uow_factory),
) -> dict[str, Any]:
    try:
        options = dict(request.options or {})
        options.setdefault("trace_level", "summary")
        case_id = _resolve_case_id(
            project_id,
            request.study_case_id,
            uow_factory,
        )
        klucz = klucz_twin_z_uow(case_id, uow_factory)
        run = create_canonical_run(
            case_id=case_id,
            klucz_twin=klucz,
            project_id=str(project_id),
            analysis_type="PF",
            options=options,
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc

    return canonicalize_json(
        {
            "id": str(run.id),
            "deterministic_id": run.input_hash,
            "project_id": str(project_id),
            "study_case_id": case_id,
            "analysis_type": run.analysis_type,
            "status": run.status,
            "result_status": run.result_status,
            "created_at": run.created_at.isoformat() if run.created_at else None,
            "input_hash": run.input_hash,
        }
    )


@router.post("/power-flow-runs/{run_id}/execute")
def execute_power_flow_run(run_id: UUID) -> dict[str, Any]:
    _require_canonical_run(run_id)
    run = execute_canonical_run(run_id)
    result_v1 = (run.raw_result or {}).get("result_v1") or {}
    return canonicalize_json(
        {
            "id": str(run.id),
            "status": run.status,
            "converged": result_v1.get("converged"),
            "iterations": result_v1.get("iterations_count"),
            "error_message": run.error_message,
        }
    )


@router.get("/power-flow-runs/{run_id}")
def get_power_flow_run(run_id: UUID) -> dict[str, Any]:
    try:
        return canonicalize_json(build_power_flow_run_header(_require_canonical_run(run_id)))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/results")
def get_power_flow_results(run_id: UUID) -> dict[str, Any]:
    try:
        return canonicalize_json(get_canonical_power_flow_result(_require_canonical_run(run_id)))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/trace")
def get_power_flow_trace(run_id: UUID) -> dict[str, Any]:
    try:
        return canonicalize_json(get_canonical_power_flow_trace(_require_canonical_run(run_id)))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/export/json")
def export_power_flow_run_json(run_id: UUID) -> Response:
    try:
        return export_run_json_response(
            _require_canonical_run(run_id),
            filename_stem="power_flow_run",
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc


@router.get("/power-flow-runs/{run_id}/export/docx")
def export_power_flow_run_docx(run_id: UUID) -> Response:
    """Eksport wyniku rozplywu mocy do DOCX.

    Deterministyczny: ten sam zapisany wynik eksportowany wielokrotnie daje
    identyczne bajty (`docx_determinism.make_docx_bytes_deterministic` — sam
    python-docx NIE zeruje znacznikow czasu wpisow ZIP ani docProps/core.xml,
    dwa kolejne `Document().save()` roznia sie bajtowo, gdy wywolania przetna
    granice sekundy).
    """
    import io
    import json

    from fastapi.responses import Response
    from network_model.reporting.docx_determinism import make_docx_bytes_deterministic

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx",
        )

    bundle = _build_export_bundle(run_id)
    result = bundle["result"]
    metadata = bundle["metadata"]
    catalog_context_lines = _catalog_context_lines(bundle)
    white_box_trace = bundle.get("white_box_trace") or []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading("Raport rozplywu mocy", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_parts = [
        f"Status: {'Zbiezny' if result.get('converged') else 'Niezbiezny'}",
        f"Iteracje: {result.get('iterations_count', '—')}",
        f"Run: {metadata.get('run_id', '—')[:8]}...",
    ]
    subtitle = doc.add_paragraph(" | ".join(status_parts))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Podsumowanie", level=1)
    summary = result.get("summary", {})
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Parametr"
    hdr[1].text = "Wartosc"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    def add_row(label: str, value: Any) -> None:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(value) if value is not None else "—"

    add_row("Status zbieznosci", "Zbiezny" if result.get("converged") else "Niezbiezny")
    add_row("Liczba iteracji", result.get("iterations_count"))
    add_row("Wezel bilansujacy", result.get("slack_bus_id"))
    add_row("Calkowite straty P [MW]", format_wynik(summary.get("total_losses_p_mw"), ".4g"))
    add_row("Calkowite straty Q [Mvar]", format_wynik(summary.get("total_losses_q_mvar"), ".4g"))
    add_row("Min. napiecie [pu]", format_wynik(summary.get("min_v_pu"), ".4g"))
    add_row("Max. napiecie [pu]", format_wynik(summary.get("max_v_pu"), ".4g"))
    add_row("Elementy z katalogiem", metadata.get("catalog_context_count"))

    doc.add_paragraph()
    doc.add_heading("Kontekst katalogowy", level=1)
    if catalog_context_lines:
        doc.add_paragraph(
            "Format: element_id | typ | katalog | pochodzenie parametrow | materialized_params"
        )
        for line in catalog_context_lines:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("Brak jawnego kontekstu katalogowego.")

    doc.add_paragraph()
    doc.add_heading("White Box", level=1)
    if white_box_trace:
        for step in white_box_trace[:12]:
            title = step.get("title") or step.get("key") or "Krok"
            doc.add_paragraph(f"{title}: {json.dumps(step, ensure_ascii=False, sort_keys=True)}")
        if len(white_box_trace) > 12:
            doc.add_paragraph(f"... oraz {len(white_box_trace) - 12} kolejnych krokow")
    else:
        doc.add_paragraph("Brak jawnego sladu White Box.")

    doc.add_paragraph()
    doc.add_heading("Wyniki wezlowe (szyny)", level=1)
    bus_results = result.get("bus_results", [])
    if bus_results:
        bus_table = doc.add_table(rows=1, cols=5)
        bus_table.style = "Table Grid"
        header = bus_table.rows[0].cells
        header[0].text = "ID szyny"
        header[1].text = "V [pu]"
        header[2].text = "Kat [deg]"
        header[3].text = "P_inj [MW]"
        header[4].text = "Q_inj [Mvar]"
        for cell in header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for bus in bus_results[:30]:
            row = bus_table.add_row().cells
            row[0].text = str(bus.get("bus_id", "—"))[:16]
            row[1].text = format_wynik(bus.get("v_pu"), ".4g")
            row[2].text = format_wynik(bus.get("angle_deg"), ".2f")
            row[3].text = format_wynik(bus.get("p_injected_mw"), ".3g")
            row[4].text = format_wynik(bus.get("q_injected_mvar"), ".3g")
        if len(bus_results) > 30:
            doc.add_paragraph(f"... oraz {len(bus_results) - 30} dodatkowych wezlow")
    else:
        doc.add_paragraph("Brak wynikow wezlowych.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(
        content=make_docx_bytes_deterministic(buffer.getvalue()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="power_flow_run_{run_id}.docx"'},
    )


@router.get("/power-flow-runs/{run_id}/export/xlsx")
def export_power_flow_run_xlsx(run_id: UUID) -> Response:
    from fastapi.responses import Response

    bundle = _build_export_bundle(run_id)
    workbook_bytes = _build_power_flow_xlsx(bundle)
    return Response(
        content=workbook_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="power_flow_run_{run_id}.xlsx"'},
    )


@router.get("/power-flow-runs/{run_id}/export/pdf")
def export_power_flow_run_pdf(run_id: UUID) -> Response:
    import json

    from fastapi.responses import Response

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab",
        )

    bundle = _build_export_bundle(run_id)
    result = bundle["result"]
    metadata = bundle["metadata"]
    catalog_context_lines = _catalog_context_lines(bundle)
    white_box_trace = bundle.get("white_box_trace") or []

    buffer = io.BytesIO()
    zarejestruj_czcionki()
    canvas_obj = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 25 * mm
    top_margin = page_height - 25 * mm
    y = top_margin
    line_height = 5 * mm

    canvas_obj.setFont("DejaVuSans-Bold", 16)
    title = "Raport rozplywu mocy"
    canvas_obj.drawString(
        (page_width - canvas_obj.stringWidth(title, "DejaVuSans-Bold", 16)) / 2,
        y,
        title,
    )
    y -= 10 * mm

    canvas_obj.setFont("DejaVuSans", 10)
    status_text = (
        f"Status: {'Zbiezny' if result.get('converged') else 'Niezbiezny'} | "
        f"Iteracje: {result.get('iterations_count', '—')} | "
        f"Run: {metadata.get('run_id', '—')[:8]}..."
    )
    canvas_obj.drawString(left_margin, y, status_text)
    y -= 8 * mm

    canvas_obj.setFont("DejaVuSans-Bold", 14)
    canvas_obj.drawString(left_margin, y, "Podsumowanie")
    y -= 6 * mm

    canvas_obj.setFont("DejaVuSans", 10)
    summary = result.get("summary", {})
    summary_lines = [
        f"Wezel bilansujacy: {result.get('slack_bus_id', '—')}",
        f"Calkowite straty P: {format_wynik(summary.get('total_losses_p_mw'), '.4g')} MW",
        f"Calkowite straty Q: {format_wynik(summary.get('total_losses_q_mvar'), '.4g')} Mvar",
        f"Min. napiecie: {format_wynik(summary.get('min_v_pu'), '.4g')} pu",
        f"Max. napiecie: {format_wynik(summary.get('max_v_pu'), '.4g')} pu",
        f"Elementy z katalogiem: {metadata.get('catalog_context_count', 0)}",
    ]
    for line in summary_lines:
        canvas_obj.drawString(left_margin, y, line)
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Kontekst katalogowy")
    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans", 8)
    if catalog_context_lines:
        canvas_obj.drawString(
            left_margin,
            y,
            "Format: element_id | typ | katalog | pochodzenie | materialized_params",
        )
        y -= line_height
        for line in catalog_context_lines:
            canvas_obj.drawString(left_margin, y, line[:160])
            y -= line_height
            if y < 30 * mm:
                canvas_obj.showPage()
                y = top_margin
                canvas_obj.setFont("DejaVuSans", 8)
    else:
        canvas_obj.drawString(left_margin, y, "Brak jawnego kontekstu katalogowego.")
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "White Box")
    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans", 8)
    if white_box_trace:
        for step in white_box_trace[:10]:
            title = step.get("title") or step.get("key") or "Krok"
            canvas_obj.drawString(
                left_margin,
                y,
                f"{title}: {json.dumps(step, ensure_ascii=False, sort_keys=True)[:150]}",
            )
            y -= line_height
            if y < 30 * mm:
                canvas_obj.showPage()
                y = top_margin
                canvas_obj.setFont("DejaVuSans", 8)
        if len(white_box_trace) > 10:
            canvas_obj.drawString(
                left_margin,
                y,
                f"... oraz {len(white_box_trace) - 10} kolejnych krokow",
            )
            y -= line_height
    else:
        canvas_obj.drawString(left_margin, y, "Brak jawnego sladu White Box.")
        y -= line_height

    y -= 5 * mm
    canvas_obj.setFont("DejaVuSans-Bold", 12)
    canvas_obj.drawString(left_margin, y, "Wyniki wezlowe (top 20)")
    y -= 5 * mm

    canvas_obj.setFont("DejaVuSans", 9)
    for bus in result.get("bus_results", [])[:20]:
        text = (
            f"{str(bus.get('bus_id', '—'))[:12]}: "
            f"V={format_wynik(bus.get('v_pu'), '.4g')} pu, "
            f"kat={format_wynik(bus.get('angle_deg'), '.2f')} deg"
        )
        canvas_obj.drawString(left_margin, y, text)
        y -= line_height
        if y < 30 * mm:
            canvas_obj.showPage()
            y = top_margin

    canvas_obj.save()
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="power_flow_run_{run_id}.pdf"'},
    )


@router.get("/power-flow-runs/{run_id}/export/proof/json")
def export_power_flow_proof_json(run_id: UUID) -> dict[str, Any]:
    _ = run_id
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Eksport proof dla power-flow został wycofany z toru produkcyjnego do czasu "
            "pełnego przepięcia na kanoniczny White Box."
        ),
    )


@router.get("/power-flow-runs/{run_id}/export/proof/latex")
def export_power_flow_proof_latex(run_id: UUID) -> dict[str, Any]:
    _ = run_id
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Eksport proof dla power-flow został wycofany z toru produkcyjnego do czasu "
            "pełnego przepięcia na kanoniczny White Box."
        ),
    )


@router.get("/power-flow-runs/{run_id}/export/proof/pdf")
def export_power_flow_proof_pdf(run_id: UUID) -> dict[str, Any]:
    _ = run_id
    raise HTTPException(
        status_code=status.HTTP_410_GONE,
        detail=(
            "Eksport proof dla power-flow został wycofany z toru produkcyjnego do czasu "
            "pełnego przepięcia na kanoniczny White Box."
        ),
    )


_interpretation_cache: dict[str, dict[str, Any]] = {}


@router.get("/power-flow-runs/{run_id}/interpretation")
def get_power_flow_interpretation(run_id: UUID) -> dict[str, Any]:
    run_id_str = str(run_id)
    if run_id_str in _interpretation_cache:
        return canonicalize_json(_interpretation_cache[run_id_str])

    try:
        result_dict = build_power_flow_interpretation(_require_canonical_run(run_id))
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(exc),
        ) from exc
    _interpretation_cache[run_id_str] = result_dict
    return canonicalize_json(result_dict)


@router.post("/power-flow-runs/{run_id}/interpretation")
def create_power_flow_interpretation(run_id: UUID) -> dict[str, Any]:
    return get_power_flow_interpretation(run_id)
