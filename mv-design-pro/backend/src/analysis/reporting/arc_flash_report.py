"""
Arc Flash Report module — deterministyczny raport (JSON / tekst / PDF / DOCX / LaTeX)
dla analizy zagrożenia łukiem elektrycznym (IEEE 1584-2018).

Czyta zserializowany ``ArcFlashView`` (analysis.arc_flash.serializer.view_to_dict):
per szyna — energia incydentu [cal/cm²], granica łuku AFB [mm], kategoria ŚOI (PPE),
czas łuku, napięcie, odległość robocza, status i proweniencja.

Wzorzec 1:1 z ``audit2_report`` (Phase 5 backend integration):
- JSON: strukturalne dane (frontend ReportSurface / eksport).
- Tekst PL: dla inżyniera bez Acrobata.
- PDF: reportlab (deterministyczny — identyczny widok → identyczny PDF).
- DOCX: python-docx (deterministyczny).
- LaTeX: źródło do dołączenia do dowodów inżynierskich.

INVARIANTY:
- Identyczny widok → identyczny output (deterministic; sort po ``bus_ref``).
- Polskie etykiety, brak terminów zakazanych.
- Brak generacji z pustego stanu (brak wyników → „Brak szyn do uwzględnienia").
- ZERO fizyki: raport TYLKO interpretuje wartości policzone przez solver arc flash.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any

from network_model.reporting.czcionki import ustaw_czcionki_stylow, zarejestruj_czcionki

# Publiczny próg granicy łuku (definicja fizyczna IEEE 1584): E = 1,2 cal/cm².
_AFB_THRESHOLD_CAL_CM2 = 1.2


@dataclass(frozen=True)
class ArcFlashReportContext:
    """Kontekst raportu arc flash."""

    project_name: str
    station_id: str
    arc_flash_view_dict: dict[str, Any]
    operator_pl: str = ""
    generated_at_iso: str = "1970-01-01T00:00:00Z"


def _results(ctx: ArcFlashReportContext) -> list[dict[str, Any]]:
    """Wyniki per szyna, deterministycznie posortowane po ``bus_ref``."""
    results = ctx.arc_flash_view_dict.get("results", []) or []
    return sorted(results, key=lambda r: str(r.get("bus_ref", "")))


def _fmt(value: Any, digits: int = 2) -> str:
    """Deterministyczne formatowanie liczby (lub „—" dla braku)."""
    if isinstance(value, int | float):
        return f"{float(value):.{digits}f}"
    return "—"


def _summary(results: list[dict[str, Any]]) -> dict[str, Any]:
    """Podsumowanie: najgorszy przypadek + rozkład kategorii ŚOI + braki danych."""
    energies = [
        (r.get("bus_ref"), float(e))
        for r in results
        if isinstance((e := r.get("incident_energy_cal_cm2")), int | float)
    ]
    worst_bus_ref: str | None = None
    worst_energy: float | None = None
    if energies:
        worst_bus_ref, worst_energy = max(energies, key=lambda t: t[1])

    ppe_distribution: dict[str, int] = {}
    for r in results:
        category = r.get("ppe_category")
        key = str(category) if category is not None else "—"
        ppe_distribution[key] = ppe_distribution.get(key, 0) + 1

    with_missing = sum(1 for r in results if r.get("missing_data"))
    return {
        "bus_count": len(results),
        "worst_bus_ref": worst_bus_ref,
        "worst_incident_energy_cal_cm2": worst_energy,
        "ppe_distribution": dict(sorted(ppe_distribution.items())),
        "buses_with_missing_data": with_missing,
    }


def render_arc_flash_report_json(ctx: ArcFlashReportContext) -> dict[str, Any]:
    """Strukturalny raport JSON (deterministic)."""
    results = _results(ctx)
    view = ctx.arc_flash_view_dict
    return {
        "report_kind": "arc_flash_report",
        "format_version": "1.0",
        "project_name": ctx.project_name,
        "station_id": ctx.station_id,
        "operator_pl": ctx.operator_pl,
        "generated_at": ctx.generated_at_iso,
        "analysis_id": view.get("analysis_id"),
        "status": view.get("status"),
        "status_label_pl": view.get("status_label_pl"),
        "afb_threshold_cal_cm2": _AFB_THRESHOLD_CAL_CM2,
        "summary": _summary(results),
        "results": results,
    }


def _text_rows(results: list[dict[str, Any]]) -> list[str]:
    lines: list[str] = []
    for r in results:
        lines.append(
            f"  [{r.get('status_label_pl', r.get('status', ''))}] "
            f"{r.get('bus_ref', '')} — "
            f"E={_fmt(r.get('incident_energy_cal_cm2'))} cal/cm², "
            f"AFB={_fmt(r.get('arc_flash_boundary_mm'), 0)} mm, "
            f"ŚOI={r.get('ppe_category') or '—'}, "
            f"t_łuku={_fmt(r.get('arc_time_s'), 3)} s, "
            f"U={_fmt(r.get('voltage_kv'), 2)} kV"
        )
    return lines


def render_arc_flash_report_text(ctx: ArcFlashReportContext) -> str:
    """Plain text PL (dla inżyniera / audytora)."""
    results = _results(ctx)
    if not results:
        return (
            "Raport zagrożenia łukiem elektrycznym — brak szyn do uwzględnienia.\n"
            f"Stacja: {ctx.station_id}\n"
            f"Projekt: {ctx.project_name}\n"
        )

    summary = _summary(results)
    lines: list[str] = []
    lines.append("RAPORT ZAGROŻENIA ŁUKIEM ELEKTRYCZNYM (IEEE 1584)")
    lines.append("=" * 52)
    lines.append(f"Projekt: {ctx.project_name}")
    lines.append(f"Stacja: {ctx.station_id}")
    if ctx.operator_pl:
        lines.append(f"Operator: {ctx.operator_pl}")
    lines.append(f"Wygenerowano: {ctx.generated_at_iso}")
    lines.append("")
    lines.append(
        f"PODSUMOWANIE: {summary['bus_count']} szyn, "
        f"najwyższa energia incydentu {_fmt(summary['worst_incident_energy_cal_cm2'])} cal/cm²"
        + (f" (szyna {summary['worst_bus_ref']})" if summary["worst_bus_ref"] else "")
        + f", szyny z brakami danych: {summary['buses_with_missing_data']}."
    )
    lines.append(f"Granica łuku (próg AFB): {_AFB_THRESHOLD_CAL_CM2} cal/cm².")
    lines.append("")
    lines.append("--- Wyniki per szyna ---")
    lines.extend(_text_rows(results))
    lines.append("")
    return "\n".join(lines) + "\n"


def render_arc_flash_report_pdf(ctx: ArcFlashReportContext) -> bytes:
    """PDF z reportlab (deterministyczny). Fallback: tekst, gdy reportlab niedostępny."""
    try:
        from reportlab.lib import colors as rl_colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError:
        return render_arc_flash_report_text(ctx).encode("utf-8")

    results = _results(ctx)
    summary = _summary(results)
    buffer = BytesIO()
    zarejestruj_czcionki()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        title=f"Raport arc flash — {ctx.station_id}",
        author="MV-DESIGN-PRO",
        invariant=1,
        pageCompression=0,
    )
    styles = getSampleStyleSheet()
    ustaw_czcionki_stylow(styles)
    elements: list = []

    elements.append(Paragraph("Raport zagrożenia łukiem elektrycznym", styles["Title"]))
    elements.append(Spacer(1, 0.3 * cm))
    elements.append(Paragraph("IEEE 1584-2018", styles["Normal"]))
    elements.append(Spacer(1, 0.4 * cm))
    elements.append(Paragraph(f"Projekt: {ctx.project_name}", styles["Normal"]))
    elements.append(Paragraph(f"Stacja: {ctx.station_id}", styles["Normal"]))
    if ctx.operator_pl:
        elements.append(Paragraph(f"Operator: {ctx.operator_pl}", styles["Normal"]))
    elements.append(Paragraph(f"Wygenerowano: {ctx.generated_at_iso}", styles["Normal"]))
    elements.append(Spacer(1, 0.5 * cm))

    elements.append(Paragraph("Podsumowanie", styles["Heading1"]))
    summary_data = [
        ["Liczba szyn", str(summary["bus_count"])],
        [
            "Najwyższa energia incydentu [cal/cm²]",
            _fmt(summary["worst_incident_energy_cal_cm2"]),
        ],
        ["Szyna krytyczna", str(summary["worst_bus_ref"] or "—")],
        ["Szyny z brakami danych", str(summary["buses_with_missing_data"])],
        ["Próg granicy łuku [cal/cm²]", _fmt(_AFB_THRESHOLD_CAL_CM2)],
    ]
    summary_table = Table(summary_data, colWidths=[9 * cm, 7 * cm])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), rl_colors.HexColor("#cccccc")),
                ("GRID", (0, 0), (-1, -1), 0.25, rl_colors.black),
                ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]
        )
    )
    elements.append(summary_table)
    elements.append(Spacer(1, 0.5 * cm))

    if not results:
        elements.append(Paragraph("Brak szyn do uwzględnienia w raporcie.", styles["Normal"]))
    else:
        elements.append(Paragraph("Wyniki per szyna", styles["Heading2"]))
        rows: list[list[str]] = [
            ["Szyna", "U [kV]", "E [cal/cm²]", "AFB [mm]", "ŚOI", "t łuku [s]", "Status"]
        ]
        for r in results:
            rows.append(
                [
                    str(r.get("bus_ref", ""))[:28],
                    _fmt(r.get("voltage_kv"), 2),
                    _fmt(r.get("incident_energy_cal_cm2")),
                    _fmt(r.get("arc_flash_boundary_mm"), 0),
                    str(r.get("ppe_category") or "—"),
                    _fmt(r.get("arc_time_s"), 3),
                    str(r.get("status_label_pl") or r.get("status") or "")[:16],
                ]
            )
        tbl = Table(
            rows,
            colWidths=[4.4 * cm, 1.6 * cm, 2.4 * cm, 2.1 * cm, 1.4 * cm, 2.0 * cm, 2.5 * cm],
        )
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#dddddd")),
                    ("GRID", (0, 0), (-1, -1), 0.25, rl_colors.black),
                    ("FONTNAME", (0, 0), (-1, -1), "DejaVuSans"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        elements.append(tbl)

    doc.build(elements)
    return buffer.getvalue()


def render_arc_flash_report_docx(ctx: ArcFlashReportContext) -> bytes:
    """DOCX z python-docx (deterministyczny). Fallback: tekst, gdy docx niedostępny.

    Determinizm binarny wymaga normalizacji `docx_determinism` — sam python-docx
    NIE zeruje znacznikow czasu wpisow ZIP ani docProps/core.xml (dwa kolejne
    `Document().save()` roznia sie bajtowo, gdy wywolania przetna granice sekundy).
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return render_arc_flash_report_text(ctx).encode("utf-8")

    from network_model.reporting.docx_determinism import make_docx_bytes_deterministic

    results = _results(ctx)
    summary = _summary(results)
    doc = Document()
    title = doc.add_heading("Raport zagrożenia łukiem elektrycznym", level=0)
    title.alignment = 1  # center
    doc.add_paragraph("IEEE 1584-2018")

    doc.add_paragraph(f"Projekt: {ctx.project_name}")
    doc.add_paragraph(f"Stacja: {ctx.station_id}")
    if ctx.operator_pl:
        doc.add_paragraph(f"Operator: {ctx.operator_pl}")
    doc.add_paragraph(f"Wygenerowano: {ctx.generated_at_iso}")

    doc.add_heading("Podsumowanie", level=1)
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = "Light Shading"
    rows_data = [
        ("Liczba szyn", str(summary["bus_count"])),
        (
            "Najwyższa energia incydentu [cal/cm²]",
            _fmt(summary["worst_incident_energy_cal_cm2"]),
        ),
        ("Szyna krytyczna", str(summary["worst_bus_ref"] or "—")),
        ("Szyny z brakami danych", str(summary["buses_with_missing_data"])),
        ("Próg granicy łuku [cal/cm²]", _fmt(_AFB_THRESHOLD_CAL_CM2)),
    ]
    for i, (label, value) in enumerate(rows_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    if not results:
        doc.add_paragraph("Brak szyn do uwzględnienia w raporcie.")
    else:
        doc.add_heading("Wyniki per szyna", level=2)
        headers = ["Szyna", "U [kV]", "E [cal/cm²]", "AFB [mm]", "ŚOI", "t łuku [s]", "Status"]
        tbl = doc.add_table(rows=1 + len(results), cols=len(headers))
        tbl.style = "Light Grid"
        for c, head in enumerate(headers):
            tbl.rows[0].cells[c].text = head
        for i, r in enumerate(results, start=1):
            cells = tbl.rows[i].cells
            cells[0].text = str(r.get("bus_ref", ""))
            cells[1].text = _fmt(r.get("voltage_kv"), 2)
            cells[2].text = _fmt(r.get("incident_energy_cal_cm2"))
            cells[3].text = _fmt(r.get("arc_flash_boundary_mm"), 0)
            cells[4].text = str(r.get("ppe_category") or "—")
            cells[5].text = _fmt(r.get("arc_time_s"), 3)
            cells[6].text = str(r.get("status_label_pl") or r.get("status") or "")

    for para in doc.paragraphs:
        for run in para.runs:
            if not run.font.name:
                run.font.name = "Calibri"
            if not run.font.size:
                run.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    return make_docx_bytes_deterministic(buffer.getvalue())


def render_arc_flash_report_latex(ctx: ArcFlashReportContext) -> str:
    """LaTeX source (do dołączenia do dowodów inżynierskich)."""
    results = _results(ctx)
    summary = _summary(results)

    def esc(value: Any) -> str:
        return str(value).replace("&", r"\&").replace("_", r"\_").replace("%", r"\%")

    lines: list[str] = []
    lines.append(r"\documentclass[11pt,a4paper]{article}")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage[polish]{babel}")
    lines.append(r"\usepackage{amsmath,amssymb,longtable}")
    lines.append(r"\title{Raport zagrożenia łukiem elektrycznym (IEEE 1584)}")
    lines.append(rf"\author{{Stacja {esc(ctx.station_id)}}}")
    lines.append(rf"\date{{{esc(ctx.generated_at_iso)}}}")
    lines.append(r"\begin{document}")
    lines.append(r"\maketitle")
    lines.append(rf"\section{{Projekt}} {esc(ctx.project_name)}")
    lines.append(rf"\section{{Operator}} {esc(ctx.operator_pl) or 'brak'}")
    lines.append(r"\section{Podsumowanie}")
    lines.append(
        rf"\textbf{{Liczba szyn:}} {summary['bus_count']}, "
        rf"najwyższa energia incydentu {_fmt(summary['worst_incident_energy_cal_cm2'])} cal/cm\textsuperscript{{2}}"
        + (rf" (szyna {esc(summary['worst_bus_ref'])})" if summary["worst_bus_ref"] else "")
        + rf", szyny z brakami danych: {summary['buses_with_missing_data']}.\par"
    )
    if results:
        lines.append(r"\section{Wyniki per szyna}")
        lines.append(r"\begin{longtable}{l r r r c r}")
        lines.append(
            r"\textbf{Szyna} & \textbf{U [kV]} & \textbf{E [cal/cm\textsuperscript{2}]} "
            r"& \textbf{AFB [mm]} & \textbf{ŚOI} & \textbf{t [s]} \\ \hline"
        )
        for r in results:
            lines.append(
                rf"{esc(r.get('bus_ref', ''))} & {_fmt(r.get('voltage_kv'), 2)} & "
                rf"{_fmt(r.get('incident_energy_cal_cm2'))} & "
                rf"{_fmt(r.get('arc_flash_boundary_mm'), 0)} & "
                rf"{esc(r.get('ppe_category') or '—')} & "
                rf"{_fmt(r.get('arc_time_s'), 3)} \\"
            )
        lines.append(r"\end{longtable}")
    lines.append(r"\end{document}")
    return "\n".join(lines)


def render_arc_flash_report_all_formats(ctx: ArcFlashReportContext) -> dict[str, Any]:
    """Wszystkie formaty w jednym dictie (do enkapsulacji w API response)."""
    return {
        "json": render_arc_flash_report_json(ctx),
        "text_pl": render_arc_flash_report_text(ctx),
        "latex": render_arc_flash_report_latex(ctx),
        "pdf_size_bytes": len(render_arc_flash_report_pdf(ctx)),
        "docx_size_bytes": len(render_arc_flash_report_docx(ctx)),
    }
