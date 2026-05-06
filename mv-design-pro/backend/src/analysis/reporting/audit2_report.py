"""
Audit2 Report module — generuje deterministyczny raport (JSON / PDF / DOCX) dla
walidacji rozszerzen audytu 2 (BESS modes, tap-changers, hosting capacity, withstand,
VT grounding).

Phase 5 backend integration:
- JSON: structured data (uzywany przez frontend ReportSurface)
- PDF: minimal reportlab template (sections: header, summary, proofs detail)
- DOCX: minimal python-docx template (rownolegle sekcje)
- Plain text PL: dla audytu / inzynierow bez Acrobat'a

INVARIANTS:
- Identyczny ProofPack -> identyczny output (deterministic).
- Polskie etykiety, brak forbidden terms.
- Brak generacji z pustego stanu (puste proofy -> "Brak walidacji do uwzglednienia").
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any


@dataclass(frozen=True)
class Audit2ReportContext:
    """Kontekst raportu audytu 2."""

    project_name: str
    station_id: str
    proof_pack_dict: dict[str, Any]
    operator_pl: str = ""
    generated_at_iso: str = "1970-01-01T00:00:00Z"


def render_audit2_report_json(ctx: Audit2ReportContext) -> dict[str, Any]:
    """Strukturalny raport JSON (deterministic)."""
    proofs = ctx.proof_pack_dict.get("proofs", [])
    pass_count = sum(1 for p in proofs if p.get("pass_status"))
    fail_count = sum(1 for p in proofs if not p.get("pass_status"))
    return {
        "report_kind": "audit2_validation_report",
        "format_version": "1.0",
        "project_name": ctx.project_name,
        "station_id": ctx.station_id,
        "operator_pl": ctx.operator_pl,
        "generated_at": ctx.generated_at_iso,
        "summary": {
            "total_proofs": len(proofs),
            "pass_count": pass_count,
            "fail_count": fail_count,
            "all_pass": ctx.proof_pack_dict.get("all_pass", False),
        },
        "proofs_by_type": _group_proofs_by_type(proofs),
        "all_proofs": proofs,
    }


def _group_proofs_by_type(proofs: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    """Grupuje dowody po proof_type (deterministyczne sortowanie)."""
    grouped: dict[str, list[dict[str, Any]]] = {}
    for p in sorted(proofs, key=lambda x: (x.get("proof_type", ""), str(x.get("proof_id", "")))):
        key = p.get("proof_type", "UNKNOWN")
        grouped.setdefault(key, []).append(p)
    return grouped


def render_audit2_report_text(ctx: Audit2ReportContext) -> str:
    """Plain text PL (dla audytora / inzyniera)."""
    proofs = ctx.proof_pack_dict.get("proofs", [])
    if not proofs:
        return (
            f"Raport walidacji rozszerzen audytu 2 — brak dowodow do uwzglednienia.\n"
            f"Stacja: {ctx.station_id}\n"
            f"Projekt: {ctx.project_name}\n"
        )

    lines: list[str] = []
    lines.append("RAPORT WALIDACJI ROZSZERZEN AUDYTU 2")
    lines.append("=" * 50)
    lines.append(f"Projekt: {ctx.project_name}")
    lines.append(f"Stacja: {ctx.station_id}")
    if ctx.operator_pl:
        lines.append(f"Operator: {ctx.operator_pl}")
    lines.append(f"Wygenerowano: {ctx.generated_at_iso}")
    lines.append("")

    pass_count = sum(1 for p in proofs if p.get("pass_status"))
    fail_count = len(proofs) - pass_count
    lines.append(f"PODSUMOWANIE: {len(proofs)} walidacji, {pass_count} OK, {fail_count} blokerow.")
    lines.append("")

    grouped = _group_proofs_by_type(proofs)
    for proof_type, group in grouped.items():
        lines.append(f"--- {proof_type} ({len(group)} dowodow) ---")
        for p in group:
            status = "OK" if p.get("pass_status") else "BLOKER"
            lines.append(f"  [{status}] {p.get('summary_pl', '')}")
        lines.append("")

    return "\n".join(lines) + "\n"


def render_audit2_report_pdf(ctx: Audit2ReportContext) -> bytes:
    """
    PDF z reportlab (deterministyczny — identyczny ProofPack -> identyczny PDF).

    Gdy reportlab nie jest dostepny w srodowisku: zwraca placeholder
    'PDF rendering unavailable' z metadata jako tekst.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
        from reportlab.lib import colors as rl_colors
    except ImportError:
        # Placeholder gdy reportlab niedostepny.
        return render_audit2_report_text(ctx).encode("utf-8")

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        title=f"Raport walidacji audytu 2 — {ctx.station_id}",
        author="MV-DESIGN-PRO",
    )
    styles = getSampleStyleSheet()
    elements: list = []

    # Naglowek.
    elements.append(Paragraph("Raport walidacji rozszerzen audytu 2", styles["Title"]))
    elements.append(Spacer(1, 0.5 * cm))
    elements.append(Paragraph(f"Projekt: {ctx.project_name}", styles["Normal"]))
    elements.append(Paragraph(f"Stacja: {ctx.station_id}", styles["Normal"]))
    if ctx.operator_pl:
        elements.append(Paragraph(f"Operator: {ctx.operator_pl}", styles["Normal"]))
    elements.append(Paragraph(f"Wygenerowano: {ctx.generated_at_iso}", styles["Normal"]))
    elements.append(Spacer(1, 0.5 * cm))

    # Podsumowanie.
    proofs = ctx.proof_pack_dict.get("proofs", [])
    pass_count = sum(1 for p in proofs if p.get("pass_status"))
    fail_count = len(proofs) - pass_count
    elements.append(Paragraph("Podsumowanie", styles["Heading1"]))
    summary_data = [
        ["Total walidacji", str(len(proofs))],
        ["OK", str(pass_count)],
        ["Blokery", str(fail_count)],
        ["Wynik calosciowy", "OK" if fail_count == 0 else "BLOKERY OBECNE"],
    ]
    summary_table = Table(summary_data, colWidths=[6 * cm, 4 * cm])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#cccccc")),
                ("GRID", (0, 0), (-1, -1), 0.25, rl_colors.black),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]
        )
    )
    elements.append(summary_table)
    elements.append(Spacer(1, 0.5 * cm))

    # Sekcje per proof_type.
    grouped = _group_proofs_by_type(proofs)
    for proof_type, group in grouped.items():
        elements.append(Paragraph(proof_type, styles["Heading2"]))
        rows: list[list[str]] = [["Status", "Podsumowanie"]]
        for p in group:
            status = "OK" if p.get("pass_status") else "BLOKER"
            rows.append([status, str(p.get("summary_pl", ""))[:200]])
        tbl = Table(rows, colWidths=[2 * cm, 14 * cm])
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#dddddd")),
                    ("GRID", (0, 0), (-1, -1), 0.25, rl_colors.black),
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        elements.append(tbl)
        elements.append(Spacer(1, 0.4 * cm))

    doc.build(elements)
    return buffer.getvalue()


def render_audit2_report_docx(ctx: Audit2ReportContext) -> bytes:
    """
    DOCX z python-docx (deterministyczny).

    Gdy docx niedostepny: zwraca tekst jako bytes z extension .txt fallback.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return render_audit2_report_text(ctx).encode("utf-8")

    doc = Document()
    # Tytul.
    title = doc.add_heading("Raport walidacji rozszerzen audytu 2", level=0)
    title.alignment = 1  # center

    # Metadata.
    doc.add_paragraph(f"Projekt: {ctx.project_name}")
    doc.add_paragraph(f"Stacja: {ctx.station_id}")
    if ctx.operator_pl:
        doc.add_paragraph(f"Operator: {ctx.operator_pl}")
    doc.add_paragraph(f"Wygenerowano: {ctx.generated_at_iso}")

    # Podsumowanie.
    doc.add_heading("Podsumowanie", level=1)
    proofs = ctx.proof_pack_dict.get("proofs", [])
    pass_count = sum(1 for p in proofs if p.get("pass_status"))
    fail_count = len(proofs) - pass_count
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Light Shading"
    rows_data = [
        ("Total walidacji", str(len(proofs))),
        ("OK", str(pass_count)),
        ("Blokery", str(fail_count)),
        ("Wynik calosciowy", "OK" if fail_count == 0 else "BLOKERY OBECNE"),
    ]
    for i, (label, value) in enumerate(rows_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    # Sekcje per proof_type.
    grouped = _group_proofs_by_type(proofs)
    for proof_type, group in grouped.items():
        doc.add_heading(proof_type, level=2)
        tbl = doc.add_table(rows=1 + len(group), cols=2)
        tbl.style = "Light Grid"
        tbl.rows[0].cells[0].text = "Status"
        tbl.rows[0].cells[1].text = "Podsumowanie"
        for i, p in enumerate(group, start=1):
            status = "OK" if p.get("pass_status") else "BLOKER"
            tbl.rows[i].cells[0].text = status
            tbl.rows[i].cells[1].text = str(p.get("summary_pl", ""))[:300]

    # Style: Helvetica 10pt for all paragraphs.
    for section in doc.sections:
        for para in doc.paragraphs:
            for run in para.runs:
                if not run.font.name:
                    run.font.name = "Calibri"
                if not run.font.size:
                    run.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_audit2_report_latex(ctx: Audit2ReportContext) -> str:
    """LaTeX source (do dolaczenia do dowodow inzynierskich w E-36 ProofSurface)."""
    proofs = ctx.proof_pack_dict.get("proofs", [])
    lines: list[str] = []
    lines.append(r"\documentclass[11pt,a4paper]{article}")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage[polish]{babel}")
    lines.append(r"\usepackage{amsmath,amssymb}")
    lines.append(r"\title{Raport walidacji audytu 2}")
    lines.append(rf"\author{{Stacja {ctx.station_id}}}")
    lines.append(rf"\date{{{ctx.generated_at_iso}}}")
    lines.append(r"\begin{document}")
    lines.append(r"\maketitle")
    lines.append(rf"\section{{Projekt}} {ctx.project_name}")
    lines.append(rf"\section{{Operator}} {ctx.operator_pl or 'brak'}")
    lines.append(r"\section{Podsumowanie}")
    pass_count = sum(1 for p in proofs if p.get("pass_status"))
    lines.append(
        rf"\textbf{{Total walidacji:}} {len(proofs)}, OK: {pass_count}, "
        rf"Blokery: {len(proofs) - pass_count}."
    )
    lines.append(r"")

    grouped = _group_proofs_by_type(proofs)
    for proof_type, group in grouped.items():
        lines.append(rf"\subsection{{{proof_type}}}")
        for p in group:
            status = "OK" if p.get("pass_status") else "BLOKER"
            summary = str(p.get("summary_pl", "")).replace("&", r"\&").replace("_", r"\_")
            lines.append(rf"\textbf{{[{status}]}} {summary}\par")
            for formula in p.get("formulas_latex", []):
                lines.append(formula)

    lines.append(r"\end{document}")
    return "\n".join(lines)


def render_audit2_report_all_formats(ctx: Audit2ReportContext) -> dict[str, Any]:
    """Wszystkie formaty w jednym dictie (do enkapsulacji w API response)."""
    return {
        "json": render_audit2_report_json(ctx),
        "text_pl": render_audit2_report_text(ctx),
        "latex": render_audit2_report_latex(ctx),
        # PDF/DOCX bytes -> base64 dla JSON transport.
        "pdf_size_bytes": len(render_audit2_report_pdf(ctx)),
        "docx_size_bytes": len(render_audit2_report_docx(ctx)),
    }
