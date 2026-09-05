"""Serwis aplikacyjny: certyfikat zgodności projektu z wymaganiami NC RfG.

Warstwa APPLICATION — czysta KOMPOZYCJA gotowych werdyktów, NIE fizyka i NIE
ocena. Certyfikat powstaje WYŁĄCZNIE z gotowego wyniku macierzy zgodności NC RfG
(``NcRfgPtpireeRunResult`` z solvera ``network_model.solvers.ncrfg_ptpiree``),
czyli tej samej ścieżki, którą liczy macierz frontendu:
``POST /api/ncrfg-tests/run`` → ``NcRfgPtpireeSolver.run`` (patrz
``api/ncrfg_ptpiree_tests.py:41`` oraz ``engine.py:208``). Serwis NIE przelicza
żadnego testu — cytuje istniejące werdykty, ich podsumowania i odcisk wejścia.

Uczciwa bramka kompletności (wzór „lista braków przed generacją"):
- werdykt ``no_data`` na teście wymaganym, brak klasy modułu, lub moduł BEZ
  żadnego testu wymaganego I BEZ zweryfikowanego certyfikatu PTPiREE
  (``certificate_status != "ptpiree_verified"``) → certyfikat NIE powstaje,
  zwracana jest lista braków po polsku;
- moduł bez testu wymaganego, ale ZE zweryfikowanym certyfikatem PTPiREE
  (karta FAB-K), NIE jest brakiem — certyfikat z wykazu producenta jest
  SAMODZIELNĄ, silniejszą podstawą niż bieg testów NC RfG w warsztacie; zero
  testów wymaganych jest tu WNIOSKIEM klasyfikacji (typ modułu A, PPM), nie
  luką dowodową. Precedens defektu: przed tą kartą front czytał
  `ptpiree_certificate_ref` z pola, którego zapis nigdy go nie wypełniał —
  `certificate_status` był więc ZAWSZE `"unknown"`, więc ta gałąź nigdy się
  nie uruchamiała i luka była niewidoczna;
- werdykt negatywny (``fail``, moduł ``niezgodny``) NIE blokuje — dokument
  stwierdza stan (certyfikat z jednoznacznym werdyktem negatywnym).

Determinizm: identyczne wejście → identyczny widok JSON, identyczny
``odcisk_wejscia_sha256`` i bajtowo identyczny eksport DOCX.
"""

from __future__ import annotations

from collections.abc import Sequence
from io import BytesIO

from application.analyses.dowod_certyfikatu import (
    TYTUL_DOWODU,
    NcRfgCertificateEvidence,
    sekcja_dowodu,
    wiersze_dowodu_pl,
)
from network_model.reporting.czcionki import zarejestruj_czcionki
from network_model.reporting.docx_determinism import make_docx_bytes_deterministic
from network_model.solvers.ncrfg_ptpiree import (
    NcRfgPtpireeRunRequest,
    NcRfgPtpireeRunResult,
)
from pydantic import BaseModel, Field

try:  # pragma: no cover - zależy od środowiska
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False

try:  # pragma: no cover - zależy od środowiska
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfgen import canvas

    _PDF_AVAILABLE = True
except ImportError:  # pragma: no cover
    _PDF_AVAILABLE = False

CERTYFIKAT_CONTRACT = "CertyfikatZgodnosciNcRfgV1"
CERTYFIKAT_TYTUL = "Certyfikat zgodności projektu z wymaganiami NC RfG"

# Etykiety werdyktów testów (kontrakt solvera: pass/fail/no_data/not_required).
_WERDYKT_PL: dict[str, str] = {
    "pass": "spełnia",
    "fail": "nie spełnia",
    "no_data": "brak danych",
    "not_required": "nie dotyczy",
}
# Etykiety statusu zbiorczego modułu (kontrakt solvera).
_STATUS_MODULU_PL: dict[str, str] = {
    "zgodny": "Zgodny",
    "niezgodny": "Niezgodny",
    "brak_danych": "Brak danych",
}


class CertyfikatZgodnosciRequest(BaseModel):
    """Wejście końcówki certyfikatu: pakiet macierzy + identyfikacja projektu."""

    run_request: NcRfgPtpireeRunRequest
    nazwa_projektu: str = Field(min_length=1)
    nazwa_przypadku: str | None = None


class CertyfikatBrakiError(Exception):
    """Certyfikat nie powstaje z powodu braków kompletności (lista po polsku)."""

    def __init__(self, braki: list[str]) -> None:
        self.braki = braki
        super().__init__("Certyfikat nie może powstać — dane zgodności niekompletne.")


def zbierz_braki(run_result: NcRfgPtpireeRunResult) -> list[str]:
    """Zbierz braki kompletności blokujące generację (deterministyczna kolejność).

    Braki (per moduł, w kolejności macierzy):
    - brak klasy modułu (``module_type == '?'``),
    - moduł bez żadnego testu wymaganego I bez zweryfikowanego certyfikatu
      PTPiREE (nie ma czego certyfikować — ani biegu, ani wykazu producenta),
    - test wymagany z werdyktem ``no_data`` (brak danych do oceny).
    """
    braki: list[str] = []
    for module in run_result.modules:
        etykieta = module.der_name or module.der_ref
        if module.module_type == "?":
            braki.append(
                f"Moduł „{etykieta}”: brak klasyfikacji klasy modułu "
                "(moc/napięcie poza zakresem profilu operatora)."
            )
            continue
        if module.required_count == 0 and module.certificate_status != "ptpiree_verified":
            braki.append(
                f"Moduł „{etykieta}”: brak testów wymaganych — brak podstawy " "do certyfikacji."
            )
        for test in module.tests:
            if test.required and test.verdict == "no_data":
                braki.append(
                    f"Moduł „{etykieta}”: test {test.test_id} "
                    f"({test.ability_pl}) — brak danych do oceny."
                )
    return braki


def build_certyfikat_view(
    run_result: NcRfgPtpireeRunResult,
    *,
    nazwa_projektu: str,
    nazwa_przypadku: str | None = None,
    dowody: Sequence[NcRfgCertificateEvidence] | None = None,
) -> dict:
    """Zbuduj widok JSON certyfikatu z gotowego wyniku macierzy zgodności.

    Rzuca ``CertyfikatBrakiError`` gdy dane są niekompletne (bramka kompletności).

    ``dowody`` (opcjonalne) to dowód certyfikacji PTPiREE per urządzenie z
    tabliczek modelu — dokładnie ten sam blok, który niesie bieg NC RfG. Bez
    dowodów (wywołanie bez wskazanego przypadku) widok jest IDENTYCZNY jak przed
    dodaniem sekcji: klucz ``dowod_certyfikatu`` w ogóle nie powstaje.
    """
    braki = zbierz_braki(run_result)
    if braki:
        raise CertyfikatBrakiError(braki)

    indeks_dowodow = {dowod.der_ref: dowod for dowod in dowody} if dowody is not None else None

    moduly_view: list[dict] = []
    modulow_zgodnych = 0
    for module in run_result.modules:
        if module.overall_status == "zgodny":
            modulow_zgodnych += 1
        testy_view = [
            {
                "test_id": test.test_id,
                "nazwa_pl": test.ability_pl,
                "wymagany": test.required,
                "werdykt": test.verdict,
                "werdykt_pl": _WERDYKT_PL.get(test.verdict, test.verdict),
                "wartosci_pl": test.summary_pl,
            }
            for test in module.tests
            if test.required or test.verdict != "not_required"
        ]
        modul_view: dict = {
            "der_ref": module.der_ref,
            "der_name": module.der_name,
            "operator_pl": module.operator_name_pl,
            "klasa": module.module_type,
            "rodzina": module.module_family,
            "p_max_kw": module.p_max_kw,
            "voltage_kv": module.voltage_kv,
            "status": module.overall_status,
            "status_pl": _STATUS_MODULU_PL.get(module.overall_status, module.overall_status),
            "podsumowanie": {
                "wymagane": module.required_count,
                "spelnia": module.pass_count,
                "nie_spelnia": module.fail_count,
            },
            "testy": testy_view,
        }
        if indeks_dowodow is not None:
            dowod = indeks_dowodow.get(module.der_ref)
            modul_view["dowod_certyfikatu"] = sekcja_dowodu(
                dowod if dowod is not None else NcRfgCertificateEvidence(der_ref=module.der_ref)
            )
        moduly_view.append(modul_view)

    modulow = len(run_result.modules)
    modulow_niezgodnych = modulow - modulow_zgodnych
    projekt_zgodny = modulow_niezgodnych == 0
    operatorzy = sorted({m.operator_name_pl for m in run_result.modules})

    return {
        "kontrakt": CERTYFIKAT_CONTRACT,
        "tytul": CERTYFIKAT_TYTUL,
        "identyfikacja": {
            "projekt": nazwa_projektu,
            "przypadek": nazwa_przypadku,
            "procedura": run_result.procedure_version,
            "wersja_narzedzia": run_result.solver_version,
        },
        "werdykt_zbiorczy": {
            "status": "zgodny" if projekt_zgodny else "niezgodny",
            "etykieta_pl": (
                "Projekt zgodny z wymaganiami NC RfG"
                if projekt_zgodny
                else "Projekt niezgodny z wymaganiami NC RfG"
            ),
            "liczba_modulow": modulow,
            "modulow_zgodnych": modulow_zgodnych,
            "modulow_niezgodnych": modulow_niezgodnych,
        },
        "moduly": moduly_view,
        "zalozenia_i_zrodla": [
            "Certyfikat zestawia gotowe werdykty z macierzy zgodności NC RfG — "
            "nie przelicza testów i nie zastępuje badań terenowych ani sprawdzeń "
            "u operatora systemu.",
            "Procedura testowania: " + run_result.procedure_version + ".",
            "Wersja narzędzia obliczeniowego: " + run_result.solver_version + ".",
            "Profil(e) operatora: " + ", ".join(operatorzy) + ".",
        ],
        "odcisk_wejscia_sha256": run_result.input_hash,
        "odcisk_wyniku_sha256": run_result.deterministic_hash,
    }


def render_certyfikat_docx(view: dict) -> bytes:
    """Zrenderuj deterministyczny DOCX certyfikatu z widoku JSON.

    Zwraca bajty znormalizowane przez ``make_docx_bytes_deterministic`` — dwa
    wywołania na tym samym widoku dają identyczne bajty.
    """
    if not _DOCX_AVAILABLE:  # pragma: no cover
        raise ImportError("Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx")

    doc = Document()

    tytul = doc.add_heading(view["tytul"], level=0)
    tytul.alignment = WD_ALIGN_PARAGRAPH.CENTER

    identyfikacja = view["identyfikacja"]
    id_para = doc.add_paragraph()
    id_para.add_run("Projekt: ").bold = True
    id_para.add_run(str(identyfikacja["projekt"]))
    if identyfikacja.get("przypadek"):
        id_para.add_run("  |  Przypadek: ").bold = True
        id_para.add_run(str(identyfikacja["przypadek"]))
    proc_para = doc.add_paragraph()
    proc_para.add_run("Procedura: ").bold = True
    proc_para.add_run(str(identyfikacja["procedura"]))
    proc_para.add_run("  |  Narzędzie: ").bold = True
    proc_para.add_run(str(identyfikacja["wersja_narzedzia"]))

    werdykt = view["werdykt_zbiorczy"]
    doc.add_paragraph()
    w_para = doc.add_paragraph()
    w_para.add_run("Werdykt zbiorczy: ").bold = True
    w_run = w_para.add_run(werdykt["etykieta_pl"])
    w_run.bold = True
    if werdykt["status"] == "zgodny":
        w_run.font.color.rgb = RGBColor(22, 163, 74)
    else:
        w_run.font.color.rgb = RGBColor(220, 38, 38)
    licz_para = doc.add_paragraph()
    licz_para.add_run(
        f"Moduły: {werdykt['liczba_modulow']} "
        f"(zgodnych: {werdykt['modulow_zgodnych']}, "
        f"niezgodnych: {werdykt['modulow_niezgodnych']})"
    )

    for module in view["moduly"]:
        doc.add_paragraph()
        naglowek = module["der_name"] or module["der_ref"]
        doc.add_heading(f"Moduł: {naglowek} (klasa {module['klasa']})", level=1)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Operator: {module['operator_pl']}  |  "
            f"Pmax: {module['p_max_kw']} kW  |  "
            f"Napięcie: {module['voltage_kv']} kV  |  "
            f"Status: {module['status_pl']}"
        )

        tabela = doc.add_table(rows=1, cols=4)
        tabela.style = "Table Grid"
        naglowki = ["Test", "Zdolność", "Werdykt", "Wartości"]
        hdr = tabela.rows[0].cells
        for i, tekst in enumerate(naglowki):
            hdr[i].text = tekst
            hdr[i].paragraphs[0].runs[0].bold = True
        for test in module["testy"]:
            row = tabela.add_row().cells
            row[0].text = test["test_id"]
            row[1].text = test["nazwa_pl"]
            row[2].text = test["werdykt_pl"]
            row[3].text = test["wartosci_pl"]

        dowod = module.get("dowod_certyfikatu")
        if dowod is not None:
            dow_para = doc.add_paragraph()
            dow_para.add_run(f"{TYTUL_DOWODU}: ").bold = True
            wiersze = wiersze_dowodu_pl(dowod)
            if wiersze:
                dow_para.add_run(
                    "  |  ".join(f"{etykieta}: {wartosc}" for etykieta, wartosc in wiersze)
                )
            else:
                dow_para.add_run(str(dowod["stan_pl"]))

    doc.add_paragraph()
    doc.add_heading("Założenia i źródła", level=1)
    for pozycja in view["zalozenia_i_zrodla"]:
        doc.add_paragraph(pozycja, style="List Bullet")

    doc.add_paragraph()
    stopka = doc.add_paragraph()
    stopka.add_run(f"Odcisk SHA-256 wejścia: {view['odcisk_wejscia_sha256']}").font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    return make_docx_bytes_deterministic(buffer.getvalue())


def render_certyfikat_pdf(view: dict) -> bytes:
    """Zrenderuj deterministyczny PDF certyfikatu z widoku JSON (układ 1:1 z DOCX).

    Determinizm bajtowy: canvas z ``invariant=1`` (stały ``CreationDate`` i ``ID``
    dokumentu) oraz ``pageCompression=0`` — dwa wywołania na tym samym widoku dają
    identyczne bajty. Czcionki DejaVu Sans (polskie diakrytyki) rejestrowane wspólnym
    modułem ``network_model.reporting.czcionki`` (subset TTF jest deterministyczny).
    """
    if not _PDF_AVAILABLE:  # pragma: no cover
        raise ImportError("Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab")

    buffer = BytesIO()
    zarejestruj_czcionki()
    c = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 25 * mm
    right_edge = page_width - 25 * mm
    top_margin = page_height - 25 * mm
    bottom_margin = 25 * mm
    line_height = 5 * mm
    content_width = right_edge - left_margin

    y = top_margin

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = top_margin

    def ensure(needed: float) -> None:
        nonlocal y
        if y - needed < bottom_margin:
            new_page()

    def para(text: str, *, size: int = 10, bold: bool = False, indent: float = 0.0) -> None:
        nonlocal y
        font = "DejaVuSans-Bold" if bold else "DejaVuSans"
        for line in simpleSplit(text, font, size, content_width - indent):
            ensure(line_height)
            c.setFont(font, size)
            c.setFillColorRGB(0, 0, 0)
            c.drawString(left_margin + indent, y, line)
            y -= line_height

    # 1) Tytuł (wyśrodkowany).
    title = str(view["tytul"])
    c.setFont("DejaVuSans-Bold", 16)
    c.drawCentredString(page_width / 2, y, title)
    y -= 10 * mm

    # 2) Identyfikacja.
    identyfikacja = view["identyfikacja"]
    id_line = f"Projekt: {identyfikacja['projekt']}"
    if identyfikacja.get("przypadek"):
        id_line += f"  |  Przypadek: {identyfikacja['przypadek']}"
    para(id_line)
    para(
        f"Procedura: {identyfikacja['procedura']}  |  "
        f"Narzędzie: {identyfikacja['wersja_narzedzia']}"
    )
    y -= line_height

    # 3) Werdykt zbiorczy (kolor zależny od statusu).
    werdykt = view["werdykt_zbiorczy"]
    ensure(line_height * 2)
    c.setFont("DejaVuSans-Bold", 12)
    c.drawString(left_margin, y, "Werdykt zbiorczy: ")
    prefix_width = c.stringWidth("Werdykt zbiorczy: ", "DejaVuSans-Bold", 12)
    if werdykt["status"] == "zgodny":
        c.setFillColorRGB(22 / 255, 163 / 255, 74 / 255)
    else:
        c.setFillColorRGB(220 / 255, 38 / 255, 38 / 255)
    c.drawString(left_margin + prefix_width, y, str(werdykt["etykieta_pl"]))
    c.setFillColorRGB(0, 0, 0)
    y -= line_height
    para(
        f"Moduły: {werdykt['liczba_modulow']} "
        f"(zgodnych: {werdykt['modulow_zgodnych']}, "
        f"niezgodnych: {werdykt['modulow_niezgodnych']})"
    )
    y -= line_height

    # 4) Moduły — nagłówek, meta i tabela testów.
    module_cols = [22 * mm, 55 * mm, 25 * mm, content_width - 102 * mm]
    for module in view["moduly"]:
        ensure(line_height * 4)
        naglowek = module["der_name"] or module["der_ref"]
        para(f"Moduł: {naglowek} (klasa {module['klasa']})", size=12, bold=True)
        para(
            f"Operator: {module['operator_pl']}  |  "
            f"Pmax: {module['p_max_kw']} kW  |  "
            f"Napięcie: {module['voltage_kv']} kV  |  "
            f"Status: {module['status_pl']}"
        )
        header = ["Test", "Zdolność", "Werdykt", "Wartości"]
        ensure(line_height)
        c.setFont("DejaVuSans-Bold", 9)
        col_x = left_margin
        for i, label in enumerate(header):
            c.drawString(col_x, y, label)
            col_x += module_cols[i]
        y -= line_height
        for test in module["testy"]:
            cells = [
                str(test["test_id"]),
                str(test["nazwa_pl"]),
                str(test["werdykt_pl"]),
                str(test["wartosci_pl"]),
            ]
            wrapped = [
                simpleSplit(cell, "DejaVuSans", 9, module_cols[i] - 2 * mm)
                for i, cell in enumerate(cells)
            ]
            row_lines = max(len(w) for w in wrapped)
            ensure(line_height * row_lines)
            c.setFont("DejaVuSans", 9)
            col_x = left_margin
            for i, lines in enumerate(wrapped):
                for j, line in enumerate(lines):
                    c.drawString(col_x, y - j * line_height, line)
                col_x += module_cols[i]
            y -= line_height * row_lines
        dowod = module.get("dowod_certyfikatu")
        if dowod is not None:
            wiersze = wiersze_dowodu_pl(dowod)
            tresc = (
                "  |  ".join(f"{etykieta}: {wartosc}" for etykieta, wartosc in wiersze)
                if wiersze
                else str(dowod["stan_pl"])
            )
            para(f"{TYTUL_DOWODU}: {tresc}", size=9)
        y -= line_height

    # 5) Założenia i źródła.
    ensure(line_height * 2)
    para("Założenia i źródła", size=12, bold=True)
    for pozycja in view["zalozenia_i_zrodla"]:
        para(f"• {pozycja}", indent=4 * mm)

    # 6) Stopka — odcisk wejścia.
    y -= line_height
    para(f"Odcisk SHA-256 wejścia: {view['odcisk_wejscia_sha256']}", size=8)

    c.showPage()
    c.save()
    return buffer.getvalue()
