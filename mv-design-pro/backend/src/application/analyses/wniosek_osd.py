"""Serwis aplikacyjny: wniosek o określenie warunków przyłączenia (OSD).

Warstwa APPLICATION — czysta KOMPOZYCJA gotowych wyników, NIE fizyka i NIE
ocena. Wniosek powstaje WYŁĄCZNIE z istniejących źródeł, każda sekcja cytuje
swoje źródło (run_id / odcisk wejścia):

- **bilans mocy** ← widok walidacji energetycznej przebiegu rozpływu (``PF``)
  (``application.analyses.energy_validation.build_energy_validation_view`` —
  obciążenia gałęzi/transformatorów, bilans mocy biernej slack, straty sieciowe)
  uzupełniony o moc zainstalowaną źródeł falownikowych ze snapshotu (wzorzec
  ``grid_strength._installed_mva_by_bus`` z krotnością ``n_parallel``),
- **zwarcia w punkcie przyłączenia** ← przebieg zwarciowy (``short_circuit_sn``)
  zakończony: Ik'' / S_k'' wskazanego węzła
  (wzorzec ``grid_strength._sk_mva_by_bus`` → ``build_short_circuit_results``),
- **zgodność NC RfG** ← ``NcRfgPtpireeSolver`` przez ``build_certyfikat_view``
  (ta sama ścieżka co macierz i certyfikat D14) — werdykt zbiorczy + odesłanie
  do certyfikatu zgodności.

Poza zakresem (uczciwe adnotacje w ``zalozenia_pl``): schemat elektryczny
(światło techniczne SLD — osobny wątek) oraz zestawienia materiałowe (jeszcze
nie generowane przez narzędzie).

Uczciwa bramka kompletności (lista braków przed generacją, kolejność
deterministyczna: bilans → zwarcia → zgodność):
- przebieg rozpływu złego rodzaju lub niezakończony,
- przebieg zwarciowy złego rodzaju lub niezakończony,
- wskazany węzeł przyłączenia bez wyniku zwarciowego,
- braki modułów NC RfG (ta sama lista co certyfikat D14).
Przy jakimkolwiek braku wniosek NIE powstaje — zwracana jest lista po polsku.

Determinizm: identyczne wejście → identyczny widok JSON, identyczne odciski
sekcji, identyczny ``input_hash`` i bajtowo identyczny eksport DOCX.
"""

from __future__ import annotations

import hashlib
import json
from collections.abc import Sequence
from io import BytesIO
from typing import Any

from application.analyses.certyfikat_zgodnosci import (
    CertyfikatBrakiError,
    build_certyfikat_view,
)
from application.analyses.dowod_certyfikatu import (
    TYTUL_DOWODU,
    NcRfgCertificateEvidence,
    sekcje_dowodow,
    wiersze_dowodu_pl,
)
from application.analyses.energy_validation.service import build_energy_validation_view
from application.analyses.grid_strength import _installed_mva_by_bus
from enm.canonical_analysis import CanonicalRun, build_short_circuit_results
from network_model.reporting.czcionki import zarejestruj_czcionki
from network_model.reporting.docx_determinism import make_docx_bytes_deterministic
from network_model.solvers.ncrfg_ptpiree import NcRfgPtpireeRunResult
from pydantic import BaseModel, Field

try:  # pragma: no cover - zależy od środowiska
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False

try:  # pragma: no cover - zależy od środowiska
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfgen import canvas

    _PDF_AVAILABLE = True
except ImportError:  # pragma: no cover
    _PDF_AVAILABLE = False

WNIOSEK_OSD_CONTRACT = "WniosekOkresleniaWarunkowPrzylaczeniaV1"
WNIOSEK_OSD_TYTUL = "Wniosek o określenie warunków przyłączenia do sieci OSD"

_ADNOTACJA_SCHEMAT = (
    "Schemat elektryczny (światło techniczne) dołączany jest do wniosku odrębnie "
    "— poza zakresem niniejszego dokumentu."
)
_ADNOTACJA_ZESTAWIENIA = (
    "Zestawienia materiałowe nie są jeszcze generowane przez narzędzie — sekcja "
    "pominięta (bez danych zastępczych)."
)


class WniosekOsdIdentyfikacja(BaseModel):
    """Dane identyfikacyjne wniosku (pola tekstowe — opcjonalne poza nazwą)."""

    nazwa_projektu: str = Field(min_length=1)
    nazwa_przypadku: str | None = None
    wnioskodawca: str | None = None
    adres_przylaczenia: str | None = None


class WniosekOsdBrakiError(Exception):
    """Wniosek nie powstaje z powodu braków kompletności (lista po polsku)."""

    def __init__(self, braki: list[str]) -> None:
        self.braki = braki
        super().__init__("Wniosek nie może powstać — dane przyłączeniowe niekompletne.")


def _odcisk(payload: Any) -> str:
    """Deterministyczny odcisk SHA-256 sekcji (kanoniczny JSON, sortowane klucze)."""
    canonical = json.dumps(payload, sort_keys=True, ensure_ascii=False, default=str)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def _sc_row_for_bus(sc_run: CanonicalRun, bus_ref: str) -> dict[str, Any] | None:
    """Wiersz wyniku zwarciowego dla węzła przyłączenia (``element_id == bus_ref``)."""
    for row in build_short_circuit_results(sc_run).get("rows", []):
        if row.get("element_id") == bus_ref:
            return row
    return None


def zbierz_braki_wniosku(
    pf_run: CanonicalRun,
    sc_run: CanonicalRun,
    bus_ref: str,
    ncrfg_run_result: NcRfgPtpireeRunResult,
) -> list[str]:
    """Zbierz braki kompletności blokujące generację (kolejność deterministyczna).

    Kolejność: bilans mocy (przebieg PF) → zwarcia (przebieg SC + węzeł) →
    zgodność NC RfG (braki modułów, ta sama lista co certyfikat D14).
    """
    braki: list[str] = []

    # 1. Bilans mocy — przebieg rozpływu.
    if pf_run.analysis_type != "PF":
        braki.append(
            "Bilans mocy: wskazany przebieg nie jest rozpływem mocy (PF); "
            f"otrzymano rodzaj analizy: {pf_run.analysis_type}."
        )
    elif pf_run.status != "FINISHED":
        braki.append(
            f"Bilans mocy: przebieg rozpływu {pf_run.id} nie jest zakończony "
            f"(status={pf_run.status})."
        )

    # 2. Zwarcia — przebieg zwarciowy + obecność węzła przyłączenia.
    if sc_run.analysis_type != "short_circuit_sn":
        braki.append(
            "Zwarcia w punkcie przyłączenia: wskazany przebieg nie jest zwarciowy "
            f"(short_circuit_sn); otrzymano rodzaj analizy: {sc_run.analysis_type}."
        )
    elif sc_run.status != "FINISHED":
        braki.append(
            f"Zwarcia w punkcie przyłączenia: przebieg zwarciowy {sc_run.id} nie "
            f"jest zakończony (status={sc_run.status})."
        )
    elif _sc_row_for_bus(sc_run, bus_ref) is None:
        braki.append(
            f"Zwarcia w punkcie przyłączenia: węzeł „{bus_ref}” nie występuje "
            "w wynikach zwarciowych wskazanego przebiegu."
        )

    # 3. Zgodność NC RfG — braki modułów (ta sama ścieżka co certyfikat D14).
    try:
        build_certyfikat_view(ncrfg_run_result, nazwa_projektu="—")
    except CertyfikatBrakiError as exc:
        braki.extend(f"Zgodność NC RfG: {b}" for b in exc.braki)

    return braki


def _bilans_mocy_sekcja(pf_run: CanonicalRun, bus_ref: str) -> dict[str, Any]:
    """Zbuduj sekcję bilansu mocy z widoku walidacji energetycznej + snapshotu."""
    view = build_energy_validation_view(pf_run)
    items = view.get("items", [])
    summary = view.get("summary", {})

    obciazenia = [
        item
        for item in items
        if item.get("check_type") in {"BRANCH_LOADING", "TRANSFORMER_LOADING"}
        and item.get("observed_value") is not None
    ]
    najwyzsze = max(obciazenia, key=lambda i: i["observed_value"], default=None)
    reactive = next((i for i in items if i.get("check_type") == "REACTIVE_BALANCE"), None)
    straty = next((i for i in items if i.get("check_type") == "LOSS_BUDGET"), None)

    installed_by_bus = _installed_mva_by_bus(pf_run.snapshot or {})
    moc_calkowita = round(sum(installed_by_bus.values()), 6) if installed_by_bus else 0.0
    moc_w_punkcie = installed_by_bus.get(bus_ref)

    return {
        "moc_zainstalowana_zrodel_mva": moc_calkowita,
        "moc_zainstalowana_w_punkcie_mva": (
            round(moc_w_punkcie, 6) if moc_w_punkcie is not None else None
        ),
        "liczba_wezlow_ze_zrodlami": len(installed_by_bus),
        "obciazenie_najwyzsze_pct": (round(najwyzsze["observed_value"], 4) if najwyzsze else None),
        "obciazenie_element": (
            (najwyzsze.get("target_name") or najwyzsze.get("target_id")) if najwyzsze else None
        ),
        "wspolczynnik_mocy_slack": (
            round(reactive["observed_value"], 4)
            if reactive and reactive.get("observed_value") is not None
            else None
        ),
        "bilans_q_status": reactive.get("status") if reactive else "NOT_COMPUTED",
        "straty_pct": (
            round(straty["observed_value"], 4)
            if straty and straty.get("observed_value") is not None
            else None
        ),
        "straty_status": straty.get("status") if straty else "NOT_COMPUTED",
        "podsumowanie_walidacji": {
            "spelnione": summary.get("pass_count", 0),
            "ostrzezenia": summary.get("warning_count", 0),
            "niespelnione": summary.get("fail_count", 0),
            "nieobliczone": summary.get("not_computed_count", 0),
        },
    }


def _zwarcia_sekcja(sc_run: CanonicalRun, bus_ref: str) -> dict[str, Any]:
    """Zbuduj sekcję zwarć w punkcie przyłączenia z wiersza wyniku zwarciowego."""
    row = _sc_row_for_bus(sc_run, bus_ref)
    assert row is not None  # gwarantowane przez bramkę braków
    return {
        "bus_ref": bus_ref,
        "nazwa_wezla": row.get("target_name") or bus_ref,
        "ik_ss_ka": row.get("ikss_ka"),
        "sk_mva": row.get("sk_mva"),
        "ip_ka": row.get("ip_ka"),
        "ith_ka": row.get("ith_ka"),
        "rodzaj_zwarcia": row.get("fault_type"),
    }


def _zgodnosc_sekcja(
    ncrfg_run_result: NcRfgPtpireeRunResult,
    dowody: Sequence[NcRfgCertificateEvidence] | None,
) -> dict[str, Any]:
    """Zbuduj sekcję zgodności NC RfG z werdyktu zbiorczego certyfikatu (D14).

    Z dowodami (wskazany przypadek) sekcja niesie dodatkowo ``dowody_certyfikatu``
    — po jednej pozycji na moduł biegu, w jego kolejności. Bez dowodów klucz nie
    powstaje, więc odcisk sekcji jest identyczny jak przed dodaniem dowodu.
    """
    certyfikat = build_certyfikat_view(ncrfg_run_result, nazwa_projektu="—")
    werdykt = certyfikat["werdykt_zbiorczy"]
    sekcja: dict[str, Any] = {
        "status": werdykt["status"],
        "etykieta_pl": werdykt["etykieta_pl"],
        "liczba_modulow": werdykt["liczba_modulow"],
        "modulow_zgodnych": werdykt["modulow_zgodnych"],
        "modulow_niezgodnych": werdykt["modulow_niezgodnych"],
        "procedura": ncrfg_run_result.procedure_version,
        "odeslanie_pl": (
            "Szczegółowe werdykty testów zdolności podano w certyfikacie zgodności "
            "NC RfG (dokument odrębny, ta sama ścieżka obliczeniowa)."
        ),
        "odcisk_wejscia_nc_rfg_sha256": ncrfg_run_result.input_hash,
    }
    if dowody is not None:
        sekcja["dowody_certyfikatu"] = sekcje_dowodow(dowody)
    return sekcja


def build_wniosek_osd_view(
    pf_run: CanonicalRun,
    sc_run: CanonicalRun,
    ncrfg_run_result: NcRfgPtpireeRunResult,
    *,
    bus_ref: str,
    identyfikacja: WniosekOsdIdentyfikacja,
    dowody: Sequence[NcRfgCertificateEvidence] | None = None,
) -> dict[str, Any]:
    """Zbuduj widok JSON wniosku OSD z istniejących wyników.

    Rzuca ``WniosekOsdBrakiError`` gdy dane są niekompletne (bramka kompletności).

    ``dowody`` (opcjonalne) to dowód certyfikacji PTPiREE per moduł z tabliczek
    modelu. Bez dowodów (wywołanie bez wskazanego przypadku) widok jest IDENTYCZNY
    jak przed dodaniem sekcji — łącznie z odciskami sekcji i ``input_hash``.
    """
    braki = zbierz_braki_wniosku(pf_run, sc_run, bus_ref, ncrfg_run_result)
    if braki:
        raise WniosekOsdBrakiError(braki)

    bilans = _bilans_mocy_sekcja(pf_run, bus_ref)
    zwarcia = _zwarcia_sekcja(sc_run, bus_ref)
    zgodnosc = _zgodnosc_sekcja(ncrfg_run_result, dowody)

    zalozenia_pl = [
        "Wniosek zestawia gotowe wyniki obliczeń — nie przelicza żadnej wielkości "
        "i nie zastępuje uzgodnień z operatorem systemu dystrybucyjnego.",
        f"Bilans mocy pochodzi z przebiegu rozpływu mocy (run_id: {pf_run.id}).",
        "Zwarcia w punkcie przyłączenia pochodzą z przebiegu zwarciowego "
        f"(run_id: {sc_run.id}).",
        "Zgodność NC RfG wyznaczono tą samą ścieżką co certyfikat zgodności "
        f"(odcisk wejścia: {ncrfg_run_result.input_hash}).",
        _ADNOTACJA_SCHEMAT,
        _ADNOTACJA_ZESTAWIENIA,
    ]

    odciski_sekcji = {
        "bilans_mocy": _odcisk(bilans),
        "zwarcia_punkt_przylaczenia": _odcisk(zwarcia),
        "zgodnosc_nc_rfg": _odcisk(zgodnosc),
    }
    input_hash = _odcisk(
        {
            "pf_run_id": str(pf_run.id),
            "sc_run_id": str(sc_run.id),
            "bus_ref": bus_ref,
            "nc_rfg_input_hash": ncrfg_run_result.input_hash,
            "identyfikacja": identyfikacja.model_dump(mode="json"),
        }
    )

    return {
        "kontrakt": WNIOSEK_OSD_CONTRACT,
        "tytul": WNIOSEK_OSD_TYTUL,
        "identyfikacja": {
            "projekt": identyfikacja.nazwa_projektu,
            "przypadek": identyfikacja.nazwa_przypadku,
            "wnioskodawca": identyfikacja.wnioskodawca,
            "adres_przylaczenia": identyfikacja.adres_przylaczenia,
            "wezel_przylaczenia": bus_ref,
        },
        "bilans_mocy": bilans,
        "zwarcia_punkt_przylaczenia": zwarcia,
        "zgodnosc_nc_rfg": zgodnosc,
        "zalozenia_pl": zalozenia_pl,
        "odciski_sekcji_sha256": odciski_sekcji,
        "zrodla": {
            "pf_run_id": str(pf_run.id),
            "sc_run_id": str(sc_run.id),
            "nc_rfg_input_hash": ncrfg_run_result.input_hash,
        },
        "input_hash": input_hash,
    }


def _status_pl(status: str) -> str:
    """Etykieta polska statusu walidacji energetycznej."""
    return {
        "PASS": "spełnia",
        "WARNING": "ostrzeżenie",
        "FAIL": "nie spełnia",
        "NOT_COMPUTED": "brak danych",
    }.get(status, status)


def _fmt(value: Any, unit: str = "") -> str:
    """Sformatuj wartość liczbową dla DOCX (myślnik gdy brak)."""
    if value is None:
        return "—"
    return f"{value} {unit}".strip()


def render_wniosek_osd_docx(view: dict) -> bytes:
    """Zrenderuj deterministyczny DOCX wniosku OSD z widoku JSON.

    Zwraca bajty znormalizowane przez ``make_docx_bytes_deterministic`` — dwa
    wywołania na tym samym widoku dają identyczne bajty.
    """
    if not _DOCX_AVAILABLE:  # pragma: no cover
        raise ImportError("Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx")

    doc = Document()

    tytul = doc.add_heading(view["tytul"], level=0)
    tytul.alignment = WD_ALIGN_PARAGRAPH.CENTER

    identyfikacja = view["identyfikacja"]
    id_para = doc.add_paragraph()
    id_para.add_run("Projekt: ").bold = True
    id_para.add_run(str(identyfikacja["projekt"]))
    if identyfikacja.get("przypadek"):
        id_para.add_run("  |  Przypadek: ").bold = True
        id_para.add_run(str(identyfikacja["przypadek"]))
    if identyfikacja.get("wnioskodawca"):
        wn_para = doc.add_paragraph()
        wn_para.add_run("Wnioskodawca: ").bold = True
        wn_para.add_run(str(identyfikacja["wnioskodawca"]))
    if identyfikacja.get("adres_przylaczenia"):
        ad_para = doc.add_paragraph()
        ad_para.add_run("Adres przyłączenia: ").bold = True
        ad_para.add_run(str(identyfikacja["adres_przylaczenia"]))
    wezel_para = doc.add_paragraph()
    wezel_para.add_run("Węzeł przyłączenia: ").bold = True
    wezel_para.add_run(str(identyfikacja["wezel_przylaczenia"]))

    # Sekcja 1 — bilans mocy.
    bilans = view["bilans_mocy"]
    doc.add_heading("1. Bilans mocy", level=1)
    doc.add_paragraph(
        "Moc zainstalowana źródeł: "
        f"{_fmt(bilans['moc_zainstalowana_zrodel_mva'], 'MVA')}"
        "  (w węźle przyłączenia: "
        f"{_fmt(bilans['moc_zainstalowana_w_punkcie_mva'], 'MVA')})"
    )
    doc.add_paragraph(
        "Najwyższe obciążenie elementu: "
        f"{_fmt(bilans['obciazenie_najwyzsze_pct'], '%')}"
        f" ({bilans['obciazenie_element'] or '—'})"
    )
    doc.add_paragraph(
        "Współczynnik mocy w punkcie bilansowym: "
        f"{_fmt(bilans['wspolczynnik_mocy_slack'])}"
        f" — {_status_pl(bilans['bilans_q_status'])}"
    )
    doc.add_paragraph(
        f"Straty sieciowe: {_fmt(bilans['straty_pct'], '%')}"
        f" — {_status_pl(bilans['straty_status'])}"
    )

    # Sekcja 2 — zwarcia w punkcie przyłączenia.
    zwarcia = view["zwarcia_punkt_przylaczenia"]
    doc.add_heading("2. Zwarcia w punkcie przyłączenia", level=1)
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Table Grid"
    hdr = tabela.rows[0].cells
    hdr[0].text = "Wielkość"
    hdr[1].text = "Wartość"
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].paragraphs[0].runs[0].bold = True
    wiersze = [
        ("Węzeł", zwarcia["nazwa_wezla"]),
        ("Początkowy prąd zwarciowy Ik'' [kA]", _fmt(zwarcia["ik_ss_ka"])),
        ("Moc zwarciowa S_k'' [MVA]", _fmt(zwarcia["sk_mva"])),
        ("Prąd udarowy ip [kA]", _fmt(zwarcia["ip_ka"])),
        ("Prąd cieplny Ith [kA]", _fmt(zwarcia["ith_ka"])),
        ("Rodzaj zwarcia", zwarcia["rodzaj_zwarcia"] or "—"),
    ]
    for etykieta, wartosc in wiersze:
        row = tabela.add_row().cells
        row[0].text = str(etykieta)
        row[1].text = str(wartosc)

    # Sekcja 3 — zgodność NC RfG.
    zgodnosc = view["zgodnosc_nc_rfg"]
    doc.add_heading("3. Zgodność z wymaganiami NC RfG", level=1)
    z_para = doc.add_paragraph()
    z_para.add_run("Werdykt zbiorczy: ").bold = True
    z_para.add_run(str(zgodnosc["etykieta_pl"]))
    doc.add_paragraph(
        f"Moduły: {zgodnosc['liczba_modulow']} "
        f"(zgodnych: {zgodnosc['modulow_zgodnych']}, "
        f"niezgodnych: {zgodnosc['modulow_niezgodnych']})"
    )
    doc.add_paragraph(str(zgodnosc["odeslanie_pl"]))

    dowody = zgodnosc.get("dowody_certyfikatu")
    if dowody is not None:
        doc.add_heading(TYTUL_DOWODU, level=2)
        for dowod in dowody:
            dow_para = doc.add_paragraph()
            dow_para.add_run(f"{dowod['der_ref']}: ").bold = True
            wiersze = wiersze_dowodu_pl(dowod)
            if wiersze:
                dow_para.add_run(
                    "  |  ".join(f"{etykieta}: {wartosc}" for etykieta, wartosc in wiersze)
                )
            else:
                dow_para.add_run(str(dowod["stan_pl"]))

    # Założenia i źródła.
    doc.add_heading("Założenia i źródła", level=1)
    for pozycja in view["zalozenia_pl"]:
        doc.add_paragraph(pozycja, style="List Bullet")

    # Stopka — odciski źródeł.
    doc.add_paragraph()
    stopka = doc.add_paragraph()
    stopka.add_run(f"Odcisk SHA-256 wejścia wniosku: {view['input_hash']}").font.size = Pt(8)
    odciski = view["odciski_sekcji_sha256"]
    for nazwa, odcisk in odciski.items():
        p = doc.add_paragraph()
        p.add_run(f"Odcisk sekcji {nazwa}: {odcisk}").font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    return make_docx_bytes_deterministic(buffer.getvalue())


def render_wniosek_pdf(view: dict) -> bytes:
    """Zrenderuj deterministyczny PDF wniosku OSD z widoku JSON (układ 1:1 z DOCX).

    Determinizm bajtowy: canvas z ``invariant=1`` (stały ``CreationDate`` i ``ID``
    dokumentu) oraz ``pageCompression=0`` — dwa wywołania na tym samym widoku dają
    identyczne bajty. Czcionki DejaVu Sans (polskie diakrytyki) rejestrowane wspólnym
    modułem ``network_model.reporting.czcionki`` (subset TTF jest deterministyczny).
    """
    if not _PDF_AVAILABLE:  # pragma: no cover
        raise ImportError("Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab")

    buffer = BytesIO()
    zarejestruj_czcionki()
    c = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 25 * mm
    right_edge = page_width - 25 * mm
    top_margin = page_height - 25 * mm
    bottom_margin = 25 * mm
    line_height = 5 * mm
    content_width = right_edge - left_margin

    y = top_margin

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = top_margin

    def ensure(needed: float) -> None:
        nonlocal y
        if y - needed < bottom_margin:
            new_page()

    def para(text: str, *, size: int = 10, bold: bool = False, indent: float = 0.0) -> None:
        nonlocal y
        font = "DejaVuSans-Bold" if bold else "DejaVuSans"
        for line in simpleSplit(text, font, size, content_width - indent):
            ensure(line_height)
            c.setFont(font, size)
            c.setFillColorRGB(0, 0, 0)
            c.drawString(left_margin + indent, y, line)
            y -= line_height

    # Tytuł (wyśrodkowany).
    c.setFont("DejaVuSans-Bold", 16)
    c.drawCentredString(page_width / 2, y, str(view["tytul"]))
    y -= 10 * mm

    # Identyfikacja.
    identyfikacja = view["identyfikacja"]
    id_line = f"Projekt: {identyfikacja['projekt']}"
    if identyfikacja.get("przypadek"):
        id_line += f"  |  Przypadek: {identyfikacja['przypadek']}"
    para(id_line)
    if identyfikacja.get("wnioskodawca"):
        para(f"Wnioskodawca: {identyfikacja['wnioskodawca']}")
    if identyfikacja.get("adres_przylaczenia"):
        para(f"Adres przyłączenia: {identyfikacja['adres_przylaczenia']}")
    para(f"Węzeł przyłączenia: {identyfikacja['wezel_przylaczenia']}")
    y -= line_height

    # Sekcja 1 — bilans mocy.
    bilans = view["bilans_mocy"]
    para("1. Bilans mocy", size=12, bold=True)
    para(
        "Moc zainstalowana źródeł: "
        f"{_fmt(bilans['moc_zainstalowana_zrodel_mva'], 'MVA')}"
        "  (w węźle przyłączenia: "
        f"{_fmt(bilans['moc_zainstalowana_w_punkcie_mva'], 'MVA')})"
    )
    para(
        "Najwyższe obciążenie elementu: "
        f"{_fmt(bilans['obciazenie_najwyzsze_pct'], '%')}"
        f" ({bilans['obciazenie_element'] or '—'})"
    )
    para(
        "Współczynnik mocy w punkcie bilansowym: "
        f"{_fmt(bilans['wspolczynnik_mocy_slack'])}"
        f" — {_status_pl(bilans['bilans_q_status'])}"
    )
    para(
        f"Straty sieciowe: {_fmt(bilans['straty_pct'], '%')}"
        f" — {_status_pl(bilans['straty_status'])}"
    )
    y -= line_height

    # Sekcja 2 — zwarcia w punkcie przyłączenia (tabela).
    zwarcia = view["zwarcia_punkt_przylaczenia"]
    para("2. Zwarcia w punkcie przyłączenia", size=12, bold=True)
    wiersze = [
        ("Węzeł", zwarcia["nazwa_wezla"]),
        ("Początkowy prąd zwarciowy Ik'' [kA]", _fmt(zwarcia["ik_ss_ka"])),
        ("Moc zwarciowa S_k'' [MVA]", _fmt(zwarcia["sk_mva"])),
        ("Prąd udarowy ip [kA]", _fmt(zwarcia["ip_ka"])),
        ("Prąd cieplny Ith [kA]", _fmt(zwarcia["ith_ka"])),
        ("Rodzaj zwarcia", zwarcia["rodzaj_zwarcia"] or "—"),
    ]
    label_col = content_width * 0.6
    for etykieta, wartosc in wiersze:
        ensure(line_height)
        c.setFont("DejaVuSans", 10)
        c.setFillColorRGB(0, 0, 0)
        c.drawString(left_margin, y, str(etykieta))
        c.drawString(left_margin + label_col, y, str(wartosc))
        y -= line_height
    y -= line_height

    # Sekcja 3 — zgodność NC RfG.
    zgodnosc = view["zgodnosc_nc_rfg"]
    para("3. Zgodność z wymaganiami NC RfG", size=12, bold=True)
    para(f"Werdykt zbiorczy: {zgodnosc['etykieta_pl']}", bold=True)
    para(
        f"Moduły: {zgodnosc['liczba_modulow']} "
        f"(zgodnych: {zgodnosc['modulow_zgodnych']}, "
        f"niezgodnych: {zgodnosc['modulow_niezgodnych']})"
    )
    para(str(zgodnosc["odeslanie_pl"]))
    dowody = zgodnosc.get("dowody_certyfikatu")
    if dowody is not None:
        para(TYTUL_DOWODU, size=10, bold=True)
        for dowod in dowody:
            wiersze = wiersze_dowodu_pl(dowod)
            tresc = (
                "  |  ".join(f"{etykieta}: {wartosc}" for etykieta, wartosc in wiersze)
                if wiersze
                else str(dowod["stan_pl"])
            )
            para(f"{dowod['der_ref']}: {tresc}", size=9, indent=4 * mm)
    y -= line_height

    # Założenia i źródła.
    para("Założenia i źródła", size=12, bold=True)
    for pozycja in view["zalozenia_pl"]:
        para(f"• {pozycja}", indent=4 * mm)

    # Stopka — odciski źródeł.
    y -= line_height
    para(f"Odcisk SHA-256 wejścia wniosku: {view['input_hash']}", size=8)
    for nazwa, odcisk in view["odciski_sekcji_sha256"].items():
        para(f"Odcisk sekcji {nazwa}: {odcisk}", size=8)

    c.showPage()
    c.save()
    return buffer.getvalue()
