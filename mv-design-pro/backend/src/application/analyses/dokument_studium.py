"""Serwis aplikacyjny: dokument studium przyłączeniowego OZE (E13).

Warstwa APPLICATION — serwerowa KOMPOZYCJA tej samej sekwencji, którą frontend
liczy w kreatorze studium (``ui2/oze/studium/studiumModel.ts``): dla każdego
wariantu (węzła przyłączenia) kolejno zdolność przyłączeniowa → obszar pracy
P–Q → pokrycie wymagań P–Q. Serwisy wołane BEZPOŚREDNIO (nie przez HTTP):

- ``application.analyses.hosting_capacity.build_hosting_capacity_view``,
- ``application.analyses.pq_area.build_pq_area_view``,
- ``application.analyses.pq_coverage.build_pq_coverage_view``.

ZERO nowej fizyki i ZERO nowych ocen — dokument tylko zestawia gotowe widoki.

Klasa NC RfG wariantu: użyta ISTNIEJĄCA klasyfikacja backendowa
``NcRfgProfile.classify_module`` (``catalog/profiles/nc_rfg/loader.py``) —
ZERO dwóch prawd. Frontendowy ``klasaNcRfg`` (``ranking/rankingModel.ts``) jest
udokumentowanym lustrem 1:1 tej samej klasyfikacji; źródłem prawdy jest backend,
więc dokument nie dubluje mapowania — wywołuje istniejącą funkcję. Klasa liczona
jak w kreatorze studium: moc przyłączalna wariantu [MW→kW] + napięcie węzła [kV].

Bramka braków twardych (przed generacją, kolejność deterministyczna: przebieg →
typ katalogowy → operator → warianty): przebieg złego rodzaju lub niezakończony,
nieznany typ katalogowy przekształtnika, nieznany profil operatora, pusta lista
wariantów. Przy jakimkolwiek braku dokument NIE powstaje — zwracana jest lista
po polsku (422).

Błąd pojedynczego wariantu (np. węzeł spoza modelu, typ bez krzywej P–Q) NIE
przerywa dokumentu — sekcja wariantu niesie uczciwy opis błędu po polsku,
a pozostałe warianty są liczone dalej (wzorzec runnera frontendu). Kolejność
wariantów jest zachowana (kolejność wyboru).

Determinizm: identyczne wejście → identyczny widok JSON, identyczne odciski
sekcji, identyczny ``input_hash`` i bajtowo identyczne eksporty DOCX oraz PDF.
"""

from __future__ import annotations

import hashlib
import json
from collections.abc import Sequence
from io import BytesIO
from typing import Any

from application.analyses.dowod_certyfikatu import (
    BRAK_URZADZEN_TYPU_PL,
    TYTUL_DOWODU,
    NcRfgCertificateEvidence,
    sekcje_dowodow,
    wiersze_dowodu_pl,
)
from application.analyses.hosting_capacity import build_hosting_capacity_view
from application.analyses.pq_area import build_pq_area_view
from application.analyses.pq_coverage import build_pq_coverage_view
from catalog.profiles.nc_rfg.loader import NcRfgProfile
from enm.canonical_analysis import CanonicalRun
from network_model.catalog.types import ConverterType
from network_model.reporting.czcionki import zarejestruj_czcionki
from network_model.reporting.docx_determinism import make_docx_bytes_deterministic
from pydantic import BaseModel, Field

try:  # pragma: no cover - zależy od środowiska
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False

try:  # pragma: no cover - zależy od środowiska
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfgen import canvas

    _PDF_AVAILABLE = True
except ImportError:  # pragma: no cover
    _PDF_AVAILABLE = False

DOKUMENT_STUDIUM_CONTRACT = "DokumentStudiumPrzylaczeniowegoV1"
DOKUMENT_STUDIUM_TYTUL = "Dokument studium przyłączeniowego OZE"


class DokumentStudiumIdentyfikacja(BaseModel):
    """Dane identyfikacyjne dokumentu (pola tekstowe — opcjonalne poza nazwą)."""

    nazwa_projektu: str = Field(min_length=1)
    nazwa_przypadku: str | None = None
    wnioskodawca: str | None = None
    adres_przylaczenia: str | None = None


class DokumentStudiumBrakiError(Exception):
    """Dokument nie powstaje z powodu braków twardych (lista po polsku)."""

    def __init__(self, braki: list[str]) -> None:
        self.braki = braki
        super().__init__("Dokument studium nie może powstać — dane wejściowe niekompletne.")


def _odcisk(payload: Any) -> str:
    """Deterministyczny odcisk SHA-256 sekcji (kanoniczny JSON, sortowane klucze)."""
    canonical = json.dumps(payload, sort_keys=True, ensure_ascii=False, default=str)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def zbierz_braki_dokumentu(
    run: CanonicalRun,
    converter: ConverterType | None,
    profile: NcRfgProfile | None,
    *,
    catalog_item_id: str,
    operator_id: str,
    warianty: list[str],
) -> list[str]:
    """Zbierz braki twarde blokujące generację (kolejność deterministyczna).

    Kolejność: przebieg bazowy → typ katalogowy → profil operatora → warianty.
    """
    braki: list[str] = []

    if run.analysis_type != "PF":
        braki.append(
            "Przebieg bazowy: wskazany przebieg nie jest rozpływem mocy (PF); "
            f"otrzymano rodzaj analizy: {run.analysis_type}."
        )
    elif run.status != "FINISHED":
        braki.append(
            f"Przebieg bazowy: przebieg {run.id} nie jest zakończony (status={run.status})."
        )

    if converter is None:
        braki.append(
            f"Typ katalogowy: typ „{catalog_item_id}” nie istnieje " "w katalogu przekształtników."
        )

    if profile is None:
        braki.append(f"Profil operatora: operator „{operator_id}” nie ma profilu NC RfG.")

    if not warianty:
        braki.append("Warianty: nie wskazano żadnego wariantu (węzła przyłączenia).")

    return braki


def _bus_voltage_index(snapshot: dict[str, Any]) -> dict[str, float]:
    """Mapa ref_id → napięcie znamionowe [kV] węzłów ze snapshotu przebiegu."""
    index: dict[str, float] = {}
    for bus in snapshot.get("buses") or []:
        ref = bus.get("ref_id")
        napiecie = bus.get("voltage_kv")
        if ref is not None and napiecie is not None:
            index[str(ref)] = float(napiecie)
    return index


def _bus_name_index(snapshot: dict[str, Any]) -> dict[str, str | None]:
    return {
        str(bus["ref_id"]): bus.get("name")
        for bus in (snapshot.get("buses") or [])
        if bus.get("ref_id")
    }


def _ograniczenie_pl(binding: dict[str, Any]) -> str:
    """Opis PL kryterium wiążącego zdolności przyłączeniowej (wzorzec D3)."""
    kind = binding.get("kind")
    if kind == "none":
        return "Brak ograniczenia w zakresie przeglądu."
    if kind == "non_convergence":
        return "Rozpływ mocy nie osiągnął zbieżności."
    element = binding.get("element_name") or binding.get("element_id") or "—"
    if kind == "voltage":
        return f"Kryterium napięciowe — węzeł „{element}”."
    if kind == "loading":
        return f"Kryterium obciążeniowe — element „{element}”."
    return "Ograniczenie nieokreślone."


def _wierzcholek_najblizszy(
    wierzcholki: list[dict[str, Any]], p_docelowe_mw: float
) -> dict[str, Any] | None:
    """Wierzchołek obszaru P–Q najbliższy zadanej mocy [MW] (bez interpolacji).

    Remis (równa odległość) rozstrzyga niższy indeks (deterministycznie), zgodnie
    z ``wierzcholekNajblizszy`` w ``ui2/oze/studium/studiumModel.ts``.
    """
    if not wierzcholki:
        return None
    najlepszy = wierzcholki[0]
    najlepsza_odleglosc = abs(najlepszy["p_mw"] - p_docelowe_mw)
    for wierzcholek in wierzcholki[1:]:
        odleglosc = abs(wierzcholek["p_mw"] - p_docelowe_mw)
        if odleglosc < najlepsza_odleglosc:
            najlepszy = wierzcholek
            najlepsza_odleglosc = odleglosc
    return najlepszy


def _pasmo_q_pl(pq_view: dict[str, Any], moc_zrodla_mw: float) -> str:
    """Pasmo Q [Mvar] w punkcie mocy źródła — odczyt wierzchołka najbliższego mocy.

    Odwzorowuje ``pasmoQwPunkcie`` z kreatora studium: brak danych/pasma → tekst
    (bez oceny lokalnej).
    """
    wierzcholek = _wierzcholek_najblizszy(pq_view.get("vertices") or [], moc_zrodla_mw)
    if wierzcholek is None:
        return "—"
    if (
        not wierzcholek.get("feasible")
        or wierzcholek.get("q_min_dop_mvar") is None
        or wierzcholek.get("q_max_dop_mvar") is None
    ):
        return "brak pasma pracy"
    return f"{wierzcholek['q_min_dop_mvar']}…{wierzcholek['q_max_dop_mvar']} Mvar"


def _klasa_nc_rfg(
    profile: NcRfgProfile, max_moc_mw: float | None, napiecie_kv: float | None
) -> dict[str, Any]:
    """Klasa NC RfG wariantu z ISTNIEJĄCEJ klasyfikacji backendowej (art. 5).

    Wywołuje ``NcRfgProfile.classify_module`` (jedyne źródło prawdy). Moc
    przyłączalną wariantu przelicza MW→kW (jak kreator studium). Brak dodatniej
    mocy przyłączalnej → klasa nieokreślona (bez zgadywania).
    """
    if max_moc_mw is None or max_moc_mw <= 0.0:
        return {
            "klasa": None,
            "opis_pl": "Klasa nieokreślona — brak dodatniej mocy przyłączalnej.",
        }
    p_max_kw = max_moc_mw * 1000.0
    modul = profile.classify_module(p_max_kw, napiecie_kv if napiecie_kv is not None else 0.0)
    if modul is None:
        return {
            "klasa": None,
            "opis_pl": "Klasa nieokreślona — profil operatora nie zawiera kategorii.",
        }
    return {"klasa": modul.id, "opis_pl": modul.description_pl}


def _wariant_sekcja(
    run: CanonicalRun,
    converter: ConverterType,
    profile: NcRfgProfile,
    bus_ref: str,
    *,
    bus_names: dict[str, str | None],
    bus_voltages: dict[str, float],
) -> dict[str, Any]:
    """Zbuduj sekcję jednego wariantu (trzy fazy odporne na błąd pojedynczej fazy)."""
    napiecie_kv = bus_voltages.get(bus_ref)

    # Faza 1 — zdolność przyłączeniowa (hosting-capacity) dla wariantu.
    max_moc_mw: float | None = None
    try:
        hc_view = build_hosting_capacity_view(run, candidate_bus_refs=[bus_ref])
        node = hc_view["nodes"][0]
        max_moc_mw = node["max_hosting_capacity_mw"]
        zdolnosc = {
            "status": "ok",
            "max_moc_mw": max_moc_mw,
            "ograniczenie_pl": _ograniczenie_pl(node["binding_criterion"]),
            "komunikat_bledu": None,
        }
    except Exception as exc:  # noqa: BLE001 — błąd wariantu nie przerywa dokumentu
        zdolnosc = {
            "status": "blad",
            "max_moc_mw": None,
            "ograniczenie_pl": None,
            "komunikat_bledu": str(exc),
        }

    # Faza 2 — obszar bezpiecznej pracy P–Q (pq-area, parametry domyślne backendu).
    try:
        pq_view = build_pq_area_view(run, bus_ref=bus_ref)
        obszar = {
            "status": "ok",
            "pasmo_q_pl": _pasmo_q_pl(pq_view, float(converter.pmax_mw)),
            "liczba_wierzcholkow": len(pq_view.get("vertices") or []),
            "komunikat_bledu": None,
        }
    except Exception as exc:  # noqa: BLE001
        obszar = {
            "status": "blad",
            "pasmo_q_pl": None,
            "liczba_wierzcholkow": None,
            "komunikat_bledu": str(exc),
        }

    # Faza 3 — pokrycie wymagania operatora (pq-coverage, typ + operator).
    try:
        cov_view = build_pq_coverage_view(converter, profile)
        werdykt = cov_view["werdykt"]
        pokrycie = {
            "status": "ok",
            "pokryty": werdykt["pokryty"],
            "werdykt_pl": ("Pokryte" if werdykt["pokryty"] else "Niepokryte"),
            "opis_pl": werdykt["opis_pl"],
            "komunikat_bledu": None,
        }
    except Exception as exc:  # noqa: BLE001
        pokrycie = {
            "status": "blad",
            "pokryty": None,
            "werdykt_pl": None,
            "opis_pl": None,
            "komunikat_bledu": str(exc),
        }

    klasa = _klasa_nc_rfg(profile, max_moc_mw, napiecie_kv)

    sekcja = {
        "bus_ref": bus_ref,
        "nazwa_wezla": bus_names.get(bus_ref) or bus_ref,
        "napiecie_kv": napiecie_kv,
        "zdolnosc": zdolnosc,
        "obszar_pq": obszar,
        "pokrycie_pq": pokrycie,
        "klasa_nc_rfg": klasa,
    }
    sekcja["odcisk_sekcji_sha256"] = _odcisk(sekcja)
    return sekcja


def _podsumowanie_wariant(sekcja: dict[str, Any]) -> dict[str, Any]:
    """Wiersz podsumowania porównawczego dla jednego wariantu."""
    return {
        "bus_ref": sekcja["bus_ref"],
        "nazwa_wezla": sekcja["nazwa_wezla"],
        "max_moc_mw": sekcja["zdolnosc"]["max_moc_mw"],
        "klasa": sekcja["klasa_nc_rfg"]["klasa"],
        "pokrycie_pl": sekcja["pokrycie_pq"]["werdykt_pl"],
        "pasmo_q_pl": sekcja["obszar_pq"]["pasmo_q_pl"],
    }


def build_dokument_studium_view(
    run: CanonicalRun,
    converter: ConverterType | None,
    profile: NcRfgProfile | None,
    *,
    catalog_item_id: str,
    operator_id: str,
    warianty: list[str],
    identyfikacja: DokumentStudiumIdentyfikacja,
    dowody: Sequence[NcRfgCertificateEvidence] | None = None,
) -> dict[str, Any]:
    """Zbuduj widok JSON dokumentu studium przyłączeniowego.

    Rzuca ``DokumentStudiumBrakiError`` gdy dane wejściowe są niekompletne
    (bramka braków twardych przed generacją).

    ``dowody`` (opcjonalne) to dowód certyfikacji PTPiREE urządzeń modelu
    związanych z TYPEM katalogowym dokumentu (tożsamość urządzenia w studium to
    typ przekształtnika, nie moduł biegu). Bez dowodów (wywołanie bez wskazanego
    przypadku) widok jest IDENTYCZNY jak przed dodaniem sekcji — łącznie
    z odciskiem sekcji założeń i ``input_hash``.
    """
    braki = zbierz_braki_dokumentu(
        run,
        converter,
        profile,
        catalog_item_id=catalog_item_id,
        operator_id=operator_id,
        warianty=warianty,
    )
    if braki:
        raise DokumentStudiumBrakiError(braki)

    assert converter is not None  # gwarantowane przez bramkę braków
    assert profile is not None  # gwarantowane przez bramkę braków

    snapshot = run.snapshot or {}
    bus_names = _bus_name_index(snapshot)
    bus_voltages = _bus_voltage_index(snapshot)

    warianty_view = [
        _wariant_sekcja(
            run,
            converter,
            profile,
            bus_ref,
            bus_names=bus_names,
            bus_voltages=bus_voltages,
        )
        for bus_ref in warianty
    ]

    podsumowanie = [_podsumowanie_wariant(sekcja) for sekcja in warianty_view]

    zalozenia = {
        "typ_katalogowy": {
            "id": converter.id,
            "nazwa": converter.name,
            "kind": converter.kind.value,
            "pmax_mw": float(converter.pmax_mw),
            "sn_mva": float(converter.sn_mva),
        },
        "operator": {
            "id": profile.operator_id,
            "nazwa": profile.operator_name_pl,
        },
        "przebieg_bazowy": {
            "run_id": str(run.id),
            "snapshot_hash": run.snapshot_hash,
        },
        "liczba_wariantow": len(warianty_view),
    }
    if dowody is not None:
        # Brak dopasowanego urządzenia to uczciwy stan zerowy dokumentu: pusta
        # lista + jawny opis, nigdy dowód urządzenia innego typu.
        zalozenia["dowod_certyfikatu"] = {
            "catalog_item_id": catalog_item_id,
            "urzadzenia": sekcje_dowodow(dowody),
            "stan_pl": BRAK_URZADZEN_TYPU_PL if not dowody else None,
        }

    zalozenia_pl = [
        "Dokument zestawia gotowe wyniki obliczeń — nie przelicza żadnej wielkości "
        "i nie zastępuje uzgodnień z operatorem systemu dystrybucyjnego.",
        "Sekwencja per wariant: zdolność przyłączeniowa → obszar pracy P–Q → "
        "pokrycie wymagań P–Q (ta sama, którą liczy kreator studium).",
        f"Przebieg bazowy rozpływu mocy (run_id: {run.id}, odcisk snapshotu: "
        f"{run.snapshot_hash}).",
        f"Typ katalogowy przekształtnika: {converter.name} ({converter.id}).",
        f"Profil operatora: {profile.operator_name_pl} ({profile.operator_id}).",
        "Klasa NC RfG wyznaczona istniejącą klasyfikacją katalogową operatora "
        "(art. 5) na podstawie mocy przyłączalnej wariantu.",
    ]

    odciski_sekcji = {sekcja["bus_ref"]: sekcja["odcisk_sekcji_sha256"] for sekcja in warianty_view}
    odciski_sekcji["zalozenia"] = _odcisk(zalozenia)
    odciski_sekcji["podsumowanie"] = _odcisk(podsumowanie)

    input_hash = _odcisk(
        {
            "snapshot_hash": run.snapshot_hash,
            "catalog_item_id": catalog_item_id,
            "operator_id": operator_id,
            "warianty": warianty,
            "identyfikacja": identyfikacja.model_dump(mode="json"),
        }
    )

    return {
        "kontrakt": DOKUMENT_STUDIUM_CONTRACT,
        "tytul": DOKUMENT_STUDIUM_TYTUL,
        "identyfikacja": {
            "projekt": identyfikacja.nazwa_projektu,
            "przypadek": identyfikacja.nazwa_przypadku,
            "wnioskodawca": identyfikacja.wnioskodawca,
            "adres_przylaczenia": identyfikacja.adres_przylaczenia,
        },
        "zalozenia": zalozenia,
        "warianty": warianty_view,
        "podsumowanie": podsumowanie,
        "zalozenia_pl": zalozenia_pl,
        "odciski_sekcji_sha256": odciski_sekcji,
        "input_hash": input_hash,
    }


def _fmt(value: Any, unit: str = "") -> str:
    """Sformatuj wartość liczbową dla eksportu (myślnik gdy brak)."""
    if value is None:
        return "—"
    return f"{value} {unit}".strip()


def _klasa_tekst(klasa: dict[str, Any]) -> str:
    """Tekst klasy NC RfG na potrzeby eksportu (kod klasy + opis, lub myślnik)."""
    if klasa.get("klasa") is None:
        return "—"
    return f"{klasa['klasa']} — {klasa['opis_pl']}"


def render_dokument_studium_docx(view: dict) -> bytes:
    """Zrenderuj deterministyczny DOCX dokumentu studium z widoku JSON.

    Zwraca bajty znormalizowane przez ``make_docx_bytes_deterministic`` — dwa
    wywołania na tym samym widoku dają identyczne bajty.
    """
    if not _DOCX_AVAILABLE:  # pragma: no cover
        raise ImportError("Eksport DOCX wymaga python-docx. Zainstaluj: pip install python-docx")

    doc = Document()

    tytul = doc.add_heading(view["tytul"], level=0)
    tytul.alignment = WD_ALIGN_PARAGRAPH.CENTER

    identyfikacja = view["identyfikacja"]
    id_para = doc.add_paragraph()
    id_para.add_run("Projekt: ").bold = True
    id_para.add_run(str(identyfikacja["projekt"]))
    if identyfikacja.get("przypadek"):
        id_para.add_run("  |  Przypadek: ").bold = True
        id_para.add_run(str(identyfikacja["przypadek"]))
    if identyfikacja.get("wnioskodawca"):
        wn_para = doc.add_paragraph()
        wn_para.add_run("Wnioskodawca: ").bold = True
        wn_para.add_run(str(identyfikacja["wnioskodawca"]))
    if identyfikacja.get("adres_przylaczenia"):
        ad_para = doc.add_paragraph()
        ad_para.add_run("Adres przyłączenia: ").bold = True
        ad_para.add_run(str(identyfikacja["adres_przylaczenia"]))

    # Założenia.
    zal = view["zalozenia"]
    doc.add_heading("Założenia", level=1)
    doc.add_paragraph(
        f"Typ modułu: {zal['typ_katalogowy']['nazwa']} "
        f"(Pmax {_fmt(zal['typ_katalogowy']['pmax_mw'], 'MW')})"
    )
    doc.add_paragraph(f"Operator: {zal['operator']['nazwa']} ({zal['operator']['id']})")
    doc.add_paragraph(
        f"Przebieg bazowy (rozpływ mocy): {zal['przebieg_bazowy']['run_id']}"
        f"  |  odcisk snapshotu: {zal['przebieg_bazowy']['snapshot_hash']}"
    )
    dowod_blok = zal.get("dowod_certyfikatu")
    if dowod_blok is not None:
        doc.add_heading(TYTUL_DOWODU, level=2)
        if dowod_blok["stan_pl"]:
            doc.add_paragraph(str(dowod_blok["stan_pl"]))
        for dowod in dowod_blok["urzadzenia"]:
            dow_para = doc.add_paragraph()
            dow_para.add_run(f"{dowod['der_ref']}: ").bold = True
            wiersze = wiersze_dowodu_pl(dowod)
            if wiersze:
                dow_para.add_run(
                    "  |  ".join(f"{etykieta}: {wartosc}" for etykieta, wartosc in wiersze)
                )
            else:
                dow_para.add_run(str(dowod["stan_pl"]))

    # Warianty.
    for i, wariant in enumerate(view["warianty"], start=1):
        doc.add_heading(
            f"Wariant {i}: {wariant['nazwa_wezla']} "
            f"(napięcie: {_fmt(wariant['napiecie_kv'], 'kV')})",
            level=1,
        )
        zdolnosc = wariant["zdolnosc"]
        if zdolnosc["status"] == "ok":
            doc.add_paragraph(
                "Zdolność przyłączeniowa: maksymalna moc "
                f"{_fmt(zdolnosc['max_moc_mw'], 'MW')} — {zdolnosc['ograniczenie_pl']}"
            )
        else:
            doc.add_paragraph(f"Zdolność przyłączeniowa: błąd — {zdolnosc['komunikat_bledu']}")
        obszar = wariant["obszar_pq"]
        if obszar["status"] == "ok":
            doc.add_paragraph(f"Obszar pracy P–Q: pasmo Q {obszar['pasmo_q_pl']}")
        else:
            doc.add_paragraph(f"Obszar pracy P–Q: błąd — {obszar['komunikat_bledu']}")
        pokrycie = wariant["pokrycie_pq"]
        if pokrycie["status"] == "ok":
            doc.add_paragraph(
                f"Pokrycie wymagań P–Q: {pokrycie['werdykt_pl']} — {pokrycie['opis_pl']}"
            )
        else:
            doc.add_paragraph(f"Pokrycie wymagań P–Q: błąd — {pokrycie['komunikat_bledu']}")
        doc.add_paragraph(f"Klasa NC RfG: {_klasa_tekst(wariant['klasa_nc_rfg'])}")

    # Podsumowanie porównawcze wariantów (tabela).
    doc.add_heading("Podsumowanie porównawcze wariantów", level=1)
    tabela = doc.add_table(rows=1, cols=6)
    tabela.style = "Table Grid"
    naglowki = ["Wariant", "Węzeł", "Moc [MW]", "Klasa", "Pokrycie", "Pasmo Q"]
    hdr = tabela.rows[0].cells
    for i, tekst in enumerate(naglowki):
        hdr[i].text = tekst
        hdr[i].paragraphs[0].runs[0].bold = True
    for wiersz in view["podsumowanie"]:
        row = tabela.add_row().cells
        row[0].text = str(wiersz["nazwa_wezla"])
        row[1].text = str(wiersz["bus_ref"])
        row[2].text = _fmt(wiersz["max_moc_mw"])
        row[3].text = wiersz["klasa"] if wiersz["klasa"] is not None else "—"
        row[4].text = wiersz["pokrycie_pl"] if wiersz["pokrycie_pl"] is not None else "—"
        row[5].text = wiersz["pasmo_q_pl"] if wiersz["pasmo_q_pl"] is not None else "—"

    # Założenia i źródła.
    doc.add_heading("Założenia i źródła", level=1)
    for pozycja in view["zalozenia_pl"]:
        doc.add_paragraph(pozycja, style="List Bullet")

    # Stopka — odciski.
    doc.add_paragraph()
    stopka = doc.add_paragraph()
    stopka.add_run(f"Odcisk SHA-256 wejścia dokumentu: {view['input_hash']}").font.size = Pt(8)
    for nazwa, odcisk in view["odciski_sekcji_sha256"].items():
        p = doc.add_paragraph()
        p.add_run(f"Odcisk sekcji {nazwa}: {odcisk}").font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    return make_docx_bytes_deterministic(buffer.getvalue())


def render_dokument_studium_pdf(view: dict) -> bytes:
    """Zrenderuj deterministyczny PDF dokumentu studium z widoku JSON (układ 1:1 z DOCX).

    Determinizm bajtowy: canvas z ``invariant=1`` (stały ``CreationDate`` i ``ID``
    dokumentu) oraz ``pageCompression=0`` — dwa wywołania na tym samym widoku dają
    identyczne bajty. Czcionki DejaVu Sans (polskie diakrytyki) rejestrowane wspólnym
    modułem ``network_model.reporting.czcionki`` (subset TTF jest deterministyczny).
    """
    if not _PDF_AVAILABLE:  # pragma: no cover
        raise ImportError("Eksport PDF wymaga reportlab. Zainstaluj: pip install reportlab")

    buffer = BytesIO()
    zarejestruj_czcionki()
    c = canvas.Canvas(buffer, pagesize=A4, invariant=1, pageCompression=0)
    page_width, page_height = A4
    left_margin = 25 * mm
    right_edge = page_width - 25 * mm
    top_margin = page_height - 25 * mm
    bottom_margin = 25 * mm
    line_height = 5 * mm
    content_width = right_edge - left_margin

    y = top_margin

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = top_margin

    def ensure(needed: float) -> None:
        nonlocal y
        if y - needed < bottom_margin:
            new_page()

    def para(text: str, *, size: int = 10, bold: bool = False, indent: float = 0.0) -> None:
        nonlocal y
        font = "DejaVuSans-Bold" if bold else "DejaVuSans"
        for line in simpleSplit(text, font, size, content_width - indent):
            ensure(line_height)
            c.setFont(font, size)
            c.setFillColorRGB(0, 0, 0)
            c.drawString(left_margin + indent, y, line)
            y -= line_height

    # Tytuł (wyśrodkowany).
    c.setFont("DejaVuSans-Bold", 16)
    c.drawCentredString(page_width / 2, y, str(view["tytul"]))
    y -= 10 * mm

    # Identyfikacja.
    identyfikacja = view["identyfikacja"]
    id_line = f"Projekt: {identyfikacja['projekt']}"
    if identyfikacja.get("przypadek"):
        id_line += f"  |  Przypadek: {identyfikacja['przypadek']}"
    para(id_line)
    if identyfikacja.get("wnioskodawca"):
        para(f"Wnioskodawca: {identyfikacja['wnioskodawca']}")
    if identyfikacja.get("adres_przylaczenia"):
        para(f"Adres przyłączenia: {identyfikacja['adres_przylaczenia']}")
    y -= line_height

    # Założenia.
    zal = view["zalozenia"]
    para("Założenia", size=12, bold=True)
    para(
        f"Typ modułu: {zal['typ_katalogowy']['nazwa']} "
        f"(Pmax {_fmt(zal['typ_katalogowy']['pmax_mw'], 'MW')})"
    )
    para(f"Operator: {zal['operator']['nazwa']} ({zal['operator']['id']})")
    para(
        f"Przebieg bazowy (rozpływ mocy): {zal['przebieg_bazowy']['run_id']}"
        f"  |  odcisk snapshotu: {zal['przebieg_bazowy']['snapshot_hash']}"
    )
    dowod_blok = zal.get("dowod_certyfikatu")
    if dowod_blok is not None:
        para(TYTUL_DOWODU, size=10, bold=True)
        if dowod_blok["stan_pl"]:
            para(str(dowod_blok["stan_pl"]), size=9, indent=4 * mm)
        for dowod in dowod_blok["urzadzenia"]:
            wiersze = wiersze_dowodu_pl(dowod)
            tresc = (
                "  |  ".join(f"{etykieta}: {wartosc}" for etykieta, wartosc in wiersze)
                if wiersze
                else str(dowod["stan_pl"])
            )
            para(f"{dowod['der_ref']}: {tresc}", size=9, indent=4 * mm)
    y -= line_height

    # Warianty.
    for i, wariant in enumerate(view["warianty"], start=1):
        ensure(line_height * 5)
        para(
            f"Wariant {i}: {wariant['nazwa_wezla']} "
            f"(napięcie: {_fmt(wariant['napiecie_kv'], 'kV')})",
            size=12,
            bold=True,
        )
        zdolnosc = wariant["zdolnosc"]
        if zdolnosc["status"] == "ok":
            para(
                "Zdolność przyłączeniowa: maksymalna moc "
                f"{_fmt(zdolnosc['max_moc_mw'], 'MW')} — {zdolnosc['ograniczenie_pl']}"
            )
        else:
            para(f"Zdolność przyłączeniowa: błąd — {zdolnosc['komunikat_bledu']}")
        obszar = wariant["obszar_pq"]
        if obszar["status"] == "ok":
            para(f"Obszar pracy P–Q: pasmo Q {obszar['pasmo_q_pl']}")
        else:
            para(f"Obszar pracy P–Q: błąd — {obszar['komunikat_bledu']}")
        pokrycie = wariant["pokrycie_pq"]
        if pokrycie["status"] == "ok":
            para(f"Pokrycie wymagań P–Q: {pokrycie['werdykt_pl']} — {pokrycie['opis_pl']}")
        else:
            para(f"Pokrycie wymagań P–Q: błąd — {pokrycie['komunikat_bledu']}")
        para(f"Klasa NC RfG: {_klasa_tekst(wariant['klasa_nc_rfg'])}")
        y -= line_height

    # Podsumowanie porównawcze wariantów (tabela).
    ensure(line_height * 3)
    para("Podsumowanie porównawcze wariantów", size=12, bold=True)
    kolumny = [
        content_width * 0.24,
        content_width * 0.16,
        content_width * 0.14,
        content_width * 0.10,
        content_width * 0.18,
        content_width * 0.18,
    ]
    naglowki = ["Węzeł", "Identyfikator", "Moc [MW]", "Klasa", "Pokrycie", "Pasmo Q"]
    ensure(line_height)
    c.setFont("DejaVuSans-Bold", 9)
    col_x = left_margin
    for i, label in enumerate(naglowki):
        c.drawString(col_x, y, label)
        col_x += kolumny[i]
    y -= line_height
    for wiersz in view["podsumowanie"]:
        cells = [
            str(wiersz["nazwa_wezla"]),
            str(wiersz["bus_ref"]),
            _fmt(wiersz["max_moc_mw"]),
            wiersz["klasa"] if wiersz["klasa"] is not None else "—",
            wiersz["pokrycie_pl"] if wiersz["pokrycie_pl"] is not None else "—",
            wiersz["pasmo_q_pl"] if wiersz["pasmo_q_pl"] is not None else "—",
        ]
        wrapped = [
            simpleSplit(cell, "DejaVuSans", 9, kolumny[i] - 2 * mm) for i, cell in enumerate(cells)
        ]
        row_lines = max(len(w) for w in wrapped)
        ensure(line_height * row_lines)
        c.setFont("DejaVuSans", 9)
        col_x = left_margin
        for i, lines in enumerate(wrapped):
            for j, line in enumerate(lines):
                c.drawString(col_x, y - j * line_height, line)
            col_x += kolumny[i]
        y -= line_height * row_lines
    y -= line_height

    # Założenia i źródła.
    ensure(line_height * 2)
    para("Założenia i źródła", size=12, bold=True)
    for pozycja in view["zalozenia_pl"]:
        para(f"• {pozycja}", indent=4 * mm)

    # Stopka — odciski.
    y -= line_height
    para(f"Odcisk SHA-256 wejścia dokumentu: {view['input_hash']}", size=8)
    for nazwa, odcisk in view["odciski_sekcji_sha256"].items():
        para(f"Odcisk sekcji {nazwa}: {odcisk}", size=8)

    c.showPage()
    c.save()
    return buffer.getvalue()
