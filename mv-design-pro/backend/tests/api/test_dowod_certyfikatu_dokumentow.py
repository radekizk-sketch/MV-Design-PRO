"""Formalne dokumenty OZE niosą DOWÓD certyfikacji PTPiREE (klasa, nie instancja).

PO CO. Certyfikat zgodności, wniosek o warunki przyłączenia i dokument studium
przyłączeniowego są składane u operatora systemu dystrybucyjnego. Każdy z nich
deklaruje zgodność urządzenia z wymaganiami NC RfG, więc musi dać się z niego
odczytać, NA CO ta deklaracja się powołuje: numer dokumentu WOŚ/WiPWC, data
akceptacji, zakres PPM, adres wpisu w wykazie i — krytyczny — WARUNEK ważności
certyfikatu. Bieg NC RfG cytował te dane od karty dodającej
``certificate_evidence``; trzy dokumenty budowane z tego samego biegu ich nie
niosły, bo nie dostawały wskazania przypadku obliczeniowego.

INWENTARZ KLASY (9 końcówek, 3 dokumenty × JSON/DOCX/PDF):
``/api/oze-analysis/compliance-certificate``, ``…/osd-application``,
``…/connection-study`` (każda w wariancie bazowym oraz ``.docx`` i ``.pdf``).
Dwie pierwsze identyfikują urządzenie referencją modułu biegu (``der_ref``);
dokument studium NIE uruchamia macierzy NC RfG — jego urządzeniem jest TYP
katalogowy przekształtnika, więc dopasowanie idzie po ``catalog_item_id``
tabliczki. Obie drogi czytają tę samą tabliczkę tym samym kodem
(``application.analyses.dowod_certyfikatu``).

ILOCZYN CECH pokryty niżej: dokument (3) × stan dowodu (komplet danych / pusta
tabliczka / urządzenie spoza modelu / bez wskazanego przypadku) × postać wyniku
(widok JSON / DOCX / PDF). Osobno przypięta jest RÓWNOŚĆ kontraktu: żądanie bez
``case_id`` daje widok bajtowo identyczny z tym sprzed dodania sekcji — łącznie
z odciskami sekcji i ``input_hash``.
"""

from __future__ import annotations

import json
import re
from io import BytesIO
from uuid import uuid4

import pytest
from application.analyses.certyfikat_zgodnosci import build_certyfikat_view, render_certyfikat_docx
from application.analyses.dokument_studium import (
    DokumentStudiumIdentyfikacja,
    build_dokument_studium_view,
    render_dokument_studium_docx,
)
from application.analyses.dowod_certyfikatu import (
    BRAK_TABLICZKI_PL,
    BRAK_URZADZEN_TYPU_PL,
    TYTUL_DOWODU,
    dowody_certyfikatu,
    dowody_certyfikatu_typu,
)
from application.analyses.wniosek_osd import (
    WniosekOsdIdentyfikacja,
    build_wniosek_osd_view,
    render_wniosek_osd_docx,
)
from catalog.profiles.nc_rfg.loader import load_nc_rfg_profile
from docx import Document
from enm.canonical_analysis import CanonicalRun, create_run, execute_run, reset_canonical_runs
from enm.models import Bus, EnergyNetworkModel, ENMHeader, Generator
from enm.store import reset_enm_store, set_enm
from network_model.catalog.repository import get_default_mv_catalog
from network_model.solvers.ncrfg_ptpiree import (
    NcRfgPtpireeModuleInput,
    NcRfgPtpireeRunRequest,
    NcRfgPtpireeSolver,
)

from tests.cgmes.golden_enm import build_golden_enm

CERT_JSON = "/api/oze-analysis/compliance-certificate"
CERT_DOCX = "/api/oze-analysis/compliance-certificate.docx"
CERT_PDF = "/api/oze-analysis/compliance-certificate.pdf"
OSD_JSON = "/api/oze-analysis/osd-application"
OSD_DOCX = "/api/oze-analysis/osd-application.docx"
OSD_PDF = "/api/oze-analysis/osd-application.pdf"
STUDY_JSON = "/api/oze-analysis/connection-study"
STUDY_DOCX = "/api/oze-analysis/connection-study.docx"
STUDY_PDF = "/api/oze-analysis/connection-study.pdf"

#: Typ katalogowy z krzywą producenta — ten sam, którego używa dokument studium.
_CATALOG_ITEM = "conv-pv-card-sungrow-sg3150u-mv"
_OPERATOR_STUDIUM = "pge"

#: Referencja urządzenia w modelu = referencja modułu biegu NC RfG.
_DER_REF = "pv-1"

#: Wartości WPROST z rekordu wykazu PTPiREE (`mv_ptpiree_catalog.py`).
TABLICZKA_PELNA: dict = {
    "catalog_item_id": _CATALOG_ITEM,
    "ptpiree_status": "POWIAZANY",
    "ptpiree_certificate_ref": "ptpiree-pv-1",
    "ptpiree_document_number": "WOS/2024/PV-900",
    "ptpiree_document_acceptance_date": "2024-11-04",
    "ptpiree_wos_version": "WOS 2021",
    "ptpiree_wipwc_version": "1.3",
    "ptpiree_ppm_scope": "PPM typu B",
    "ptpiree_source_url": "https://ptpiree.pl/wykaz/pv-900",
    "ptpiree_certificate_condition": "Tylko z modułem sterowania SG-CTRL.",
}

_MODULE_FULL: dict = {
    "der_ref": _DER_REF,
    "der_name": "PV 2 MW",
    "der_kind": "PV",
    "operator_id": "enea",
    "p_max_kw": 2000,
    "p_min_kw": 100,
    "voltage_kv": 15,
    "certificate_status": "ptpiree_verified",
    "has_lvrt_curve": True,
    "has_hvrt_curve": True,
    "has_pf_droop": True,
    "has_qu_curve": True,
    "has_dynamic_model": True,
    "has_scada_communication": True,
    "has_disturbance_recorder": True,
    "active_power_control_enabled": True,
    "droop_percent": 5,
    "dead_band_hz": 0.2,
    "ramp_rate_pct_per_min": 10,
    "cos_phi_min": 0.95,
    "q_range_pct_pn_min": -0.33,
    "q_range_pct_pn_max": 0.33,
    "reactive_current_gain": 2,
    "p_recovery_time_s": 0.8,
    "harmonic_thdu_percent": 3,
}


@pytest.fixture(autouse=True)
def _reset() -> None:
    reset_canonical_runs()
    reset_enm_store()
    yield
    reset_canonical_runs()
    reset_enm_store()


def _model_z_urzadzeniem(tabliczka: dict | None, *, ref_id: str = _DER_REF) -> EnergyNetworkModel:
    return EnergyNetworkModel(
        header=ENMHeader(name="Dowod certyfikatu w dokumentach"),
        buses=[Bus(ref_id="bus_sn", name="Szyna SN", voltage_kv=15.0)],
        generators=[
            Generator(
                ref_id=ref_id,
                name="PV 2 MW",
                bus_ref="bus_sn",
                p_mw=2.0,
                gen_type="pv_inverter",
                materialized_params=tabliczka,
            )
        ],
    )


def _przypadek_z_urzadzeniem(tabliczka: dict | None, *, ref_id: str = _DER_REF) -> str:
    """Przypadek z jednym źródłem DER o zadanej tabliczce. Zwraca ``case_id``.

    Ten `case_id` jest tu SUROWYM kluczem magazynu (bez bazy) — testy poniżej
    przekazują go WPROST jako `klucz_twin` do `dowody_certyfikatu*` (wołanie
    bezpośrednie, nie przez API), więc nie potrzebują tłumaczenia CV-1. Testy
    REALNEJ ścieżki HTTP (`app_client.post(...case_id=...)`) używają zamiast
    tego `_przypadek_z_urzadzeniem_http` poniżej.
    """
    case_id = str(uuid4())
    set_enm(case_id, _model_z_urzadzeniem(tabliczka, ref_id=ref_id))
    return case_id


def _nowy_przypadek(client) -> str:
    """Utwórz REALNY projekt + przypadek przez API; zwróć ``case_id``.

    CV-1-W: przypadek bez wiersza w bazie dostaje teraz 404 z magazynu ENM
    (inwariant I-2) — końcówki ``/api/oze-analysis/...`` tłumaczą ``case_id``
    na klucz projektu (``api/oze_analysis_runs.py::_klucz_opcjonalny``).
    """
    project_resp = client.post("/api/projects", json={"name": "Dowod certyfikatu — test"})
    assert project_resp.status_code == 201, project_resp.text
    project_id = project_resp.json()["id"]
    case_resp = client.post(
        "/api/study-cases", json={"project_id": project_id, "name": "Przypadek testu"}
    )
    assert case_resp.status_code == 201, case_resp.text
    return str(case_resp.json()["id"])


def _klucz(client, case_id: str) -> str:
    """Klucz magazynu ENM dla ``case_id`` — TO SAMO tłumaczenie co warstwa API (CV-1)."""
    from application.twin_key import klucz_twin_dla_przypadku

    return klucz_twin_dla_przypadku(case_id, client.app.state.uow_factory)


def _przypadek_z_urzadzeniem_http(client, tabliczka: dict | None, *, ref_id: str = _DER_REF) -> str:
    """Jak ``_przypadek_z_urzadzeniem``, ale dla testów REALNEJ ścieżki HTTP:
    tworzy prawdziwy projekt+przypadek i zasiewa model pod kluczem
    PRZETŁUMACZONYM (nie surowym ``case_id``, którego żaden odczyt API już nie
    widzi — magazyn jest kluczowany kluczem PROJEKTU, CV-1-W)."""
    case_id = _nowy_przypadek(client)
    set_enm(_klucz(client, case_id), _model_z_urzadzeniem(tabliczka, ref_id=ref_id))
    return case_id


def _ncrfg():
    return NcRfgPtpireeSolver().run(
        NcRfgPtpireeRunRequest(modules=[NcRfgPtpireeModuleInput(**dict(_MODULE_FULL))])
    )


def _pf_run() -> CanonicalRun:
    set_enm("c-pf", build_golden_enm())
    return execute_run(create_run(case_id="c-pf", klucz_twin="c-pf", analysis_type="PF").id)


def _sc_run() -> CanonicalRun:
    set_enm("c-sc", build_golden_enm())
    return execute_run(
        create_run(case_id="c-sc", klucz_twin="c-sc", analysis_type="short_circuit_sn").id
    )


def _identyfikacja_wniosku() -> WniosekOsdIdentyfikacja:
    return WniosekOsdIdentyfikacja(nazwa_projektu="Farma PV Wschód")


def _identyfikacja_studium() -> DokumentStudiumIdentyfikacja:
    return DokumentStudiumIdentyfikacja(nazwa_projektu="Farma PV Wschód")


def _widok_studium(case_id: str | None, run: CanonicalRun | None = None) -> dict:
    return build_dokument_studium_view(
        run or _pf_run(),
        get_default_mv_catalog().get_converter_type(_CATALOG_ITEM),
        load_nc_rfg_profile(_OPERATOR_STUDIUM),
        catalog_item_id=_CATALOG_ITEM,
        operator_id=_OPERATOR_STUDIUM,
        warianty=["bus_sn_c"],
        identyfikacja=_identyfikacja_studium(),
        dowody=(None if case_id is None else dowody_certyfikatu_typu(case_id, _CATALOG_ITEM)),
    )


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _pdf_text(data: bytes) -> str:
    """Wyciągnij tekst z operatorów PDF (Tj/TJ) — PDF bez kompresji strony."""
    parts = [m.group(0) for m in re.finditer(rb"\((?:[^()\\]|\\.)*\)\s*Tj", data)]
    parts += [m.group(1) for m in re.finditer(rb"\[(.*?)\]\s*TJ", data, re.DOTALL)]
    return b" ".join(parts).decode("latin-1", "replace")


def _payload_certyfikatu() -> dict:
    return {
        "run_request": {"modules": [dict(_MODULE_FULL)]},
        "nazwa_projektu": "Farma PV Wschód",
    }


def _payload_wniosku(pf: CanonicalRun, sc: CanonicalRun) -> dict:
    return {
        "nazwa_projektu": "Farma PV Wschód",
        "bus_ref": "bus_nn",
        "pf_run_id": str(pf.id),
        "sc_run_id": str(sc.id),
        "run_request": {"modules": [dict(_MODULE_FULL)]},
    }


def _payload_studium(run: CanonicalRun) -> dict:
    return {
        "nazwa_projektu": "Farma PV Wschód",
        "run_id": str(run.id),
        "catalog_item_id": _CATALOG_ITEM,
        "operator_id": _OPERATOR_STUDIUM,
        "warianty": ["bus_sn_c"],
    }


# --------------------------------------------------------------------------- #
# Wspólne źródło dowodu — tabliczka urządzenia w modelu
# --------------------------------------------------------------------------- #
def test_dowod_z_referencji_cytuje_komplet_tabliczki() -> None:
    """Certyfikat i wniosek identyfikują urządzenie referencją modułu biegu."""
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    dowod = dowody_certyfikatu(case_id, [_DER_REF])[0]

    assert dowod.document_number == "WOS/2024/PV-900"
    assert dowod.wipwc_version == "1.3"
    assert dowod.acceptance_date == "2024-11-04"
    assert dowod.certificate_condition == "Tylko z modułem sterowania SG-CTRL."
    assert dowod.source_url == "https://ptpiree.pl/wykaz/pv-900"


def test_dowod_z_typu_katalogowego_dopasowuje_po_tabliczce() -> None:
    """Dokument studium identyfikuje urządzenie TYPEM katalogowym, nie modułem."""
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA, ref_id="pv-w-modelu")
    dowody = dowody_certyfikatu_typu(case_id, _CATALOG_ITEM)

    assert [d.der_ref for d in dowody] == ["pv-w-modelu"]
    assert dowody[0].document_number == "WOS/2024/PV-900"


def test_dowod_typu_nie_przyklei_sie_do_innego_typu() -> None:
    """Dopasowanie po typie — nie „pierwsze lepsze urządzenie w modelu"."""
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    assert dowody_certyfikatu_typu(case_id, "conv-innego-projektu") == []


def test_dowod_referencji_spoza_modelu_jest_pusty_a_nie_cudzy() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    dowod = dowody_certyfikatu(case_id, ["gen-innego-projektu"])[0]

    assert dowod.der_ref == "gen-innego-projektu"
    assert dowod.document_number is None, "dowod przykleil sie do niewlasciwego urzadzenia"


# --------------------------------------------------------------------------- #
# Certyfikat zgodności — widok, stan zerowy, kontrakt bez przypadku
# --------------------------------------------------------------------------- #
def test_certyfikat_niesie_dowod_per_modul() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    view = build_certyfikat_view(
        _ncrfg(),
        nazwa_projektu="Farma PV Wschód",
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )

    dowod = view["moduly"][0]["dowod_certyfikatu"]
    assert dowod["der_ref"] == _DER_REF
    assert dowod["numer_dokumentu"] == "WOS/2024/PV-900"
    assert dowod["wersja_wipwc"] == "1.3"
    assert dowod["data_akceptacji"] == "2024-11-04"
    assert dowod["warunek_waznosci"] == "Tylko z modułem sterowania SG-CTRL."
    assert dowod["adres_zrodla"] == "https://ptpiree.pl/wykaz/pv-900"
    assert dowod["stan_pl"] is None


def test_certyfikat_bez_tabliczki_ma_jawny_stan_zerowy() -> None:
    case_id = _przypadek_z_urzadzeniem(None)
    view = build_certyfikat_view(
        _ncrfg(),
        nazwa_projektu="Farma PV Wschód",
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )

    dowod = view["moduly"][0]["dowod_certyfikatu"]
    assert dowod["stan_pl"] == BRAK_TABLICZKI_PL
    assert dowod["numer_dokumentu"] is None


def test_certyfikat_niepelna_tabliczka_niesie_tylko_to_co_jest() -> None:
    case_id = _przypadek_z_urzadzeniem(
        {"ptpiree_document_number": "WOS/2024/PV-900", "ptpiree_wos_version": "WOS 2021"}
    )
    view = build_certyfikat_view(
        _ncrfg(),
        nazwa_projektu="Farma PV Wschód",
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )

    dowod = view["moduly"][0]["dowod_certyfikatu"]
    assert dowod["numer_dokumentu"] == "WOS/2024/PV-900"
    assert dowod["wersja_wos"] == "WOS 2021"
    assert dowod["wersja_wipwc"] is None
    assert dowod["warunek_waznosci"] is None
    assert dowod["stan_pl"] is None, "czesciowa tabliczka nie jest stanem zerowym"


def test_certyfikat_bez_przypadku_ma_kontrakt_sprzed_dowodu() -> None:
    """Bez wskazanego przypadku widok jest IDENTYCZNY jak przed dodaniem sekcji."""
    run_result = _ncrfg()
    bez_dowodu = build_certyfikat_view(run_result, nazwa_projektu="Farma PV Wschód")

    assert "dowod_certyfikatu" not in bez_dowodu["moduly"][0]
    assert bez_dowodu["odcisk_wejscia_sha256"] == run_result.input_hash
    assert bez_dowodu["odcisk_wyniku_sha256"] == run_result.deterministic_hash


# --------------------------------------------------------------------------- #
# Wniosek OSD — widok, stan zerowy, kontrakt bez przypadku
# --------------------------------------------------------------------------- #
def test_wniosek_niesie_dowod_w_sekcji_zgodnosci() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    view = build_wniosek_osd_view(
        _pf_run(),
        _sc_run(),
        _ncrfg(),
        bus_ref="bus_nn",
        identyfikacja=_identyfikacja_wniosku(),
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )

    dowody = view["zgodnosc_nc_rfg"]["dowody_certyfikatu"]
    assert [d["der_ref"] for d in dowody] == [_DER_REF]
    assert dowody[0]["numer_dokumentu"] == "WOS/2024/PV-900"
    assert dowody[0]["warunek_waznosci"] == "Tylko z modułem sterowania SG-CTRL."


def test_wniosek_bez_tabliczki_ma_jawny_stan_zerowy() -> None:
    case_id = _przypadek_z_urzadzeniem(None)
    view = build_wniosek_osd_view(
        _pf_run(),
        _sc_run(),
        _ncrfg(),
        bus_ref="bus_nn",
        identyfikacja=_identyfikacja_wniosku(),
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )

    assert view["zgodnosc_nc_rfg"]["dowody_certyfikatu"][0]["stan_pl"] == BRAK_TABLICZKI_PL


def test_wniosek_bez_przypadku_ma_kontrakt_sprzed_dowodu() -> None:
    """Równość widoku i ODCISKÓW — dowód nie może ruszyć hasha bez przypadku."""
    pf, sc, ncrfg = _pf_run(), _sc_run(), _ncrfg()
    wspolne = {"bus_ref": "bus_nn", "identyfikacja": _identyfikacja_wniosku()}
    bez_dowodu = build_wniosek_osd_view(pf, sc, ncrfg, **wspolne)
    bez_dowodu_jawnie = build_wniosek_osd_view(pf, sc, ncrfg, **wspolne, dowody=None)

    assert "dowody_certyfikatu" not in bez_dowodu["zgodnosc_nc_rfg"]
    assert json.dumps(bez_dowodu, sort_keys=True) == json.dumps(bez_dowodu_jawnie, sort_keys=True)


def test_wniosek_dowod_zmienia_wylacznie_odcisk_swojej_sekcji() -> None:
    """Dowód wchodzi do odcisku sekcji zgodności — i tylko tam."""
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    pf, sc, ncrfg = _pf_run(), _sc_run(), _ncrfg()
    wspolne = {"bus_ref": "bus_nn", "identyfikacja": _identyfikacja_wniosku()}
    bez = build_wniosek_osd_view(pf, sc, ncrfg, **wspolne)
    z_dowodem = build_wniosek_osd_view(
        pf, sc, ncrfg, **wspolne, dowody=dowody_certyfikatu(case_id, [_DER_REF])
    )

    odciski_bez = bez["odciski_sekcji_sha256"]
    odciski_z = z_dowodem["odciski_sekcji_sha256"]
    assert odciski_z["zgodnosc_nc_rfg"] != odciski_bez["zgodnosc_nc_rfg"]
    assert odciski_z["bilans_mocy"] == odciski_bez["bilans_mocy"]
    assert odciski_z["zwarcia_punkt_przylaczenia"] == odciski_bez["zwarcia_punkt_przylaczenia"]
    assert z_dowodem["input_hash"] == bez["input_hash"]


# --------------------------------------------------------------------------- #
# Dokument studium — widok, stan zerowy, kontrakt bez przypadku
# --------------------------------------------------------------------------- #
def test_studium_niesie_dowod_urzadzen_typu() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA, ref_id="pv-w-modelu")
    blok = _widok_studium(case_id)["zalozenia"]["dowod_certyfikatu"]

    assert blok["catalog_item_id"] == _CATALOG_ITEM
    assert blok["stan_pl"] is None
    assert [u["der_ref"] for u in blok["urzadzenia"]] == ["pv-w-modelu"]
    assert blok["urzadzenia"][0]["numer_dokumentu"] == "WOS/2024/PV-900"


def test_studium_bez_urzadzenia_tego_typu_ma_jawny_stan_zerowy() -> None:
    case_id = _przypadek_z_urzadzeniem({"catalog_item_id": "conv-innego-projektu"})
    blok = _widok_studium(case_id)["zalozenia"]["dowod_certyfikatu"]

    assert blok["urzadzenia"] == []
    assert blok["stan_pl"] == BRAK_URZADZEN_TYPU_PL


def test_studium_bez_przypadku_ma_kontrakt_sprzed_dowodu() -> None:
    run = _pf_run()
    bez = _widok_studium(None, run)

    assert "dowod_certyfikatu" not in bez["zalozenia"]
    # Odcisk założeń i hash wejścia są punktem odniesienia dokumentu — bez
    # wskazanego przypadku nie wolno im drgnąć.
    assert set(bez["odciski_sekcji_sha256"]) == {"bus_sn_c", "zalozenia", "podsumowanie"}


def test_studium_dowod_zmienia_wylacznie_odcisk_zalozen() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA, ref_id="pv-w-modelu")
    run = _pf_run()
    bez = _widok_studium(None, run)
    z_dowodem = _widok_studium(case_id, run)

    assert (
        z_dowodem["odciski_sekcji_sha256"]["zalozenia"] != bez["odciski_sekcji_sha256"]["zalozenia"]
    )
    assert (
        z_dowodem["odciski_sekcji_sha256"]["bus_sn_c"] == bez["odciski_sekcji_sha256"]["bus_sn_c"]
    )
    assert z_dowodem["input_hash"] == bez["input_hash"]


# --------------------------------------------------------------------------- #
# Render DOCX — dowód musi być CZYTELNY w pliku oddawanym operatorowi
# --------------------------------------------------------------------------- #
def test_docx_certyfikatu_zawiera_numer_dokumentu() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    view = build_certyfikat_view(
        _ncrfg(),
        nazwa_projektu="Farma PV Wschód",
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )
    text = _docx_text(render_certyfikat_docx(view))

    assert TYTUL_DOWODU in text
    assert "WOS/2024/PV-900" in text
    assert "Tylko z modułem sterowania SG-CTRL." in text


def test_docx_certyfikatu_bez_tabliczki_mowi_wprost_ze_dowodu_nie_ma() -> None:
    case_id = _przypadek_z_urzadzeniem(None)
    view = build_certyfikat_view(
        _ncrfg(),
        nazwa_projektu="Farma PV Wschód",
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )
    assert BRAK_TABLICZKI_PL in _docx_text(render_certyfikat_docx(view))


def test_docx_wniosku_zawiera_numer_dokumentu() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA)
    view = build_wniosek_osd_view(
        _pf_run(),
        _sc_run(),
        _ncrfg(),
        bus_ref="bus_nn",
        identyfikacja=_identyfikacja_wniosku(),
        dowody=dowody_certyfikatu(case_id, [_DER_REF]),
    )
    text = _docx_text(render_wniosek_osd_docx(view))

    assert TYTUL_DOWODU in text
    assert "WOS/2024/PV-900" in text


def test_docx_studium_zawiera_numer_dokumentu() -> None:
    case_id = _przypadek_z_urzadzeniem(TABLICZKA_PELNA, ref_id="pv-w-modelu")
    text = _docx_text(render_dokument_studium_docx(_widok_studium(case_id)))

    assert TYTUL_DOWODU in text
    assert "WOS/2024/PV-900" in text


def test_docx_dokumentow_bez_przypadku_nie_pokazuje_sekcji_dowodu() -> None:
    """Brak przypadku = brak sekcji, a nie pusta sekcja z myślnikami."""
    assert TYTUL_DOWODU not in _docx_text(
        render_certyfikat_docx(build_certyfikat_view(_ncrfg(), nazwa_projektu="Farma PV Wschód"))
    )
    assert TYTUL_DOWODU not in _docx_text(render_dokument_studium_docx(_widok_studium(None)))


# --------------------------------------------------------------------------- #
# Końcówki API — case_id dopina dowód, jego brak zachowuje kontrakt
# --------------------------------------------------------------------------- #
def test_endpoint_certyfikatu_z_przypadkiem_niesie_dowod(app_client) -> None:
    case_id = _przypadek_z_urzadzeniem_http(app_client, TABLICZKA_PELNA)
    resp = app_client.post(f"{CERT_JSON}?case_id={case_id}", json=_payload_certyfikatu())

    assert resp.status_code == 200, resp.text
    dowod = resp.json()["moduly"][0]["dowod_certyfikatu"]
    assert dowod["numer_dokumentu"] == "WOS/2024/PV-900"
    assert dowod["warunek_waznosci"] == "Tylko z modułem sterowania SG-CTRL."


def test_endpoint_certyfikatu_bez_przypadku_zachowuje_kontrakt(app_client) -> None:
    resp = app_client.post(CERT_JSON, json=_payload_certyfikatu())

    assert resp.status_code == 200, resp.text
    assert "dowod_certyfikatu" not in resp.json()["moduly"][0]


def test_endpoint_wniosku_z_przypadkiem_niesie_dowod(app_client) -> None:
    case_id = _przypadek_z_urzadzeniem_http(app_client, TABLICZKA_PELNA)
    resp = app_client.post(
        f"{OSD_JSON}?case_id={case_id}", json=_payload_wniosku(_pf_run(), _sc_run())
    )

    assert resp.status_code == 200, resp.text
    dowody = resp.json()["zgodnosc_nc_rfg"]["dowody_certyfikatu"]
    assert dowody[0]["numer_dokumentu"] == "WOS/2024/PV-900"


def test_endpoint_wniosku_bez_przypadku_zachowuje_kontrakt(app_client) -> None:
    resp = app_client.post(OSD_JSON, json=_payload_wniosku(_pf_run(), _sc_run()))

    assert resp.status_code == 200, resp.text
    assert "dowody_certyfikatu" not in resp.json()["zgodnosc_nc_rfg"]


def test_endpoint_studium_z_przypadkiem_niesie_dowod(app_client) -> None:
    case_id = _przypadek_z_urzadzeniem_http(app_client, TABLICZKA_PELNA, ref_id="pv-w-modelu")
    run = _pf_run()
    resp = app_client.post(f"{STUDY_JSON}?case_id={case_id}", json=_payload_studium(run))

    assert resp.status_code == 200, resp.text
    blok = resp.json()["zalozenia"]["dowod_certyfikatu"]
    assert blok["urzadzenia"][0]["numer_dokumentu"] == "WOS/2024/PV-900"


def test_endpoint_studium_bez_przypadku_zachowuje_kontrakt(app_client) -> None:
    resp = app_client.post(STUDY_JSON, json=_payload_studium(_pf_run()))

    assert resp.status_code == 200, resp.text
    assert "dowod_certyfikatu" not in resp.json()["zalozenia"]


@pytest.mark.parametrize(
    ("sciezka", "buduj_payload"),
    [
        (CERT_DOCX, "certyfikat"),
        (CERT_PDF, "certyfikat"),
        (OSD_DOCX, "wniosek"),
        (OSD_PDF, "wniosek"),
        (STUDY_DOCX, "studium"),
        (STUDY_PDF, "studium"),
    ],
)
def test_endpointy_plikow_przyjmuja_przypadek_i_niosa_dowod(
    app_client, sciezka: str, buduj_payload: str
) -> None:
    """Każda z 6 końcówek plikowych przyjmuje case_id i cytuje numer dokumentu."""
    case_id = _przypadek_z_urzadzeniem_http(app_client, TABLICZKA_PELNA, ref_id="pv-1")
    if buduj_payload == "certyfikat":
        payload = _payload_certyfikatu()
    elif buduj_payload == "wniosek":
        payload = _payload_wniosku(_pf_run(), _sc_run())
    else:
        # Studium dopasowuje po typie katalogowym — urządzenie modelu niesie
        # ten sam `catalog_item_id`, więc dowód powstaje tą samą tabliczką.
        payload = _payload_studium(_pf_run())

    resp = app_client.post(f"{sciezka}?case_id={case_id}", json=payload)
    assert resp.status_code == 200, resp.text

    text = _docx_text(resp.content) if sciezka.endswith(".docx") else _pdf_text(resp.content)
    assert "WOS/2024/PV-900" in text


@pytest.mark.parametrize("sciezka", [CERT_DOCX, CERT_PDF, STUDY_DOCX, STUDY_PDF])
def test_endpointy_plikow_bez_przypadku_pozostaja_deterministyczne(
    app_client, sciezka: str
) -> None:
    """Brak case_id → bajtowo ten sam plik przy dwóch wywołaniach (kontrakt)."""
    payload = (
        _payload_certyfikatu() if sciezka.startswith(CERT_JSON) else _payload_studium(_pf_run())
    )
    pierwszy = app_client.post(sciezka, json=payload)
    drugi = app_client.post(sciezka, json=payload)

    assert pierwszy.status_code == 200, pierwszy.text
    assert pierwszy.content == drugi.content
