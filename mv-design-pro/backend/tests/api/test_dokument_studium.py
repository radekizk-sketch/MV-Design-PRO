"""Testy dokumentu studium przyłączeniowego (E13, D17) — serwis + końcówki API.

Dokument studium to serwerowa KOMPOZYCJA tej samej sekwencji, którą liczy kreator
studium: per wariant zdolność przyłączeniowa → obszar pracy P–Q → pokrycie wymagań
P–Q (serwisy wołane bezpośrednio). Testy pokrywają: dokument z 2 wariantów (sekcje
per wariant + podsumowanie), zachowaną kolejność wariantów, klasę NC RfG z jedynej
klasyfikacji backendowej, odporność na błąd pojedynczego wariantu (reszta liczona),
bramkę braków twardych (422 z listą), determinizm bajtowy DOCX i PDF, stabilność
odcisków/hash, etykiety PL w obu formatach, content-type oraz 404/422.
"""

from __future__ import annotations

import re
from io import BytesIO
from uuid import uuid4

import pytest
from application.analyses.dokument_studium import (
    DokumentStudiumBrakiError,
    DokumentStudiumIdentyfikacja,
    build_dokument_studium_view,
    render_dokument_studium_docx,
    render_dokument_studium_pdf,
)
from catalog.profiles.nc_rfg.loader import load_nc_rfg_profile
from docx import Document
from enm.canonical_analysis import (
    CanonicalRun,
    create_run,
    execute_run,
    reset_canonical_runs,
)
from enm.store import reset_enm_store, set_enm
from network_model.catalog.repository import get_default_mv_catalog

from tests.cgmes.golden_enm import build_golden_enm

STUDY_JSON = "/api/oze-analysis/connection-study"
STUDY_DOCX = "/api/oze-analysis/connection-study.docx"
STUDY_PDF = "/api/oze-analysis/connection-study.pdf"

_CATALOG_ITEM = "conv-pv-card-sungrow-sg3150u-mv"  # typ z krzywą producenta
_OPERATOR = "pge"


@pytest.fixture(autouse=True)
def _reset() -> None:
    reset_canonical_runs()
    reset_enm_store()
    yield
    reset_canonical_runs()
    reset_enm_store()


def _pf_run(case_id: str = "c-pf") -> CanonicalRun:
    set_enm(case_id, build_golden_enm())
    return execute_run(create_run(case_id=case_id, analysis_type="PF").id)


def _sc_run() -> CanonicalRun:
    set_enm("c-sc", build_golden_enm())
    return execute_run(create_run(case_id="c-sc", analysis_type="short_circuit_sn").id)


def _converter():
    return get_default_mv_catalog().get_converter_type(_CATALOG_ITEM)


def _profile():
    return load_nc_rfg_profile(_OPERATOR)


def _identyfikacja() -> DokumentStudiumIdentyfikacja:
    return DokumentStudiumIdentyfikacja(
        nazwa_projektu="Farma PV Wschód",
        nazwa_przypadku="Wariant bazowy",
        wnioskodawca="OZE Sp. z o.o.",
        adres_przylaczenia="Stacja B",
    )


def _view(run: CanonicalRun | None = None, warianty: list[str] | None = None):
    return build_dokument_studium_view(
        run or _pf_run(),
        _converter(),
        _profile(),
        catalog_item_id=_CATALOG_ITEM,
        operator_id=_OPERATOR,
        warianty=warianty or ["bus_sn_c", "bus_nn"],
        identyfikacja=_identyfikacja(),
    )


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _pdf_text(data: bytes) -> str:
    """Wyciągnij tekst z operatorów PDF (Tj/TJ) — PDF bez kompresji strony."""
    parts = [m.group(0) for m in re.finditer(rb"\((?:[^()\\]|\\.)*\)\s*Tj", data)]
    parts += [m.group(1) for m in re.finditer(rb"\[(.*?)\]\s*TJ", data, re.DOTALL)]
    return b" ".join(parts).decode("latin-1", "replace")


def _payload(run: CanonicalRun, **extra) -> dict:
    body = {
        "nazwa_projektu": "Farma PV Wschód",
        "nazwa_przypadku": "Wariant bazowy",
        "wnioskodawca": "OZE Sp. z o.o.",
        "adres_przylaczenia": "Stacja B",
        "run_id": str(run.id),
        "catalog_item_id": _CATALOG_ITEM,
        "operator_id": _OPERATOR,
        "warianty": ["bus_sn_c", "bus_nn"],
    }
    body.update(extra)
    return body


# --------------------------------------------------------------------------- #
# Serwis kompozycji — sekwencja per wariant
# --------------------------------------------------------------------------- #
def test_dwa_warianty_sekcje_i_podsumowanie() -> None:
    view = _view()
    assert view["kontrakt"] == "DokumentStudiumPrzylaczeniowegoV1"
    assert len(view["warianty"]) == 2
    for wariant in view["warianty"]:
        assert {"zdolnosc", "obszar_pq", "pokrycie_pq", "klasa_nc_rfg"} <= set(wariant)
    assert len(view["podsumowanie"]) == 2
    assert view["zalozenia"]["liczba_wariantow"] == 2


def test_kolejnosc_wariantow_zachowana() -> None:
    view = _view(warianty=["bus_nn", "bus_sn_c"])
    assert [w["bus_ref"] for w in view["warianty"]] == ["bus_nn", "bus_sn_c"]
    assert [w["bus_ref"] for w in view["podsumowanie"]] == ["bus_nn", "bus_sn_c"]


def test_klasa_nc_rfg_z_klasyfikacji_backendowej() -> None:
    # Klasa liczona istniejącą klasyfikacją NcRfgProfile.classify_module.
    view = _view(warianty=["bus_sn_c"])
    wariant = view["warianty"][0]
    assert wariant["zdolnosc"]["status"] == "ok"
    assert wariant["klasa_nc_rfg"]["klasa"] in {"A", "B", "C", "D"}
    # Wartość zgodna z bezpośrednim wywołaniem jedynego źródła prawdy.
    max_moc = wariant["zdolnosc"]["max_moc_mw"]
    oczekiwana = _profile().classify_module(max_moc * 1000.0, wariant["napiecie_kv"])
    assert wariant["klasa_nc_rfg"]["klasa"] == oczekiwana.id


def test_pokrycie_pq_werdykt_w_wariancie() -> None:
    wariant = _view(warianty=["bus_sn_c"])["warianty"][0]
    assert wariant["pokrycie_pq"]["status"] == "ok"
    assert wariant["pokrycie_pq"]["werdykt_pl"] in {"Pokryte", "Niepokryte"}


# --------------------------------------------------------------------------- #
# Odporność na błąd pojedynczego wariantu
# --------------------------------------------------------------------------- #
def test_blad_jednego_wariantu_nie_przerywa_dokumentu() -> None:
    view = _view(warianty=["bus_sn_c", "bus_nieistniejacy"])
    dobry, zly = view["warianty"][0], view["warianty"][1]
    assert dobry["bus_ref"] == "bus_sn_c"
    assert dobry["zdolnosc"]["status"] == "ok"
    assert zly["bus_ref"] == "bus_nieistniejacy"
    assert zly["zdolnosc"]["status"] == "blad"
    assert zly["zdolnosc"]["komunikat_bledu"]
    assert zly["obszar_pq"]["status"] == "blad"
    # Klasa nieokreślona przy braku mocy przyłączalnej (bez zgadywania).
    assert zly["klasa_nc_rfg"]["klasa"] is None


# --------------------------------------------------------------------------- #
# Bramka braków twardych (422 z listą)
# --------------------------------------------------------------------------- #
def test_braki_twarde_nieznany_typ_katalogowy() -> None:
    with pytest.raises(DokumentStudiumBrakiError) as exc:
        build_dokument_studium_view(
            _pf_run(),
            None,  # nieznany typ katalogowy
            _profile(),
            catalog_item_id="nie-ma-takiego",
            operator_id=_OPERATOR,
            warianty=["bus_sn_c"],
            identyfikacja=_identyfikacja(),
        )
    assert any("nie istnieje" in b for b in exc.value.braki)


def test_braki_twarde_nieznany_operator() -> None:
    with pytest.raises(DokumentStudiumBrakiError) as exc:
        build_dokument_studium_view(
            _pf_run(),
            _converter(),
            None,  # nieznany profil operatora
            catalog_item_id=_CATALOG_ITEM,
            operator_id="nieistniejacy",
            warianty=["bus_sn_c"],
            identyfikacja=_identyfikacja(),
        )
    assert any("profilu NC RfG" in b for b in exc.value.braki)


def test_braki_twarde_zly_rodzaj_przebiegu() -> None:
    with pytest.raises(DokumentStudiumBrakiError) as exc:
        build_dokument_studium_view(
            _sc_run(),  # przebieg zwarciowy zamiast rozpływu
            _converter(),
            _profile(),
            catalog_item_id=_CATALOG_ITEM,
            operator_id=_OPERATOR,
            warianty=["bus_sn_c"],
            identyfikacja=_identyfikacja(),
        )
    assert any("nie jest rozpływem" in b for b in exc.value.braki)


def test_kolejnosc_brakow_deterministyczna() -> None:
    # Przebieg → typ katalogowy → operator → warianty.
    from application.analyses.dokument_studium import zbierz_braki_dokumentu

    lista = zbierz_braki_dokumentu(
        _sc_run(),
        None,
        None,
        catalog_item_id="x",
        operator_id="y",
        warianty=[],
    )
    assert "nie jest rozpływem" in lista[0]
    assert "katalog" in lista[1].lower() or "typ" in lista[1].lower()
    assert "operator" in lista[2].lower()
    assert "wariant" in lista[3].lower()


# --------------------------------------------------------------------------- #
# Determinizm i odciski
# --------------------------------------------------------------------------- #
def test_docx_determinizm_bajtowy() -> None:
    view = _view(warianty=["bus_sn_c"])
    assert render_dokument_studium_docx(view) == render_dokument_studium_docx(view)
    assert len(render_dokument_studium_docx(view)) > 0


def test_pdf_determinizm_bajtowy() -> None:
    view = _view(warianty=["bus_sn_c"])
    assert render_dokument_studium_pdf(view) == render_dokument_studium_pdf(view)
    assert render_dokument_studium_pdf(view)[:4] == b"%PDF"


def test_input_hash_stabilny_dla_tego_samego_snapshotu() -> None:
    # Dwa niezależne przebiegi z tego samego modelu → ten sam input_hash
    # (hash oparty o snapshot, nie o identyfikator przebiegu).
    first = _view(_pf_run("c-a"), warianty=["bus_sn_c"])["input_hash"]
    second = _view(_pf_run("c-b"), warianty=["bus_sn_c"])["input_hash"]
    assert first == second


def test_odciski_sekcji_obecne_dla_wariantow_i_zbiorczych() -> None:
    view = _view()
    odciski = view["odciski_sekcji_sha256"]
    assert "bus_sn_c" in odciski and "bus_nn" in odciski
    assert "zalozenia" in odciski and "podsumowanie" in odciski
    assert view["input_hash"]


# --------------------------------------------------------------------------- #
# Etykiety PL w obu formatach
# --------------------------------------------------------------------------- #
def test_docx_etykiety_pl() -> None:
    text = _docx_text(render_dokument_studium_docx(_view(warianty=["bus_sn_c"])))
    assert "Dokument studium przyłączeniowego OZE" in text
    assert "Zdolność przyłączeniowa" in text
    assert "Obszar pracy P–Q" in text
    assert "Pokrycie wymagań P–Q" in text
    assert "Klasa NC RfG" in text
    assert "Podsumowanie porównawcze wariantów" in text
    assert "Farma PV Wschód" in text


def test_pdf_etykiety_pl() -> None:
    text = _pdf_text(render_dokument_studium_pdf(_view(warianty=["bus_sn_c"])))
    assert "Zdolno" in text  # „Zdolność" (kodowanie latin-1 w ekstrakcji)
    assert "Pokrycie" in text
    assert "Podsumowanie" in text


def test_docx_bez_kodow_projektowych() -> None:
    text = _docx_text(render_dokument_studium_docx(_view(warianty=["bus_sn_c"])))
    for pattern in (r"\bP\d{2}\b", r"\bE\d{2}\b", r"\bD\d{2}\b"):
        assert re.search(pattern, text) is None, pattern


# --------------------------------------------------------------------------- #
# Końcówki API
# --------------------------------------------------------------------------- #
def test_endpoint_json_200_komplet(app_client) -> None:
    run = _pf_run()
    resp = app_client.post(STUDY_JSON, json=_payload(run))
    assert resp.status_code == 200
    data = resp.json()
    assert data["identyfikacja"]["projekt"] == "Farma PV Wschód"
    assert len(data["warianty"]) == 2


def test_endpoint_json_braki_422_nieznany_typ(app_client) -> None:
    run = _pf_run()
    resp = app_client.post(STUDY_JSON, json=_payload(run, catalog_item_id="nie-ma-takiego"))
    assert resp.status_code == 422
    detail = resp.json()["detail"]
    assert isinstance(detail["braki"], list)
    assert any("nie istnieje" in b for b in detail["braki"])


def test_endpoint_nieznany_przebieg_404(app_client) -> None:
    _pf_run()
    resp = app_client.post(STUDY_JSON, json=_payload(_pf_run(), run_id=str(uuid4())))
    assert resp.status_code == 404
    assert "nie istnieje" in resp.json()["detail"]


def test_endpoint_docx_content_type(app_client) -> None:
    run = _pf_run()
    resp = app_client.post(STUDY_DOCX, json=_payload(run))
    assert resp.status_code == 200
    assert resp.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content[:2] == b"PK"


def test_endpoint_pdf_content_type(app_client) -> None:
    run = _pf_run()
    resp = app_client.post(STUDY_PDF, json=_payload(run))
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content[:4] == b"%PDF"


def test_endpoint_docx_determinizm(app_client) -> None:
    run = _pf_run()
    payload = _payload(run, warianty=["bus_sn_c"])
    first = app_client.post(STUDY_DOCX, json=payload)
    second = app_client.post(STUDY_DOCX, json=payload)
    assert first.status_code == 200
    assert first.content == second.content


def test_endpoint_pdf_determinizm(app_client) -> None:
    run = _pf_run()
    payload = _payload(run, warianty=["bus_sn_c"])
    first = app_client.post(STUDY_PDF, json=payload)
    second = app_client.post(STUDY_PDF, json=payload)
    assert first.status_code == 200
    assert first.content == second.content
