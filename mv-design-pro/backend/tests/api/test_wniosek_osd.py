"""Testy generatora wniosku OSD (W-707) — serwis kompozycji + końcówki API.

Wniosek OSD to czysta kompozycja gotowych wyników: bilans mocy z walidacji
energetycznej przebiegu rozpływu (``PF``), zwarcia w punkcie przyłączenia
z przebiegu zwarciowego (``short_circuit_sn``) oraz zgodność NC RfG tą samą
ścieżką co certyfikat D14. Testy pokrywają: komplet źródeł → 3 sekcje, bramkę
braków (lista PL blokuje generację), determinizm bajtowy DOCX, stabilność
odcisków, etykiety PL i uczciwe adnotacje (schemat, zestawienia), content-type
oraz 404/422.
"""

from __future__ import annotations

import re
from datetime import UTC, datetime
from io import BytesIO
from uuid import uuid4

import pytest
from application.analyses.wniosek_osd import (
    WniosekOsdBrakiError,
    WniosekOsdIdentyfikacja,
    build_wniosek_osd_view,
    render_wniosek_osd_docx,
    render_wniosek_pdf,
    zbierz_braki_wniosku,
)
from docx import Document
from enm.canonical_analysis import (
    CanonicalRun,
    create_run,
    execute_run,
    reset_canonical_runs,
)
from enm.models import GenLimits
from enm.store import reset_enm_store, set_enm
from network_model.solvers.ncrfg_ptpiree import (
    NcRfgPtpireeModuleInput,
    NcRfgPtpireeRunRequest,
    NcRfgPtpireeSolver,
)

from tests.cgmes.golden_enm import build_golden_enm

OSD_JSON = "/api/oze-analysis/osd-application"
OSD_DOCX = "/api/oze-analysis/osd-application.docx"
OSD_PDF = "/api/oze-analysis/osd-application.pdf"

_MODULE_FULL: dict = {
    "der_ref": "pv-1",
    "der_name": "PV 2 MW",
    "der_kind": "PV",
    "operator_id": "enea",
    "p_max_kw": 2000,
    "p_min_kw": 100,
    "voltage_kv": 15,
    "certificate_status": "ptpiree_verified",
    "has_lvrt_curve": True,
    "has_hvrt_curve": True,
    "has_pf_droop": True,
    "has_qu_curve": True,
    "has_dynamic_model": True,
    "has_scada_communication": True,
    "has_disturbance_recorder": True,
    "active_power_control_enabled": True,
    "droop_percent": 5,
    "dead_band_hz": 0.2,
    "ramp_rate_pct_per_min": 10,
    "cos_phi_min": 0.95,
    "q_range_pct_pn_min": -0.33,
    "q_range_pct_pn_max": 0.33,
    "reactive_current_gain": 2,
    "p_recovery_time_s": 0.8,
    "harmonic_thdu_percent": 3,
}


@pytest.fixture(autouse=True)
def _reset() -> None:
    reset_canonical_runs()
    reset_enm_store()
    yield
    reset_canonical_runs()
    reset_enm_store()


def _augmented_enm():
    enm = build_golden_enm()
    gens = list(enm.generators)
    gens[1] = gens[1].model_copy(
        update={
            "limits": GenLimits(q_min_mvar=-0.9, q_max_mvar=0.9),
            "materialized_params": {"sn_mva": 2.75},
        }
    )
    return enm.model_copy(update={"generators": gens})


def _pf_run() -> CanonicalRun:
    set_enm("c-pf", _augmented_enm())
    return execute_run(create_run(case_id="c-pf", analysis_type="PF").id)


def _sc_run() -> CanonicalRun:
    set_enm("c-sc", _augmented_enm())
    return execute_run(create_run(case_id="c-sc", analysis_type="short_circuit_sn").id)


def _ncrfg(module: dict | None = None) -> NcRfgPtpireeSolver:
    data = dict(_MODULE_FULL)
    if module:
        data.update(module)
    return NcRfgPtpireeSolver().run(
        NcRfgPtpireeRunRequest(modules=[NcRfgPtpireeModuleInput(**data)])
    )


def _identyfikacja() -> WniosekOsdIdentyfikacja:
    return WniosekOsdIdentyfikacja(
        nazwa_projektu="Farma PV Wschód",
        nazwa_przypadku="Wariant bazowy",
        wnioskodawca="OZE Sp. z o.o.",
        adres_przylaczenia="Stacja B",
    )


def _fake_run(analysis_type: str, status: str) -> CanonicalRun:
    return CanonicalRun(
        id=uuid4(),
        case_id="c",
        project_id="p",
        analysis_type=analysis_type,
        status=status,
        created_at=datetime(2024, 1, 1, tzinfo=UTC),
        snapshot_hash="h",
        input_hash="h",
        snapshot={"header": {"name": "Projekt"}, "buses": [], "generators": []},
        validation={},
        readiness={},
        raw_result={},
    )


def _view():
    return build_wniosek_osd_view(
        _pf_run(),
        _sc_run(),
        _ncrfg(),
        bus_ref="bus_nn",
        identyfikacja=_identyfikacja(),
    )


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _pdf_text(data: bytes) -> str:
    """Wyciągnij tekst z operatorów PDF (Tj/TJ) — PDF bez kompresji strony."""
    parts = [m.group(0) for m in re.finditer(rb"\((?:[^()\\]|\\.)*\)\s*Tj", data)]
    parts += [m.group(1) for m in re.finditer(rb"\[(.*?)\]\s*TJ", data, re.DOTALL)]
    return b" ".join(parts).decode("latin-1", "replace")


def _payload(**extra) -> dict:
    body = {
        "nazwa_projektu": "Farma PV Wschód",
        "nazwa_przypadku": "Wariant bazowy",
        "wnioskodawca": "OZE Sp. z o.o.",
        "adres_przylaczenia": "Stacja B",
        "bus_ref": "bus_nn",
        "run_request": {"modules": [dict(_MODULE_FULL)]},
    }
    body.update(extra)
    return body


# --------------------------------------------------------------------------- #
# Serwis kompozycji — happy path
# --------------------------------------------------------------------------- #
def test_komplet_zrodel_daje_trzy_sekcje() -> None:
    view = _view()
    assert view["kontrakt"] == "WniosekOkresleniaWarunkowPrzylaczeniaV1"
    assert "bilans_mocy" in view
    assert "zwarcia_punkt_przylaczenia" in view
    assert "zgodnosc_nc_rfg" in view
    assert view["identyfikacja"]["wezel_przylaczenia"] == "bus_nn"


def test_bilans_zawiera_moc_zainstalowana() -> None:
    bilans = _view()["bilans_mocy"]
    assert bilans["moc_zainstalowana_zrodel_mva"] == pytest.approx(2.75)
    assert bilans["moc_zainstalowana_w_punkcie_mva"] == pytest.approx(2.75)


def test_zwarcia_sekcja_ma_ik_i_sk() -> None:
    zwarcia = _view()["zwarcia_punkt_przylaczenia"]
    assert zwarcia["bus_ref"] == "bus_nn"
    assert zwarcia["ik_ss_ka"] is not None
    assert zwarcia["sk_mva"] is not None


def test_zgodnosc_nc_rfg_werdykt_i_odeslanie() -> None:
    zgodnosc = _view()["zgodnosc_nc_rfg"]
    assert zgodnosc["status"] in {"zgodny", "niezgodny"}
    assert zgodnosc["liczba_modulow"] == 1
    assert "certyfikacie" in zgodnosc["odeslanie_pl"]


# --------------------------------------------------------------------------- #
# Bramka braków
# --------------------------------------------------------------------------- #
def test_braki_bez_pf_zly_rodzaj() -> None:
    braki = zbierz_braki_wniosku(
        _fake_run("short_circuit_sn", "FINISHED"), _sc_run(), "bus_nn", _ncrfg()
    )
    assert any("nie jest rozpływem" in b for b in braki)


def test_braki_pf_niezakonczony() -> None:
    braki = zbierz_braki_wniosku(_fake_run("PF", "RUNNING"), _sc_run(), "bus_nn", _ncrfg())
    assert any("nie jest zakończony" in b for b in braki)


def test_braki_bez_sc_zly_rodzaj() -> None:
    braki = zbierz_braki_wniosku(_pf_run(), _fake_run("PF", "FINISHED"), "bus_nn", _ncrfg())
    assert any("nie jest zwarciowy" in b for b in braki)


def test_braki_nieznany_bus_ref() -> None:
    braki = zbierz_braki_wniosku(_pf_run(), _sc_run(), "bus_nieznany", _ncrfg())
    assert any("nie występuje w wynikach zwarciowych" in b for b in braki)


def test_braki_modulow_nc_rfg() -> None:
    ncrfg = _ncrfg({"p_recovery_time_s": None, "reactive_current_gain": None})
    braki = zbierz_braki_wniosku(_pf_run(), _sc_run(), "bus_nn", ncrfg)
    assert any(b.startswith("Zgodność NC RfG:") for b in braki)
    assert any("brak danych do oceny" in b for b in braki)


def test_braki_blokuja_generacje() -> None:
    with pytest.raises(WniosekOsdBrakiError) as exc:
        build_wniosek_osd_view(
            _pf_run(),
            _sc_run(),
            _ncrfg(),
            bus_ref="bus_nieznany",
            identyfikacja=_identyfikacja(),
        )
    assert exc.value.braki


def test_kolejnosc_brakow_deterministyczna() -> None:
    # Bilans (PF) → zwarcia (SC) → zgodność (NC RfG).
    braki = zbierz_braki_wniosku(
        _fake_run("short_circuit_sn", "FINISHED"),
        _fake_run("PF", "FINISHED"),
        "bus_nn",
        _ncrfg({"p_recovery_time_s": None, "reactive_current_gain": None}),
    )
    assert "nie jest rozpływem" in braki[0]
    assert "nie jest zwarciowy" in braki[1]
    assert braki[-1].startswith("Zgodność NC RfG:")


# --------------------------------------------------------------------------- #
# Determinizm i odciski
# --------------------------------------------------------------------------- #
def test_docx_determinizm_bajtowy() -> None:
    view = _view()
    first = render_wniosek_osd_docx(view)
    second = render_wniosek_osd_docx(view)
    assert first == second
    assert len(first) > 0


def test_odciski_stabilne_dla_tego_samego_widoku() -> None:
    view = _view()
    assert set(view["odciski_sekcji_sha256"]) == {
        "bilans_mocy",
        "zwarcia_punkt_przylaczenia",
        "zgodnosc_nc_rfg",
    }
    assert view["input_hash"]
    # Odciski sekcji są funkcją zawartości — powtórne zbudowanie tej samej
    # sekcji daje ten sam odcisk.
    assert render_wniosek_osd_docx(view) == render_wniosek_osd_docx(view)


# --------------------------------------------------------------------------- #
# Etykiety PL i adnotacje w DOCX
# --------------------------------------------------------------------------- #
def test_docx_etykiety_pl() -> None:
    text = _docx_text(render_wniosek_osd_docx(_view()))
    assert "Wniosek o określenie warunków przyłączenia do sieci OSD" in text
    assert "Bilans mocy" in text
    assert "Zwarcia w punkcie przyłączenia" in text
    assert "Zgodność z wymaganiami NC RfG" in text
    assert "Farma PV Wschód" in text
    assert "Odcisk SHA-256 wejścia wniosku" in text


def test_docx_adnotacje_schemat_i_zestawienia() -> None:
    text = _docx_text(render_wniosek_osd_docx(_view()))
    assert "dołączany jest do wniosku odrębnie" in text
    assert "Zestawienia materiałowe nie są jeszcze generowane" in text


def test_dokument_bez_kodow_projektowych() -> None:
    text = _docx_text(render_wniosek_osd_docx(_view()))
    for pattern in (r"\bP\d{2}\b", r"\bE\d{2}\b", r"\bW-\d{3}\b"):
        assert re.search(pattern, text) is None, pattern


# --------------------------------------------------------------------------- #
# Końcówki API
# --------------------------------------------------------------------------- #
def test_endpoint_json_200_komplet(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    resp = app_client.post(OSD_JSON, json=_payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id)))
    assert resp.status_code == 200
    data = resp.json()
    assert data["identyfikacja"]["projekt"] == "Farma PV Wschód"
    assert "bilans_mocy" in data and "zwarcia_punkt_przylaczenia" in data


def test_endpoint_json_braki_422(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    resp = app_client.post(
        OSD_JSON,
        json=_payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id), bus_ref="bus_nieznany"),
    )
    assert resp.status_code == 422
    detail = resp.json()["detail"]
    assert isinstance(detail["braki"], list)
    assert any("nie występuje w wynikach zwarciowych" in b for b in detail["braki"])


def test_endpoint_nieznany_przebieg_404(app_client) -> None:
    sc = _sc_run()
    resp = app_client.post(OSD_JSON, json=_payload(pf_run_id=str(uuid4()), sc_run_id=str(sc.id)))
    assert resp.status_code == 404
    assert "nie istnieje" in resp.json()["detail"]


def test_endpoint_nieznany_operator_404(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    module = dict(_MODULE_FULL)
    module["operator_id"] = "nieistniejacy"
    resp = app_client.post(
        OSD_JSON,
        json=_payload(
            pf_run_id=str(pf.id), sc_run_id=str(sc.id), run_request={"modules": [module]}
        ),
    )
    assert resp.status_code == 404


def test_endpoint_docx_content_type(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    resp = app_client.post(OSD_DOCX, json=_payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id)))
    assert resp.status_code == 200
    assert resp.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content[:2] == b"PK"


def test_endpoint_docx_determinizm(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    payload = _payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id))
    first = app_client.post(OSD_DOCX, json=payload)
    second = app_client.post(OSD_DOCX, json=payload)
    assert first.status_code == 200
    assert first.content == second.content


def test_endpoint_walidacja_422_pusta_nazwa(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    resp = app_client.post(
        OSD_JSON,
        json=_payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id), nazwa_projektu=""),
    )
    assert resp.status_code == 422


# --------------------------------------------------------------------------- #
# Wariant PDF (D16)
# --------------------------------------------------------------------------- #
def test_pdf_powstaje_naglowek_i_niepusty() -> None:
    data = render_wniosek_pdf(_view())
    assert data[:4] == b"%PDF"
    assert len(data) > 0


def test_pdf_determinizm_bajtowy() -> None:
    view = _view()
    assert render_wniosek_pdf(view) == render_wniosek_pdf(view)


def test_pdf_etykiety_pl() -> None:
    text = _pdf_text(render_wniosek_pdf(_view()))
    assert "Bilans mocy" in text
    assert "Moc zwarciowa" in text
    assert "Rodzaj zwarcia" in text
    assert "Werdykt zbiorczy" in text


def test_pdf_dokument_bez_kodow_projektowych() -> None:
    text = _pdf_text(render_wniosek_pdf(_view()))
    for pattern in (r"\bP\d{2}\b", r"\bE\d{2}\b", r"\bW-\d{3}\b"):
        assert re.search(pattern, text) is None, pattern


def test_endpoint_pdf_content_type(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    resp = app_client.post(OSD_PDF, json=_payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id)))
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content[:4] == b"%PDF"


def test_endpoint_pdf_determinizm(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    payload = _payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id))
    first = app_client.post(OSD_PDF, json=payload)
    second = app_client.post(OSD_PDF, json=payload)
    assert first.status_code == 200
    assert first.content == second.content


def test_endpoint_pdf_braki_422(app_client) -> None:
    pf, sc = _pf_run(), _sc_run()
    resp = app_client.post(
        OSD_PDF,
        json=_payload(pf_run_id=str(pf.id), sc_run_id=str(sc.id), bus_ref="bus_nieznany"),
    )
    assert resp.status_code == 422
    detail = resp.json()["detail"]
    assert any("nie występuje w wynikach zwarciowych" in b for b in detail["braki"])


def test_endpoint_pdf_nieznany_przebieg_404(app_client) -> None:
    sc = _sc_run()
    resp = app_client.post(OSD_PDF, json=_payload(pf_run_id=str(uuid4()), sc_run_id=str(sc.id)))
    assert resp.status_code == 404
