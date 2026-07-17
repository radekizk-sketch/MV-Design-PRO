"""Testy certyfikatu zgodności projektu NC RfG (serwis kompozycji + końcówki).

Certyfikat to czysta kompozycja gotowych werdyktów macierzy NC RfG/PTPiREE
(``NcRfgPtpireeSolver``). Testy pokrywają: certyfikat pozytywny i negatywny,
bramkę braków (lista PL blokuje generację), determinizm bajtowy DOCX, stabilność
odcisku wejścia, etykiety PL w dokumencie, brak kodów projektowych, 404/422,
content-type.
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

import pytest
from application.analyses.certyfikat_zgodnosci import (
    CertyfikatBrakiError,
    build_certyfikat_view,
    render_certyfikat_docx,
    zbierz_braki,
)
from docx import Document
from fastapi.testclient import TestClient
from network_model.solvers.ncrfg_ptpiree import (
    NcRfgPtpireeModuleInput,
    NcRfgPtpireeRunRequest,
    NcRfgPtpireeSolver,
)

_MODULE_FULL: dict = {
    "der_ref": "pv-1",
    "der_name": "PV 2 MW",
    "der_kind": "PV",
    "operator_id": "enea",
    "p_max_kw": 2000,
    "p_min_kw": 100,
    "voltage_kv": 15,
    "certificate_status": "ptpiree_verified",
    "has_lvrt_curve": True,
    "has_hvrt_curve": True,
    "has_pf_droop": True,
    "has_qu_curve": True,
    "has_dynamic_model": True,
    "has_scada_communication": True,
    "has_disturbance_recorder": True,
    "active_power_control_enabled": True,
    "droop_percent": 5,
    "dead_band_hz": 0.2,
    "ramp_rate_pct_per_min": 10,
    "cos_phi_min": 0.95,
    "q_range_pct_pn_min": -0.33,
    "q_range_pct_pn_max": 0.33,
    "reactive_current_gain": 2,
    "p_recovery_time_s": 0.8,
    "harmonic_thdu_percent": 3,
}


def _module(**overrides: object) -> dict:
    data = dict(_MODULE_FULL)
    data.update(overrides)
    return data


def _run_result(module: dict):
    solver = NcRfgPtpireeSolver()
    return solver.run(NcRfgPtpireeRunRequest(modules=[NcRfgPtpireeModuleInput(**module)]))


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.fixture
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    db_path = tmp_path / "certyfikat-api.db"
    monkeypatch.setenv("DATABASE_URL", f"sqlite+pysqlite:///{db_path}")
    from api.main import app

    with TestClient(app) as test_client:
        yield test_client


def _payload(module: dict, **extra: object) -> dict:
    body = {
        "run_request": {"modules": [module]},
        "nazwa_projektu": "Farma PV Wschód",
        "nazwa_przypadku": "Wariant bazowy",
    }
    body.update(extra)
    return body


# --------------------------------------------------------------------------- #
# Serwis kompozycji
# --------------------------------------------------------------------------- #
def test_certyfikat_pozytywny_werdykt_zgodny() -> None:
    view = build_certyfikat_view(_run_result(_module()), nazwa_projektu="Projekt A")
    assert view["kontrakt"] == "CertyfikatZgodnosciNcRfgV1"
    assert view["werdykt_zbiorczy"]["status"] == "zgodny"
    assert view["werdykt_zbiorczy"]["modulow_niezgodnych"] == 0
    assert view["moduly"][0]["klasa"] == "B"
    assert view["odcisk_wejscia_sha256"]


def test_certyfikat_negatywny_powstaje_z_werdyktem_niezgodnym() -> None:
    view = build_certyfikat_view(
        _run_result(_module(reactive_current_gain=1.0)),
        nazwa_projektu="Projekt A",
    )
    assert view["werdykt_zbiorczy"]["status"] == "niezgodny"
    assert view["werdykt_zbiorczy"]["modulow_niezgodnych"] == 1
    assert view["moduly"][0]["status_pl"] == "Niezgodny"


def test_braki_blokuja_generacje_lista_pl() -> None:
    run_result = _run_result(_module(p_recovery_time_s=None, reactive_current_gain=None))
    with pytest.raises(CertyfikatBrakiError) as exc:
        build_certyfikat_view(run_result, nazwa_projektu="Projekt A")
    braki = exc.value.braki
    assert braki
    assert any("brak danych do oceny" in b for b in braki)
    assert any("T16" in b for b in braki)
    assert any("T17" in b for b in braki)


def test_zbierz_braki_pusta_lista_gdy_komplet() -> None:
    assert zbierz_braki(_run_result(_module())) == []


def test_docx_determinizm_bajtowy() -> None:
    view = build_certyfikat_view(_run_result(_module()), nazwa_projektu="Projekt A")
    first = render_certyfikat_docx(view)
    second = render_certyfikat_docx(view)
    assert first == second
    assert len(first) > 0


def test_input_hash_stabilny_dla_tego_samego_wejscia() -> None:
    view_a = build_certyfikat_view(_run_result(_module()), nazwa_projektu="Projekt A")
    view_b = build_certyfikat_view(_run_result(_module()), nazwa_projektu="Projekt A")
    assert view_a["odcisk_wejscia_sha256"] == view_b["odcisk_wejscia_sha256"]
    assert view_a == view_b


def test_docx_naglowki_i_etykiety_pl() -> None:
    view = build_certyfikat_view(
        _run_result(_module()),
        nazwa_projektu="Farma PV",
        nazwa_przypadku="Wariant bazowy",
    )
    text = _docx_text(render_certyfikat_docx(view))
    assert "Certyfikat zgodności projektu z wymaganiami NC RfG" in text
    assert "Werdykt zbiorczy" in text
    assert "Farma PV" in text
    assert "Założenia i źródła" in text
    assert "Odcisk SHA-256 wejścia" in text
    assert "spełnia" in text


def test_dokument_bez_kodow_projektowych() -> None:
    view = build_certyfikat_view(_run_result(_module()), nazwa_projektu="Projekt A")
    text = _docx_text(render_certyfikat_docx(view))
    for pattern in (r"\bP\d{2}\b", r"\bE\d{2}\b", r"\bW-\d{3}\b"):
        assert re.search(pattern, text) is None, pattern


# --------------------------------------------------------------------------- #
# Końcówki API
# --------------------------------------------------------------------------- #
def test_endpoint_json_200_zgodny(client: TestClient) -> None:
    response = client.post("/api/oze-analysis/compliance-certificate", json=_payload(_module()))
    assert response.status_code == 200
    payload = response.json()
    assert payload["werdykt_zbiorczy"]["status"] == "zgodny"
    assert payload["identyfikacja"]["projekt"] == "Farma PV Wschód"


def test_endpoint_json_negatywny_200(client: TestClient) -> None:
    response = client.post(
        "/api/oze-analysis/compliance-certificate",
        json=_payload(_module(reactive_current_gain=1.0)),
    )
    assert response.status_code == 200
    assert response.json()["werdykt_zbiorczy"]["status"] == "niezgodny"


def test_endpoint_braki_422_z_lista(client: TestClient) -> None:
    response = client.post(
        "/api/oze-analysis/compliance-certificate",
        json=_payload(_module(p_recovery_time_s=None, reactive_current_gain=None)),
    )
    assert response.status_code == 422
    detail = response.json()["detail"]
    assert isinstance(detail["braki"], list)
    assert any("brak danych do oceny" in b for b in detail["braki"])


def test_endpoint_nieznany_operator_404(client: TestClient) -> None:
    response = client.post(
        "/api/oze-analysis/compliance-certificate",
        json=_payload(_module(operator_id="nieistniejacy")),
    )
    assert response.status_code == 404
    assert "nie istnieje" in response.json()["detail"]


def test_endpoint_walidacja_422_pusta_nazwa(client: TestClient) -> None:
    response = client.post(
        "/api/oze-analysis/compliance-certificate",
        json=_payload(_module(), nazwa_projektu=""),
    )
    assert response.status_code == 422


def test_endpoint_docx_content_type(client: TestClient) -> None:
    response = client.post(
        "/api/oze-analysis/compliance-certificate.docx", json=_payload(_module())
    )
    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in response.headers["content-disposition"]
    assert response.content[:2] == b"PK"


def test_endpoint_docx_determinizm(client: TestClient) -> None:
    first = client.post("/api/oze-analysis/compliance-certificate.docx", json=_payload(_module()))
    second = client.post("/api/oze-analysis/compliance-certificate.docx", json=_payload(_module()))
    assert first.status_code == 200
    assert first.content == second.content


def test_endpoint_docx_braki_422(client: TestClient) -> None:
    response = client.post(
        "/api/oze-analysis/compliance-certificate.docx",
        json=_payload(_module(p_recovery_time_s=None, reactive_current_gain=None)),
    )
    assert response.status_code == 422
    assert "braki" in response.json()["detail"]
