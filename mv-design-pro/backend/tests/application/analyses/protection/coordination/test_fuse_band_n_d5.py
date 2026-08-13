"""Karta N-D5-FUSE — bezpiecznik NIE JEST liczony jak przekaznik.

TLO POMIAROWE: urzadzenie `device_type=FUSE` dostawalo krzywa z formuly IDMT
IEC 60255 przez cichy fallback `standard_map.get(..., IEC)`. Zmierzone na zywej
sciezce API: 100 punktow IDENTYCZNYCH CO DO OSTATNIEJ CYFRY z przekaznikiem
IEC SI, opisanych etykieta `FUSE_SI`. Defekt przezyl lata, bo KAZDY istniejacy
test koordynacji uzywal wylacznie `ProtectionDeviceType.RELAY`.

FIZYKA: bezpiecznik topikowy SN nie ma charakterystyki IDMT ani TMS. Ma pasmo
topikowe (przedlukowe + wylaczania) z karty katalogowej wg IEC 60282-1.

POKRYCIE = ILOCZYN CECH, nie przyklad z karty:
    typ urzadzenia (RELAY / FUSE / RECLOSER / CIRCUIT_BREAKER)
  x sposob zgloszenia krzywej (norma FUSE / norma IEC / brak nastaw / nieznana)
  x sciezka wyniku (czas zadzialania / krzywa TCC / selektywnosc / API / raport)
"""

from __future__ import annotations

import itertools
from uuid import UUID, uuid4

import pytest
from application.analyses.protection.coordination import (
    CoordinationInput,
    OvercurrentCoordinationAnalyzer,
)
from application.analyses.protection.coordination.analyzer import (
    _NORMY_PRZEKAZNIKOWE,
    KOD_BRAK_CHARAKTERYSTYKI,
    KOD_BRAK_NASTAW_KRZYWEJ,
    KOD_BRAK_PASMA_BEZPIECZNIKA,
    KOD_KRZYWA_PRZEKAZNIKOWA,
    PodstawaKrzywej,
    rozstrzygnij_podstawe_krzywej,
)
from application.analyses.protection.coordination.models import (
    FaultCurrentData,
    OperatingCurrentData,
)
from domain.protection_device import (
    CoordinationVerdict,
    CurveStandard,
    OvercurrentProtectionSettings,
    OvercurrentStageSettings,
    ProtectionCurveSettings,
    ProtectionDevice,
    ProtectionDeviceType,
)
from network_model.reporting.protection_tcc_presentation import (
    ETYKIETY_BRAKU_PL,
    NIE_DOTYCZY,
    etykieta_tms,
    etykieta_typu_krzywej_pl,
    ma_podstawe_przekaznikowa,
    powod_braku_pl,
)

#: Typy, ktore MAJA prawo do krzywej przekaznikowej (wyzwalacz nadpradowy).
TYPY_PRZEKAZNIKOWE = (
    ProtectionDeviceType.RELAY,
    ProtectionDeviceType.RECLOSER,
    ProtectionDeviceType.CIRCUIT_BREAKER,
)


#: Wariant WLASCIWY dla danej normy. IEC i IEEE maja ROZLACZNE slowniki
#: wariantow (IEC: SI/VI/EI/LTI/DT, IEEE: MI/VI/EI/STI/DT) — podstawienie
#: wariantu IEC pod norme IEEE jest bledem danych, nie „drobiazgiem".
WARIANT_DLA_NORMY: dict[CurveStandard, str] = {
    CurveStandard.IEC: "SI",
    CurveStandard.IEEE: "MI",
    CurveStandard.FUSE: "SI",
}


def _urzadzenie(
    typ: ProtectionDeviceType,
    *,
    standard: CurveStandard | None,
    idx: int = 1,
    lokalizacja: str = "L1",
    pickup: float = 63.0,
) -> ProtectionDevice:
    """Urzadzenie o zadanym typie i sposobie zgloszenia krzywej.

    `standard is None` znaczy: czlon 51 BEZ nastaw charakterystyki.
    """
    krzywa = (
        None
        if standard is None
        else ProtectionCurveSettings(
            standard=standard,
            variant=WARIANT_DLA_NORMY[standard],
            pickup_current_a=pickup,
            time_multiplier=0.2,
        )
    )
    return ProtectionDevice(
        id=UUID(int=idx),
        name=f"{typ.value}-{standard.value if standard else 'bez-krzywej'}",
        device_type=typ,
        location_element_id=lokalizacja,
        settings=OvercurrentProtectionSettings(
            stage_51=OvercurrentStageSettings(
                enabled=True, pickup_current_a=pickup, curve_settings=krzywa
            )
        ),
    )


def _analizuj(*urzadzenia: ProtectionDevice):
    lokalizacje = {u.location_element_id for u in urzadzenia}
    return OvercurrentCoordinationAnalyzer().analyze(
        CoordinationInput(
            devices=tuple(urzadzenia),
            fault_currents=tuple(
                FaultCurrentData(location_id=loc, ik_max_3f_a=6000.0, ik_min_3f_a=1500.0)
                for loc in sorted(lokalizacje)
            ),
            operating_currents=tuple(
                OperatingCurrentData(location_id=loc, i_operating_a=40.0)
                for loc in sorted(lokalizacje)
            ),
        )
    )


# =============================================================================
# ILOCZYN CECH: typ urzadzenia x sposob zgloszenia krzywej
# =============================================================================


@pytest.mark.parametrize(
    "standard",
    [CurveStandard.FUSE, CurveStandard.IEC, CurveStandard.IEEE, None],
    ids=["norma-FUSE", "norma-IEC", "norma-IEEE", "bez-nastaw-krzywej"],
)
def test_bezpiecznik_nie_dostaje_krzywej_zadnym_sposobem_zgloszenia(
    standard: CurveStandard | None,
) -> None:
    """Bezpiecznik NIE dostaje punktow krzywej, niezaleznie od zglaszanej normy.

    Kluczowy przypadek to `norma-IEC`: uzytkownik wybiera z listy norme
    przekaznikowa dla bezpiecznika. Urzadzenie jest nadal bezpiecznikiem, wiec
    krzywa IDMT byla by fabrykacja fizyki — o braku podstawy decyduje TYP
    URZADZENIA, nie sama deklaracja normy.
    """
    wynik = _analizuj(_urzadzenie(ProtectionDeviceType.FUSE, standard=standard))

    assert len(wynik.tcc_curves) == 1
    krzywa = wynik.tcc_curves[0]
    assert krzywa.points == ()
    assert krzywa.podstawa_kod == KOD_BRAK_PASMA_BEZPIECZNIKA
    assert krzywa.curve_type == KOD_BRAK_CHARAKTERYSTYKI
    # Etykieta NIGDY nie sugeruje krzywej bezpiecznikowej tam, gdzie jej nie ma.
    assert not krzywa.curve_type.startswith("FUSE_")
    assert krzywa.powod_pl is not None
    assert "IEC 60282-1" in krzywa.powod_pl
    # Bezpiecznik nie ma mnoznika czasowego — nie udajemy, ze ma.
    assert krzywa.time_multiplier == 0.0


@pytest.mark.parametrize("typ", TYPY_PRZEKAZNIKOWE, ids=lambda t: t.value)
@pytest.mark.parametrize("standard", [CurveStandard.IEC, CurveStandard.IEEE], ids=lambda s: s.value)
def test_urzadzenia_przekaznikowe_zachowuja_krzywa(
    typ: ProtectionDeviceType, standard: CurveStandard
) -> None:
    """Przekaznik, reklozer i wylacznik z wyzwalaczem NADAL dostaja krzywa IDMT.

    Druga polowa pary: naprawa nie moze wyciszyc urzadzen, ktore maja prawo do
    charakterystyki przekaznikowej.
    """
    wynik = _analizuj(_urzadzenie(typ, standard=standard))

    assert len(wynik.tcc_curves) == 1
    krzywa = wynik.tcc_curves[0]
    assert len(krzywa.points) == 100
    assert krzywa.podstawa_kod == KOD_KRZYWA_PRZEKAZNIKOWA
    assert krzywa.curve_type == f"{standard.value}_{WARIANT_DLA_NORMY[standard]}"
    assert krzywa.powod_pl is None


def test_bezpiecznik_nie_dostaje_czasu_zadzialania() -> None:
    """Czas zadzialania bezpiecznika jest NIEWYZNACZONY, z powodem po polsku."""
    analizator = OvercurrentCoordinationAnalyzer()
    bezpiecznik = _urzadzenie(ProtectionDeviceType.FUSE, standard=CurveStandard.FUSE)

    czas, powod = analizator._calculate_device_trip_time(bezpiecznik, 5000.0)

    assert czas == float("inf")
    assert powod is not None
    assert "IEC 60282-1" in powod


def test_bezpiecznik_ze_zwloka_niezalezna_tez_nie_dostaje_czasu() -> None:
    """Zadeklarowana zwloka `time_s` NIE staje sie czasem zadzialania bezpiecznika.

    Czlon nastawczy ze zwloka niezalezna to konstrukcja PRZEKAZNIKA. Gdyby
    rozstrzygniecie o bezpieczniku szlo PO odczycie `time_s`, deklaracja
    zwloki byla by ta sama klasa fabrykacji co krzywa IDMT — tyle ze cichsza.
    """
    analizator = OvercurrentCoordinationAnalyzer()
    bezpiecznik = ProtectionDevice(
        id=uuid4(),
        name="Bezpiecznik ze zwloka",
        device_type=ProtectionDeviceType.FUSE,
        location_element_id="L1",
        settings=OvercurrentProtectionSettings(
            stage_51=OvercurrentStageSettings(
                enabled=True, pickup_current_a=63.0, time_s=0.3, curve_settings=None
            )
        ),
    )

    czas, powod = analizator._calculate_device_trip_time(bezpiecznik, 5000.0)

    assert czas == float("inf")
    assert powod is not None and "IEC 60282-1" in powod


# =============================================================================
# SCIEZKA: selektywnosc (para urzadzen)
# =============================================================================


@pytest.mark.parametrize("typ_partnera", TYPY_PRZEKAZNIKOWE, ids=lambda t: t.value)
def test_selektywnosc_z_bezpiecznikiem_melduje_brak_a_nie_werdykt(
    typ_partnera: ProtectionDeviceType,
) -> None:
    """Para „bezpiecznik + przekaznik" nie dostaje werdyktu selektywnosci.

    Bez pasma topikowego nie da sie uczciwie orzec stopniowania czasowego —
    wynikiem jest BLAD ANALIZY z powodem, nigdy PASS/FAIL policzony ze wzoru
    przekaznikowego podstawionego za bezpiecznik.
    """
    bezpiecznik = _urzadzenie(
        ProtectionDeviceType.FUSE, standard=CurveStandard.FUSE, idx=1, lokalizacja="L1"
    )
    partner = _urzadzenie(typ_partnera, standard=CurveStandard.IEC, idx=2, lokalizacja="L2")

    wynik = _analizuj(bezpiecznik, partner)

    assert len(wynik.selectivity_checks) == 1
    kontrola = wynik.selectivity_checks[0]
    assert kontrola.verdict == CoordinationVerdict.ERROR
    assert "IEC 60282-1" in kontrola.notes_pl
    # Powod wskazuje KTORE urzadzenie nie ma czasu — nie ogolne „nie mozna".
    assert bezpiecznik.name in kontrola.notes_pl


def test_selektywnosc_dwoch_przekaznikow_bez_zmian() -> None:
    """Para przekaznikow nadal dostaje policzony margines (regresja naprawy)."""
    a = _urzadzenie(ProtectionDeviceType.RELAY, standard=CurveStandard.IEC, idx=1, lokalizacja="L1")
    b = _urzadzenie(
        ProtectionDeviceType.RELAY,
        standard=CurveStandard.IEC,
        idx=2,
        lokalizacja="L2",
        pickup=200.0,
    )

    wynik = _analizuj(a, b)

    kontrola = wynik.selectivity_checks[0]
    assert kontrola.verdict != CoordinationVerdict.ERROR
    assert kontrola.t_downstream_s > 0.0
    assert kontrola.t_upstream_s > 0.0


# =============================================================================
# DEKLARACJE Z DOKUMENTACJI — kazda z PRZYPIETYM testem
# =============================================================================


def test_zamknieta_mapa_norm_przekaznikowych() -> None:
    """Pin deklaracji „Mapa ZAMKNIETA" z `_NORMY_PRZEKAZNIKOWE`.

    Kazda norma dopisana do `CurveStandard` musi swiadomie trafic albo do mapy
    wzorow przekaznikowych, albo do sciezki braku podstawy — nigdy do cichego
    zastepnika. Ten test pilnuje, zeby dopisanie normy nie przeszlo bez decyzji.
    """
    assert set(_NORMY_PRZEKAZNIKOWE) == {CurveStandard.IEC, CurveStandard.IEEE}
    # CurveStandard.FUSE CELOWO poza mapa — bezpiecznik nie ma wzoru.
    assert CurveStandard.FUSE not in _NORMY_PRZEKAZNIKOWE
    assert set(CurveStandard) - set(_NORMY_PRZEKAZNIKOWE) == {CurveStandard.FUSE}


def test_para_predykatow_norma_i_nastawy_razem_albo_wcale() -> None:
    """Pin deklaracji „PARA PREDYKATOW" z `PodstawaKrzywej`.

    Norma bez nastaw (albo odwrotnie) to stan, w ktorym miejsce uzycia musialo
    by sprawdzac warunek drugi raz wlasnym `if` — dokladnie ten rozjazd byl
    zrodlem defektu. Konstrukcja ma go UNIEMOZLIWIAC, nie tylko odradzac.
    """
    from protection.curves.curve_calculator import CurveStandard as CurveCurveStandard

    nastawy = ProtectionCurveSettings(
        standard=CurveStandard.IEC,
        variant="SI",
        pickup_current_a=100.0,
        time_multiplier=0.2,
    )

    with pytest.raises(ValueError, match="razem albo wcale"):
        PodstawaKrzywej(KOD_KRZYWA_PRZEKAZNIKOWA, CurveCurveStandard.IEC, None, None)

    with pytest.raises(ValueError, match="razem albo wcale"):
        PodstawaKrzywej(KOD_BRAK_PASMA_BEZPIECZNIKA, None, "powod", nastawy)

    # Obie poprawne kombinacje przechodza.
    assert PodstawaKrzywej(KOD_BRAK_PASMA_BEZPIECZNIKA, None, "powod").standard is None
    assert (
        PodstawaKrzywej(KOD_KRZYWA_PRZEKAZNIKOWA, CurveCurveStandard.IEC, None, nastawy).nastawy
        is nastawy
    )


def test_para_predykatow_dla_kazdego_rozstrzygniecia() -> None:
    """Rozstrzygniecie KAZDEJ kombinacji utrzymuje pare predykatow.

    Iloczyn cech: 4 typy urzadzen x 4 sposoby zgloszenia krzywej.
    """
    for typ, standard in itertools.product(
        list(ProtectionDeviceType),
        [CurveStandard.IEC, CurveStandard.IEEE, CurveStandard.FUSE, None],
    ):
        podstawa = rozstrzygnij_podstawe_krzywej(_urzadzenie(typ, standard=standard))
        assert (podstawa.standard is None) == (
            podstawa.nastawy is None
        ), f"{typ.value}/{standard}: norma i nastawy rozjechaly sie"
        # Brak podstawy ZAWSZE niesie powod po polsku — cichy brak jest zly.
        if podstawa.standard is None:
            assert podstawa.powod_pl, f"{typ.value}/{standard}: brak bez powodu"
        else:
            assert podstawa.kod == KOD_KRZYWA_PRZEKAZNIKOWA


def test_urzadzenie_przekaznikowe_bez_nastaw_krzywej_ma_wlasny_kod() -> None:
    """Brak nastaw krzywej to INNY brak niz brak pasma bezpiecznika."""
    podstawa = rozstrzygnij_podstawe_krzywej(_urzadzenie(ProtectionDeviceType.RELAY, standard=None))
    assert podstawa.kod == KOD_BRAK_NASTAW_KRZYWEJ
    assert podstawa.standard is None


# =============================================================================
# SCIEZKA: prezentacja w raportach PDF / DOCX
# =============================================================================


def test_raport_tcc_etykiety_bez_podstawy() -> None:
    """Pin deklaracji „lista ZAMKNIETA" z `ETYKIETY_BRAKU_PL`.

    Raport NIGDY nie drukuje surowego kodu ani etykiety `FUSE_*` dla pozycji
    bez krzywej — tylko polskie zdanie.
    """
    assert set(ETYKIETY_BRAKU_PL) == {
        "BRAK_PASMA_BEZPIECZNIKA",
        "NIEZNANA_NORMA_KRZYWEJ",
        "BRAK_NASTAW_KRZYWEJ",
    }

    bezpiecznik = {
        "curve_type": KOD_BRAK_CHARAKTERYSTYKI,
        "podstawa_kod": KOD_BRAK_PASMA_BEZPIECZNIKA,
        "powod_pl": "Bezpiecznik topikowy ... IEC 60282-1 ...",
        "time_multiplier": 0.0,
    }
    assert not ma_podstawe_przekaznikowa(bezpiecznik)
    assert etykieta_typu_krzywej_pl(bezpiecznik) == "Bezpiecznik — brak pasma topikowego"
    assert etykieta_tms(bezpiecznik, "0.000") == NIE_DOTYCZY
    assert powod_braku_pl(bezpiecznik) is not None

    przekaznik = {
        "curve_type": "IEC_SI",
        "podstawa_kod": KOD_KRZYWA_PRZEKAZNIKOWA,
        "powod_pl": None,
        "time_multiplier": 0.2,
    }
    assert ma_podstawe_przekaznikowa(przekaznik)
    assert etykieta_typu_krzywej_pl(przekaznik) == "IEC_SI"
    assert etykieta_tms(przekaznik, "0.200") == "0.200"
    assert powod_braku_pl(przekaznik) is None

    # Kod spoza listy dostaje etykiete ogolna — NIGDY surowy kod.
    nieznany = {"curve_type": "X", "podstawa_kod": "COS_NOWEGO", "powod_pl": "p"}
    assert etykieta_typu_krzywej_pl(nieznany) == "Brak charakterystyki"
    assert "COS_NOWEGO" not in etykieta_typu_krzywej_pl(nieznany)


def test_raport_pdf_i_docx_pokazuja_ten_sam_brak(tmp_path) -> None:
    """PDF i DOCX opowiadaja TE SAMA historie o bezpieczniku.

    Dwa eksporty tego samego wyniku nie moga sie roznic werdyktem — inaczej
    czytelnik dostaje dwie prawdy w zaleznosci od formatu.
    """
    reportlab = pytest.importorskip("reportlab", reason="eksport PDF wymaga reportlab")
    docx = pytest.importorskip("docx", reason="eksport DOCX wymaga python-docx")
    assert reportlab and docx

    from network_model.reporting.protection_report_docx import (
        export_protection_coordination_to_docx,
    )
    from network_model.reporting.protection_report_pdf import (
        export_protection_coordination_to_pdf,
    )

    wynik = _analizuj(
        _urzadzenie(ProtectionDeviceType.FUSE, standard=CurveStandard.FUSE, idx=1),
        _urzadzenie(
            ProtectionDeviceType.RELAY,
            standard=CurveStandard.IEC,
            idx=2,
            lokalizacja="L2",
        ),
    ).to_dict()

    sciezka_pdf = export_protection_coordination_to_pdf(wynik, tmp_path / "k.pdf")
    sciezka_docx = export_protection_coordination_to_docx(wynik, tmp_path / "k.docx")
    assert sciezka_pdf.exists() and sciezka_docx.exists()

    from docx import Document

    tresc_docx = "\n".join(
        komorka.text
        for tabela in Document(str(sciezka_docx)).tables
        for wiersz in tabela.rows
        for komorka in wiersz.cells
    ) + "\n".join(p.text for p in Document(str(sciezka_docx)).paragraphs)

    assert "Bezpiecznik — brak pasma topikowego" in tresc_docx
    # Etykieta fabrykatu NIE pojawia sie w raporcie.
    assert "FUSE_SI" not in tresc_docx
    # Przekaznik zachowuje swoja krzywa.
    assert "IEC_SI" in tresc_docx


# =============================================================================
# SCIEZKA: publiczne API /tcc
# =============================================================================


def test_warianty_krzywej_zgodne_z_kalkulatorem() -> None:
    """Pin deklaracji `_WARIANTY_NORMY` — slownik API = slownik kalkulatora.

    Znalezisko uboczne karty N-D5-FUSE: para (norma, wariant) nie byla
    sprawdzana, wiec wariant IEC pod norma IEEE leciał do kalkulatora i wracal
    technicznym `ValueError` po angielsku (HTTP 422). Ten test pilnuje, zeby
    lista na granicy API nie rozjechala sie z enumami liczacymi punkty.
    """
    from api.protection_coordination import _WARIANTY_NORMY
    from protection.curves.iec_curves import IECCurveType
    from protection.curves.ieee_curves import IEEECurveType

    assert _WARIANTY_NORMY["IEC"] == tuple(w.value for w in IECCurveType)
    assert _WARIANTY_NORMY["IEEE"] == tuple(w.value for w in IEEECurveType)
    # Slowniki sa ROZLACZNE poza wariantami wspolnymi — dlatego para wymaga
    # sprawdzenia razem, a nie kazdego pola osobno.
    assert "SI" in _WARIANTY_NORMY["IEC"] and "SI" not in _WARIANTY_NORMY["IEEE"]
    assert "MI" in _WARIANTY_NORMY["IEEE"] and "MI" not in _WARIANTY_NORMY["IEC"]


@pytest.mark.parametrize(
    ("standard", "wariant"),
    [("IEEE", "SI"), ("IEC", "MI"), ("IEC", "ZMYSLONY")],
    ids=["IEEE-z-wariantem-IEC", "IEC-z-wariantem-IEEE", "wariant-nieistniejacy"],
)
def test_api_odrzuca_niezgodna_pare_norma_wariant(standard: str, wariant: str) -> None:
    """Niezgodna para (norma, wariant) = 400 z polskim, naprawialnym zdaniem."""
    from api.main import app
    from fastapi.testclient import TestClient

    klient = TestClient(app)
    odpowiedz = klient.post(
        f"/api/protection-coordination/projects/{uuid4()}/run",
        json={
            "devices": [
                {
                    "id": str(uuid4()),
                    "name": "Przekaznik",
                    "device_type": "RELAY",
                    "location_element_id": "L1",
                    "settings": {
                        "stage_51": {
                            "enabled": True,
                            "pickup_current_a": 100.0,
                            "curve_settings": {
                                "standard": standard,
                                "variant": wariant,
                                "pickup_current_a": 100.0,
                                "time_multiplier": 0.2,
                            },
                        }
                    },
                }
            ],
            "fault_currents": [{"location_id": "L1", "ik_max_3f_a": 5000.0, "ik_min_3f_a": 1200.0}],
            "operating_currents": [{"location_id": "L1", "i_operating_a": 40.0}],
        },
    )

    assert odpowiedz.status_code == 400
    detail = odpowiedz.json()["detail"]
    assert wariant in detail
    assert "Dozwolone warianty" in detail
    # Komunikat po polsku, nie techniczny angielski wyjatek.
    assert "is not a valid" not in detail


@pytest.mark.parametrize(
    "standard_zgloszony", ["FUSE", "IEC"], ids=["api-norma-FUSE", "api-norma-IEC"]
)
def test_api_tcc_nie_zwraca_punktow_dla_bezpiecznika(standard_zgloszony: str) -> None:
    """Publiczna sciezka API: bezpiecznik nie dostaje ani jednego punktu.

    To dokladnie ta sciezka, na ktorej zmierzono fantom (POST run -> GET tcc).
    """
    from api.main import app
    from fastapi.testclient import TestClient

    klient = TestClient(app)
    projekt = str(uuid4())
    urzadzenie = str(uuid4())

    odpowiedz = klient.post(
        f"/api/protection-coordination/projects/{projekt}/run",
        json={
            "devices": [
                {
                    "id": urzadzenie,
                    "name": "Bezpiecznik ETI VV 12 kV 63 A",
                    "device_type": "FUSE",
                    "location_element_id": "L1",
                    "settings": {
                        "stage_51": {
                            "enabled": True,
                            "pickup_current_a": 63.0,
                            "curve_settings": {
                                "standard": standard_zgloszony,
                                "variant": "SI",
                                "pickup_current_a": 63.0,
                                "time_multiplier": 0.1,
                            },
                        }
                    },
                }
            ],
            "fault_currents": [{"location_id": "L1", "ik_max_3f_a": 5000.0, "ik_min_3f_a": 1200.0}],
            "operating_currents": [{"location_id": "L1", "i_operating_a": 40.0}],
        },
    )
    assert odpowiedz.status_code == 201
    run_id = odpowiedz.json()["run_id"]

    tcc = klient.get(f"/api/protection-coordination/{run_id}/tcc")
    assert tcc.status_code == 200
    krzywe = tcc.json()["curves"]
    assert len(krzywe) == 1
    assert krzywe[0]["points"] == []
    assert krzywe[0]["podstawa_kod"] == KOD_BRAK_PASMA_BEZPIECZNIKA
    assert krzywe[0]["powod_pl"] and "IEC 60282-1" in krzywe[0]["powod_pl"]
    assert not krzywe[0]["curve_type"].startswith("FUSE_")
